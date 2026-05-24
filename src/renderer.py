"""Deterministic Word renderer.

Consumes:

* ``converted_template.docx`` — written by ``learn`` (template_builder).
  Carries ``{{ word_NNNN }}`` placeholders that a future renderer is
  contracted to honor against a human-confirmed mapping.
* ``run_validation.xlsx`` — written by ``run-preview``. One row per
  confirmed ``word_id`` with the deterministic generated value, raw
  Excel value, source sheet/cell, original Word Raw Token / Word Unit,
  and per-row status.

Produces (only on full success):

* ``new_report.docx`` — every ``{{ word_NNNN }}`` placeholder occurrence
  substituted by deterministic display text derived from the historical
  Word Raw Token shape applied to the run-time generated value.
* ``render_log.yml`` — one entry per ``word_id`` with source sheet/cell,
  raw Excel value, generated value, display text, occurrence count, and
  status. The full audit trail a reviewer needs to convince themselves
  the rendered docx contains every confirmed metric and nothing else.

Fail-loud gates (the CLI exits non-zero unless every condition holds):

* Every ``run_validation`` row has Status=ok.
* Every confirmed ``word_id`` has exactly one Generated Value
  (no duplicate rows, no null Generated Value).
* Every ``{{ word_NNNN }}`` placeholder occurrence in the template has
  a matching ``run_validation`` row.
* No ``run_validation`` ``word_id`` is silently unused — every row maps
  to at least one placeholder in the template.
* Display text can be deterministically inferred from the historical
  Word Raw Token + Word Unit. v1 supports the same numeric shapes the
  matcher already understands: optional accounting parens, optional
  explicit sign, comma-grouped or plain integer body, optional decimal,
  and one of the known unit suffixes (元/万元/亿元/%/‰/…). If a raw
  token doesn't parse cleanly into that pattern, the word_id is
  surfaced and the render halts.

No LLM, no GUI, no network, no Microsoft Word automation, no real data.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Dict, List, Optional

import docx
import openpyxl
import yaml

# Same Jinja-flavored placeholder family ``template_builder`` writes and
# ``artifact_validator`` scans. Compact ``{{word_NNNN}}`` is accepted so a
# manually massaged template still substitutes — the canonical spaced form
# is what learn writes, but a reviewer might collapse the spacing by hand.
PLACEHOLDER_RE = re.compile(r"\{\{\s*(word_\d{4,})\s*\}\}")

# Unit suffixes the deterministic formatter recognises. Mirrors the set in
# ``number_normalizer`` so any token that survived the matcher is also
# formattable. Sorted longest-first so ``万元`` wins over ``万`` when the
# raw token is scanned in one regex pass.
_KNOWN_UNITS = [
    "万元", "亿元", "千元", "百万元",
    "万单", "亿单",
    "万人", "万次", "万个",
    "万",
    "元", "单", "人", "次", "个",
    "%", "‰",
]
_UNIT_ALT = "|".join(sorted(_KNOWN_UNITS, key=len, reverse=True))

# Strict raw-token parser: must consume the WHOLE stripped string. Accepts
# an optional accounting-paren wrap, an optional explicit sign, a
# comma-grouped or plain integer body, an optional decimal, and an
# optional recognised unit suffix. Anything outside this shape — Chinese
# numerals, currency symbols, ranges, mixed text — fails loudly upstream.
_RAW_RE = re.compile(
    r"^"
    r"(?P<paren_open>\()?"
    r"(?P<sign>[-+])?"
    r"(?P<int>\d{1,3}(?:,\d{3})+|\d+)"
    r"(?P<frac>\.(?P<frac_digits>\d+))?"
    r"\s*"
    r"(?P<unit>" + _UNIT_ALT + r")?"
    r"(?P<paren_close>\))?"
    r"$"
)


# Per-row render statuses. ``ok`` is the only success state; everything
# else halts the render and is surfaced in the log + stderr so a reviewer
# can see which row(s) blocked the gate.
STATUS_OK = "ok"
STATUS_VALIDATION_NOT_OK = "validation_status_not_ok"
STATUS_MISSING_GENERATED_VALUE = "missing_generated_value"
STATUS_DUPLICATE_VALIDATION_ROW = "duplicate_validation_row"
STATUS_NO_PLACEHOLDER = "no_matching_placeholder_in_template"
STATUS_FORMAT_INFERENCE_FAILED = "format_inference_failed"


class FormatInferenceError(ValueError):
    """Raised when a historical raw token can't be parsed into the v1 shape."""


@dataclass(frozen=True)
class RawTokenPattern:
    """Deterministic display pattern extracted from a historical raw token."""
    unit: str
    has_grouping: bool
    decimals: int
    sign_style: str  # "paren" | "explicit_plus" | "explicit_minus" | "none"


@dataclass
class RenderRow:
    """One resolved (or failed-to-resolve) run_validation entry."""
    word_id: str
    word_location: str
    word_raw: str
    word_unit: str
    source_sheet: str
    source_cell: str
    raw_excel_value: Optional[float]
    generated_value: Optional[float]
    display_text: Optional[str] = None
    placeholder_occurrences: int = 0
    status: str = STATUS_OK
    detail: str = ""


@dataclass
class RenderReport:
    rows: List[RenderRow] = field(default_factory=list)
    fatal_errors: List[str] = field(default_factory=list)
    template_placeholder_counts: Dict[str, int] = field(default_factory=dict)
    template_path: Optional[Path] = None
    validation_path: Optional[Path] = None
    out_docx_path: Optional[Path] = None
    out_log_path: Optional[Path] = None

    @property
    def failures(self) -> List[RenderRow]:
        return [r for r in self.rows if r.status != STATUS_OK]

    @property
    def ok(self) -> bool:
        return (
            not self.fatal_errors
            and not self.failures
            and bool(self.rows)
            and self.out_docx_path is not None
        )


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def render_docx(
    template_path: Path,
    run_validation_path: Path,
    out_docx_path: Path,
) -> RenderReport:
    """Render the template using deterministic display text from the validation.

    Writes ``out_docx_path`` and a sibling ``render_log.yml`` ONLY when
    every gate passes — the CLI decides how to translate failures to
    exit codes. Pre-condition: both inputs exist (the CLI guards that
    for friendlier error messages).

    Stale-output guard: any prior ``new_report.docx``/``render_log.yml``
    at the target paths is removed BEFORE any work runs. Leaving a
    previous successful render's artifacts in place after a failed gate
    would mis-advertise the failure as a success to anyone (human or
    downstream tool) who only checks for file existence — exactly the
    "looks successful while silently omitting metrics" risk CLAUDE.md
    forbids. The success path writes fresh artifacts below.
    """
    report = RenderReport(
        template_path=template_path,
        validation_path=run_validation_path,
    )

    # Resolve all paths to canonical absolute form up front so the
    # cleanup and the collision check share one source of truth (no
    # tunneling through symlinks or relative-vs-absolute spellings).
    try:
        template_abs = template_path.resolve()
        validation_abs = run_validation_path.resolve()
        out_docx_abs = out_docx_path.resolve()
        log_abs = (out_docx_path.parent / "render_log.yml").resolve()
    except OSError as exc:
        report.fatal_errors.append(
            f"cannot resolve --out path {out_docx_path}: {exc}"
        )
        return report

    # Stale-output cleanup runs UNCONDITIONALLY against any non-input
    # target path — even when the collision check below is about to
    # fail this run. Skipping cleanup on collision would let a prior
    # successful render's leftovers at the un-colliding output target
    # survive an exit-8 failure and mis-advertise it as success to
    # anyone (human or downstream tool) reading the directory. The
    # ``protected`` set keeps the cleanup from wiping a path that
    # happens to resolve to an input.
    _clean_stale_outputs(
        targets=(out_docx_abs, log_abs),
        protected=frozenset((template_abs, validation_abs)),
    )

    # Refuse path collisions: an ``--out`` (or the sibling
    # ``render_log.yml``) that resolves to a renderer input cannot
    # safely be written into — the cleanup above already protected the
    # input, but going further would either crash or silently mis-write
    # the wrong file. Fatal exit 8.
    _refuse_input_output_collisions(
        template_abs=template_abs,
        validation_abs=validation_abs,
        out_docx_abs=out_docx_abs,
        log_abs=log_abs,
        out_docx_path=out_docx_path,
        report=report,
    )
    if report.fatal_errors:
        return report

    # --- Read inputs ----------------------------------------------------
    try:
        raw_rows = _read_run_validation(run_validation_path)
    except Exception as exc:
        report.fatal_errors.append(
            f"cannot read {run_validation_path.name}: {exc}"
        )
        return report
    if not raw_rows:
        report.fatal_errors.append(
            f"{run_validation_path.name} has no data rows — nothing to render"
        )
        return report

    try:
        placeholder_counts = _read_template_placeholders(template_path)
    except Exception as exc:
        report.fatal_errors.append(
            f"cannot read {template_path.name}: {exc}"
        )
        return report
    report.template_placeholder_counts = dict(placeholder_counts)
    if not placeholder_counts:
        report.fatal_errors.append(
            f"{template_path.name} contains no {{{{ word_NNNN }}}} "
            "placeholders — nothing to render"
        )
        return report

    # --- Build per-row state from the validation rows ------------------
    seen_ids: Dict[str, int] = {}
    for d in raw_rows:
        row = _row_from_validation_dict(d)
        report.rows.append(row)
        if row.word_id:
            seen_ids[row.word_id] = seen_ids.get(row.word_id, 0) + 1

    # --- Gate checks ----------------------------------------------------
    _apply_gate_checks(report, placeholder_counts, seen_ids)
    if report.fatal_errors or report.failures:
        return report

    # --- Format inference ----------------------------------------------
    for row in report.rows:
        try:
            row.display_text = format_display(
                raw=row.word_raw,
                unit=row.word_unit,
                value=row.generated_value or 0.0,
            )
        except FormatInferenceError as exc:
            row.status = STATUS_FORMAT_INFERENCE_FAILED
            row.detail = str(exc)

    if report.failures:
        return report

    # --- Substitute and write -------------------------------------------
    out_dir = out_docx_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    replacement_counts = _substitute_placeholders(
        template_path=template_path,
        rows=report.rows,
        out_path=out_docx_path,
    )
    for row in report.rows:
        row.placeholder_occurrences = replacement_counts.get(row.word_id, 0)
    report.out_docx_path = out_docx_path

    log_path = out_dir / "render_log.yml"
    _write_log(report, log_path)
    report.out_log_path = log_path
    return report


# ---------------------------------------------------------------------------
# Gate checks
# ---------------------------------------------------------------------------

def _apply_gate_checks(
    report: RenderReport,
    placeholder_counts: Dict[str, int],
    seen_ids: Dict[str, int],
) -> None:
    placeholder_ids = set(placeholder_counts.keys())
    validation_ids = set(seen_ids.keys())

    # Gate: every ``run_validation`` ``word_id`` appears exactly once.
    # Marks every occurrence of a duplicated word_id so the reviewer sees
    # both rows in the log rather than one silent winner.
    for row in report.rows:
        if row.status != STATUS_OK:
            continue
        if row.word_id and seen_ids.get(row.word_id, 0) > 1:
            row.status = STATUS_DUPLICATE_VALIDATION_ROW
            row.detail = (
                f"validation row {row.word_id} appears "
                f"{seen_ids[row.word_id]} times — every confirmed "
                "word_id must have exactly one Generated Value"
            )

    # Gate: no extra ``run_validation`` ``word_id`` silently unused.
    for row in report.rows:
        if row.status != STATUS_OK:
            continue
        if row.word_id and row.word_id not in placeholder_ids:
            row.status = STATUS_NO_PLACEHOLDER
            row.detail = (
                f"validation row {row.word_id} has no matching "
                f"{{{{ {row.word_id} }}}} placeholder in the template"
            )

    # Gate: every template placeholder has a matching validation row.
    # The orphan-placeholder case is a top-level concern (there's no row
    # to mark as failed for it), so surface it as a fatal_error rather
    # than a per-row status.
    orphan_placeholders = sorted(placeholder_ids - validation_ids)
    if orphan_placeholders:
        report.fatal_errors.append(
            "template references word_id(s) absent from run_validation: "
            + ", ".join(orphan_placeholders)
        )


# ---------------------------------------------------------------------------
# Display-text formatter
# ---------------------------------------------------------------------------

def parse_raw_pattern(raw: str, expected_unit: str) -> RawTokenPattern:
    """Extract the v1 display pattern from a historical Word raw token.

    Raises :class:`FormatInferenceError` when the token doesn't parse,
    so the caller can fail loudly with the offending word_id. The unit
    captured from the raw token must match ``expected_unit`` (typically
    the validation row's ``Word Unit`` field) — drift between them
    signals that the validation artifact has been hand-edited and the
    formatter has no safe pattern to apply.
    """
    if raw is None:
        raise FormatInferenceError("raw token is None")
    s = str(raw).strip()
    if not s:
        raise FormatInferenceError("raw token is empty")
    m = _RAW_RE.match(s)
    if not m:
        raise FormatInferenceError(
            f"raw token {raw!r} doesn't match the v1 numeric pattern"
        )

    paren_open = bool(m.group("paren_open"))
    paren_close = bool(m.group("paren_close"))
    if paren_open != paren_close:
        raise FormatInferenceError(
            f"raw token {raw!r} has unbalanced accounting parens"
        )

    sign_char = m.group("sign") or ""
    int_part = m.group("int") or ""
    frac_digits = m.group("frac_digits") or ""
    unit = m.group("unit") or ""

    if paren_open and sign_char:
        raise FormatInferenceError(
            f"raw token {raw!r} mixes accounting parens with an explicit sign"
        )

    expected = (expected_unit or "").strip() if expected_unit is not None else ""
    if unit != expected:
        raise FormatInferenceError(
            f"raw token unit {unit!r} doesn't match validation Word Unit "
            f"{expected!r} — refusing to format"
        )

    if paren_open:
        sign_style = "paren"
    elif sign_char == "-":
        sign_style = "explicit_minus"
    elif sign_char == "+":
        sign_style = "explicit_plus"
    else:
        sign_style = "none"

    return RawTokenPattern(
        unit=unit,
        has_grouping=("," in int_part),
        decimals=len(frac_digits),
        sign_style=sign_style,
    )


def format_display(raw: str, unit: str, value: float) -> str:
    """Format ``value`` so its shape matches the historical ``raw`` token.

    Preserves: percent signs, recognised Chinese units, comma grouping,
    explicit sign, and decimal precision. Uses :data:`ROUND_HALF_UP` so
    rounding matches the convention business reports follow rather than
    Python's banker's-rounding default. Raises
    :class:`FormatInferenceError` if the shape can't be parsed.
    """
    pattern = parse_raw_pattern(raw, unit)
    magnitude = abs(value)
    d = _round_half_up(magnitude, pattern.decimals)
    spec = (
        f",.{pattern.decimals}f" if pattern.has_grouping
        else f".{pattern.decimals}f"
    )
    body = format(d, spec)

    if value < 0:
        body = f"({body})" if pattern.sign_style == "paren" else f"-{body}"
    elif pattern.sign_style == "explicit_plus":
        body = f"+{body}"

    return body + pattern.unit


def _round_half_up(value: float, decimals: int) -> Decimal:
    """Round ``value`` to ``decimals`` places with conventional half-up.

    Python's float-to-string uses banker's rounding (round-half-to-even),
    which would surprise a business reader expecting ``0.5 → 1``. Using
    :class:`Decimal` with :data:`ROUND_HALF_UP` keeps the rendered text
    aligned with how the original report rounds, which the formatter
    must preserve to claim the docx is a faithful re-render.
    """
    if decimals < 0:
        raise ValueError("decimals must be >= 0")
    quant = Decimal(10) ** -decimals
    return Decimal(str(value)).quantize(quant, rounding=ROUND_HALF_UP)


# ---------------------------------------------------------------------------
# Input loaders
# ---------------------------------------------------------------------------

def _read_run_validation(path: Path) -> List[Dict[str, object]]:
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    try:
        ws = wb.active
        rows_iter = ws.iter_rows(values_only=True)
        try:
            header = list(next(rows_iter))
        except StopIteration:
            raise ValueError(f"{path.name} has no header row")
        required = (
            "Word ID", "Word Location", "Word Raw Token", "Word Unit",
            "Source Sheet", "Source Cell", "Raw Excel Value",
            "Generated Value", "Status",
        )
        for req in required:
            if req not in header:
                raise ValueError(
                    f"{path.name} missing required column {req!r}"
                )
        out: List[Dict[str, object]] = []
        for raw in rows_iter:
            if raw is None:
                continue
            if all(c is None or c == "" for c in raw):
                continue
            row_dict = {h: v for h, v in zip(header, raw)}
            if not _str_cell(row_dict.get("Word ID")):
                raise ValueError(
                    f"{path.name} has a row with no Word ID — refusing to "
                    "render against a malformed validation artifact"
                )
            out.append(row_dict)
        return out
    finally:
        wb.close()


def _read_template_placeholders(path: Path) -> Dict[str, int]:
    """Count ``{{ word_NNNN }}`` occurrences in the template, keyed by word_id.

    Walks paragraphs and top-level table cells — the same surfaces
    ``template_builder`` writes into. Headers/footers/nested tables are
    out of scope for v1; the placeholder scan must agree with what the
    writer actually puts there.
    """
    doc = docx.Document(str(path))
    counts: Dict[str, int] = {}
    for para in doc.paragraphs:
        for m in PLACEHOLDER_RE.finditer(para.text or ""):
            counts[m.group(1)] = counts.get(m.group(1), 0) + 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for m in PLACEHOLDER_RE.finditer(cell.text or ""):
                    counts[m.group(1)] = counts.get(m.group(1), 0) + 1
    return counts


def _refuse_input_output_collisions(
    *,
    template_abs: Path,
    validation_abs: Path,
    out_docx_abs: Path,
    log_abs: Path,
    out_docx_path: Path,
    report: RenderReport,
) -> None:
    """Append a ``fatal_errors`` entry for every input/output collision.

    The renderer writes ``out_docx_abs`` and ``log_abs``. If either
    resolves to a renderer input, going through with the run would
    overwrite that input. ``_clean_stale_outputs`` already refused to
    delete those paths; this helper then surfaces the situation as a
    fatal so the user is told and no further work happens.
    """
    inputs = (("--template", template_abs),
              ("--run-validation", validation_abs))
    for label, in_abs in inputs:
        if out_docx_abs == in_abs:
            report.fatal_errors.append(
                f"--out ({out_docx_path}) resolves to the same file as "
                f"{label} ({in_abs}); refusing to overwrite a renderer "
                "input"
            )
        if log_abs == in_abs:
            report.fatal_errors.append(
                f"render_log.yml would be written at {log_abs}, which "
                f"is the same file as {label} ({in_abs}); refusing to "
                "overwrite a renderer input — choose a different --out "
                "directory"
            )
    if out_docx_abs == log_abs:
        report.fatal_errors.append(
            f"--out ({out_docx_path}) is the same file as the auto-"
            "derived render_log.yml in that directory; choose a "
            "different .docx filename"
        )


def _clean_stale_outputs(
    *,
    targets: tuple,
    protected: frozenset,
) -> None:
    """Unlink any stale render output at non-input target paths.

    ``targets`` is the pair ``(out_docx_abs, log_abs)`` — the two files
    this command writes. ``protected`` is the canonical absolute paths
    of the renderer inputs. A target that resolves to a protected path
    is left alone (deleting it would destroy a renderer input); every
    other target is unlinked if it exists. Unrelated files in the
    output directory are never touched. A directory at a target path
    (typo'd ``--out``) is left in place so the later ``Document.save``
    raises a clear error rather than this helper nuking the user's
    tree.
    """
    for p_abs in targets:
        if p_abs in protected:
            continue
        try:
            p_abs.unlink()
        except FileNotFoundError:
            pass
        except IsADirectoryError:
            pass


def _row_from_validation_dict(d: Dict[str, object]) -> RenderRow:
    row = RenderRow(
        word_id=_str_cell(d.get("Word ID")),
        word_location=_str_cell(d.get("Word Location")),
        word_raw=_str_cell(d.get("Word Raw Token")),
        word_unit=_str_cell(d.get("Word Unit")),
        source_sheet=_str_cell(d.get("Source Sheet")),
        source_cell=_str_cell(d.get("Source Cell")),
        raw_excel_value=_float_cell(d.get("Raw Excel Value")),
        generated_value=_float_cell(d.get("Generated Value")),
    )
    status_in = _str_cell(d.get("Status"))
    if status_in != "ok":
        row.status = STATUS_VALIDATION_NOT_OK
        row.detail = (
            f"validation Status={status_in!r}; render-docx only accepts "
            "rows whose run-preview status is ok"
        )
        return row
    if row.generated_value is None:
        row.status = STATUS_MISSING_GENERATED_VALUE
        row.detail = (
            f"validation row {row.word_id} has no Generated Value — "
            "render-docx requires exactly one generated value per word_id"
        )
    return row


def _str_cell(v) -> str:
    if v is None:
        return ""
    return str(v).strip()


def _float_cell(v) -> Optional[float]:
    if v is None:
        return None
    if isinstance(v, bool):
        # Excel can write TRUE/FALSE; treat as missing rather than coerce
        # to 1/0, mirroring run_preview's numeric-cell discipline.
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    if not s:
        return None
    try:
        return float(s.replace(",", ""))
    except ValueError:
        return None


# ---------------------------------------------------------------------------
# Docx substitution + log writer
# ---------------------------------------------------------------------------

def _substitute_placeholders(
    template_path: Path,
    rows: List[RenderRow],
    out_path: Path,
) -> Dict[str, int]:
    """Substitute every ``{{ word_NNNN }}`` occurrence and write the docx.

    Returns the actual count of substitutions performed per word_id —
    the render_log uses this to record duplicate-placeholder accounting
    (a single confirmed word_id may legitimately appear in multiple
    sentences of the historical report).
    """
    display_by_id: Dict[str, str] = {
        r.word_id: (r.display_text or "") for r in rows
    }
    counts: Dict[str, int] = {wid: 0 for wid in display_by_id}

    def replace(match: re.Match) -> str:
        wid = match.group(1)
        if wid in display_by_id:
            counts[wid] = counts.get(wid, 0) + 1
            return display_by_id[wid]
        # The orphan-placeholder gate already blocks this path; the
        # defense-in-depth fallback keeps the original placeholder string
        # in the unexpected case so a post-hoc audit can still see what
        # slipped through rather than getting silent text loss.
        return match.group(0)

    doc = docx.Document(str(template_path))
    for para in doc.paragraphs:
        text = para.text or ""
        new_text = PLACEHOLDER_RE.sub(replace, text)
        if new_text != text:
            # python-docx collapses runs when ``.text`` is assigned. The
            # template_builder already did the same when writing the
            # placeholders, so we're not regressing run-level formatting
            # here. A future styled-renderer would replace this with a
            # run-aware substitution.
            para.text = new_text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text or ""
                new_text = PLACEHOLDER_RE.sub(replace, text)
                if new_text != text:
                    cell.text = new_text

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return counts


def _write_log(report: RenderReport, path: Path) -> None:
    entries: List[Dict] = []
    for r in report.rows:
        entries.append({
            "word_id": r.word_id,
            "location": r.word_location,
            "source_sheet": r.source_sheet,
            "source_cell": r.source_cell,
            "raw_excel_value": r.raw_excel_value,
            "generated_value": r.generated_value,
            "raw_token": r.word_raw,
            "unit": r.word_unit,
            "display_text": r.display_text,
            "placeholder_occurrences": r.placeholder_occurrences,
            "status": r.status,
            "detail": r.detail,
        })
    doc = {
        "schema_version": 1,
        "inputs": {
            "template": str(report.template_path) if report.template_path else None,
            "run_validation": (
                str(report.validation_path) if report.validation_path else None
            ),
        },
        "summary": {
            "total_rows": len(report.rows),
            "ok": sum(1 for r in report.rows if r.status == STATUS_OK),
            "failed": sum(1 for r in report.rows if r.status != STATUS_OK),
            "total_replacements": sum(
                r.placeholder_occurrences for r in report.rows
            ),
            "distinct_placeholder_word_ids": len(
                report.template_placeholder_counts
            ),
        },
        "out_docx": str(report.out_docx_path) if report.out_docx_path else None,
        "replacements": entries,
    }
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        yaml.safe_dump(doc, allow_unicode=True, sort_keys=False, width=1000),
        encoding="utf-8",
    )


# ---------------------------------------------------------------------------
# CLI summary
# ---------------------------------------------------------------------------

def format_console_summary(report: RenderReport) -> str:
    """Human-readable summary for the CLI."""
    ok_count = sum(1 for r in report.rows if r.status == STATUS_OK)
    failed = [r for r in report.rows if r.status != STATUS_OK]
    lines = [
        "render-docx summary",
        "===================",
        f"  total validation rows : {len(report.rows)}",
        f"  ok                    : {ok_count}",
        f"  failed                : {len(failed)}",
    ]
    if report.template_placeholder_counts:
        total_placeholders = sum(report.template_placeholder_counts.values())
        distinct_placeholders = len(report.template_placeholder_counts)
        lines.append(
            f"  template placeholders : {total_placeholders} "
            f"({distinct_placeholders} distinct word_id(s))"
        )
    if report.out_docx_path is not None:
        lines.append(f"  rendered docx         : {report.out_docx_path}")
    if report.out_log_path is not None:
        lines.append(f"  render log            : {report.out_log_path}")
    if failed:
        lines.append("")
        lines.append("Failed rows (first 10):")
        for row in failed[:10]:
            lines.append(f"  - {row.word_id} [{row.status}] {row.detail}")
        if len(failed) > 10:
            lines.append(f"  … and {len(failed) - 10} more")
    if report.fatal_errors:
        lines.append("")
        lines.append("FATAL:")
        for err in report.fatal_errors:
            lines.append(f"  - {err}")
    return "\n".join(lines)
