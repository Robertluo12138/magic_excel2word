"""Extract numeric tokens from a .docx report along with locating context.

Every visible number we find produces a :class:`WordNumber`. The CLAUDE.md
guidance is explicit: the system must never silently drop a Word number, so
this profiler errs on the side of capturing every match :mod:`number_normalizer`
returns. Disambiguation belongs to the matcher and the human reviewer.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

import docx

from .number_normalizer import find_numbers


@dataclass
class WordNumber:
    location: str          # e.g. "paragraph:5" or "table:1/row:2/col:3"
    snippet: str           # text of the containing paragraph or cell
    label_context: List[str] = field(default_factory=list)  # extra context (table row/col headers)
    raw: str = ""
    value: float = 0.0
    unit: Optional[str] = None
    sign: int = 1
    offset: int = 0        # character offset within ``snippet``
    exclusion_reason: Optional[str] = None  # set when the number is captured for audit but skipped by the matcher


def profile_document(path: Path) -> List[WordNumber]:
    """Return every numeric token in the document in reading order.

    Scope (per the first milestone): top-level paragraphs and top-level table
    cells. Headers/footers, footnotes, text boxes, chart data, and nested
    tables are not yet read — capturing them is a follow-up; this is an
    explicit known gap so a reviewer can decide whether a given report stays
    within the supported surfaces.
    """
    doc = docx.Document(str(path))
    found: List[WordNumber] = []

    for idx, para in enumerate(doc.paragraphs):
        text = para.text or ""
        for pn in find_numbers(text):
            found.append(WordNumber(
                location=f"paragraph:{idx}",
                snippet=text,
                label_context=[],
                raw=pn.raw, value=pn.value, unit=pn.unit, sign=pn.sign,
                offset=pn.start,
                exclusion_reason=pn.exclusion_reason,
            ))

    for t_idx, table in enumerate(doc.tables):
        rows = list(table.rows)
        if not rows:
            continue
        # Pre-read all cell texts so we can hand each numeric cell its row and
        # column header strings as supplementary context.
        cell_texts: List[List[str]] = [[c.text or "" for c in row.cells] for row in rows]
        for r_idx, row in enumerate(rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text or ""
                context: List[str] = []
                if r_idx > 0 and c_idx < len(cell_texts[0]):
                    header = cell_texts[0][c_idx].strip()
                    if header and header != text:
                        context.append(header)
                if c_idx > 0:
                    row_label = cell_texts[r_idx][0].strip()
                    if row_label and row_label != text:
                        context.append(row_label)
                for pn in find_numbers(text):
                    found.append(WordNumber(
                        location=f"table:{t_idx}/row:{r_idx}/col:{c_idx}",
                        snippet=text,
                        label_context=list(context),
                        raw=pn.raw, value=pn.value, unit=pn.unit, sign=pn.sign,
                        offset=pn.start,
                        exclusion_reason=pn.exclusion_reason,
                    ))
    return found
