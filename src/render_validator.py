"""Read-only final-artifact validator for the rendered Word output.

``render-docx`` writes three things a reviewer must trust together:

* ``new_report.docx`` — the rendered Word report. No
  ``{{ word_NNNN }}`` placeholder strings may survive.
* ``render_log.yml`` — one ``replacements`` entry per rendered
  ``word_id`` with source sheet/cell, generated value, display text,
  occurrence count, and per-row status.
* ``run_validation.xlsx`` — written upstream by ``run-preview``. One
  row per confirmed ``word_id`` with the Excel value, generated value,
  transform interpretation, and per-row Status.

The three artifacts are the audit trail a human reads to convince
themselves the rendered docx contains every confirmed metric and
nothing else. If they drift apart, the audit silently lies — exactly
the "looks successful while omitting metrics" risk CLAUDE.md forbids.

This module reloads all three from disk and proves:

* the rendered docx contains **no** ``{{ word_NNNN }}`` placeholders;
* ``render_log.yml`` has **exactly one** replacement entry per
  ``run_validation`` ``word_id`` (and vice versa);
* **every** ``run_validation`` row has ``Status=ok``;
* **every** ``render_log`` row has ``status=ok``;
* every log row carries non-empty ``generated_value``, ``source_sheet``,
  ``source_cell``, and ``display_text``;
* every log row's ``source_sheet``, ``source_cell``,
  ``raw_excel_value``, ``generated_value``, ``raw_token``, and ``unit``
  AGREE with the matching ``run_validation`` row — presence alone isn't
  enough, because a hand-edited log could claim a fabricated source or
  invented number while the validation still holds the true value. This
  is the "no invented sources, cells, or numbers" guarantee CLAUDE.md
  requires;
* every rendered log row has ``placeholder_occurrences >= 1`` (the
  contract for a confirmed metric: at least one occurrence in the
  rendered docx);
* every ``display_text`` from the log appears in the rendered docx at
  least ``placeholder_occurrences`` times — a hand-edited docx that
  swapped "100元" for "999元" would otherwise pass every other gate
  (no leftover placeholders, log and validation untouched), silently
  misrepresenting the metric;
* no ``word_id`` appears twice in either artifact.

It is read-only: it never mutates an artifact, never re-renders the
docx, never calls out to an LLM, never invents data. A failure is a
directive to a human reviewer, not something this tool can paper over.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import docx
import openpyxl
import yaml

# Same Jinja-flavored placeholder family ``template_builder`` writes and
# the deterministic renderer substitutes. Compact ``{{word_NNNN}}`` is
# accepted so a leftover written by a manually massaged docx is still
# caught — naive substring matching against the canonical spaced form
# would silently green-light an unsafe rendered report.
PLACEHOLDER_RE = re.compile(r"\{\{\s*(word_\d{4,})\s*\}\}")

# Run-validation column names. Mirrors ``run_preview.write_run_validation``;
# the validator must read them by name (column order can drift if
# run_preview ever rearranges headers).
_RUN_HEADERS_REQUIRED = (
    "Word ID",
    "Source Sheet",
    "Source Cell",
    "Raw Excel Value",
    "Generated Value",
    "Word Raw Token",
    "Word Unit",
    "Status",
)

# Float tolerance for the cross-artifact value-agreement check. The
# generated/raw Excel values pass through one ``float`` round-trip
# (Excel cell → ``run_preview`` → log YAML → load) — exact equality
# would over-fire on legitimate binary-representation noise. The
# tolerance is tight enough that any real drift (a hand-edited digit,
# an interpretation swap) still trips the gate.
_FLOAT_ABS_TOL = 1e-9
_FLOAT_REL_TOL = 1e-9


@dataclass
class ValidationIssue:
    code: str
    message: str


@dataclass
class RenderValidationReport:
    issues: List[ValidationIssue] = field(default_factory=list)
    docx_path: Optional[Path] = None
    render_log_path: Optional[Path] = None
    run_validation_path: Optional[Path] = None

    @property
    def ok(self) -> bool:
        return not self.issues

    def add(self, code: str, message: str) -> None:
        self.issues.append(ValidationIssue(code=code, message=message))


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def validate_render(
    docx_path: Path,
    render_log_path: Path,
    run_validation_path: Path,
) -> RenderValidationReport:
    """Cross-check the three rendered-output artifacts and return a report.

    Read-only. Never mutates an artifact, never re-renders the docx,
    never calls out to an LLM. The CLI decides how to translate
    failures to exit codes. Pre-condition: all three files exist (the
    CLI guards that for friendlier error messages).
    """
    report = RenderValidationReport(
        docx_path=docx_path,
        render_log_path=render_log_path,
        run_validation_path=run_validation_path,
    )

    # --- Load each artifact ------------------------------------------------
    # Each loader appends a ``cannot_read_*`` issue and returns ``None`` on
    # failure so the rest of the checks can short-circuit cleanly without
    # crashing on a missing file or malformed schema.
    docx_text = _read_docx_text(docx_path, report)
    log_doc = _read_render_log(render_log_path, report)
    run_rows = _read_run_validation(run_validation_path, report)
    if docx_text is None or log_doc is None or run_rows is None:
        return report

    log_entries = log_doc.get("replacements")
    if not isinstance(log_entries, list):
        report.add(
            "render_log_schema",
            f"{render_log_path.name} is missing a top-level "
            "'replacements' list",
        )
        return report

    # --- Run every cross-artifact check -----------------------------------
    _check_no_placeholders_in_docx(docx_text, report)
    log_ids = _check_uniqueness("render_log", log_entries, "word_id", report)
    run_ids = _check_uniqueness("run_validation", run_rows, "Word ID", report)
    _check_one_to_one_word_ids(log_ids, run_ids, report)
    _check_run_validation_status_ok(run_rows, report)
    _check_render_log_status_ok(log_entries, report)
    _check_render_log_required_fields(log_entries, report)
    _check_placeholder_occurrences(log_entries, report)
    _check_audit_value_agreement(log_entries, run_rows, report)
    _check_display_text_present_in_docx(log_entries, docx_text, report)
    return report


def format_validation_summary(
    report: RenderValidationReport,
) -> str:
    """Human-readable summary for the CLI."""
    if report.ok:
        return (
            "validate-render: OK — rendered docx, render_log.yml, and "
            "run_validation.xlsx agree on every word_id."
        )
    lines = [
        f"validate-render: FAILED — {len(report.issues)} issue(s)",
    ]
    if report.docx_path is not None:
        lines.append(f"  docx           : {report.docx_path}")
    if report.render_log_path is not None:
        lines.append(f"  render log     : {report.render_log_path}")
    if report.run_validation_path is not None:
        lines.append(f"  run validation : {report.run_validation_path}")
    for issue in report.issues:
        lines.append(f"  - [{issue.code}] {issue.message}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Loaders
# ---------------------------------------------------------------------------

def _read_docx_text(
    path: Path, report: RenderValidationReport
) -> Optional[str]:
    """Walk paragraphs + top-level table cells — the same surfaces the
    renderer writes into. Headers/footers/nested tables are out of v1
    scope for both writer and validator.
    """
    try:
        doc = docx.Document(str(path))
    except Exception as exc:  # broad: python-docx raises several types
        report.add("cannot_read_docx", f"cannot open {path.name}: {exc}")
        return None
    parts: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _read_render_log(
    path: Path, report: RenderValidationReport
) -> Optional[Dict]:
    try:
        data = yaml.safe_load(path.read_text(encoding="utf-8"))
    except Exception as exc:
        report.add("cannot_read_render_log", f"cannot parse {path.name}: {exc}")
        return None
    if not isinstance(data, dict):
        report.add(
            "render_log_schema",
            f"{path.name} top-level is not a YAML mapping",
        )
        return None
    return data


def _read_run_validation(
    path: Path, report: RenderValidationReport
) -> Optional[List[Dict[str, object]]]:
    try:
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    except Exception as exc:
        report.add(
            "cannot_read_run_validation",
            f"cannot open {path.name}: {exc}",
        )
        return None
    try:
        ws = wb.active
        rows_iter = ws.iter_rows(values_only=True)
        try:
            header = list(next(rows_iter))
        except StopIteration:
            report.add(
                "run_validation_empty",
                f"{path.name} has no header row",
            )
            return None
        for req in _RUN_HEADERS_REQUIRED:
            if req not in header:
                report.add(
                    "run_validation_schema",
                    f"{path.name} missing required column {req!r}",
                )
                return None
        rows: List[Dict[str, object]] = []
        for raw in rows_iter:
            if raw is None:
                continue
            if all(cell is None or cell == "" for cell in raw):
                continue
            rows.append({h: v for h, v in zip(header, raw)})
        return rows
    finally:
        wb.close()


# ---------------------------------------------------------------------------
# Checks
# ---------------------------------------------------------------------------

def _check_no_placeholders_in_docx(
    docx_text: str, report: RenderValidationReport
) -> None:
    leftovers = sorted(set(PLACEHOLDER_RE.findall(docx_text)))
    if leftovers:
        report.add(
            "leftover_placeholder",
            "rendered docx still contains {{ word_NNNN }} placeholder(s): "
            + ", ".join(leftovers),
        )


def _check_uniqueness(
    label: str,
    rows: List[Dict],
    id_key: str,
    report: RenderValidationReport,
) -> List[str]:
    """Return the de-duplicated id list; flag empties + duplicates."""
    ids: List[str] = []
    seen: Dict[str, int] = {}
    for r in rows:
        wid = r.get(id_key)
        wid_str = "" if wid is None else str(wid).strip()
        if not wid_str:
            report.add(
                f"{label}_missing_word_id",
                f"{label} has a row with empty {id_key}",
            )
            continue
        seen[wid_str] = seen.get(wid_str, 0) + 1
        ids.append(wid_str)
    dupes = sorted(k for k, n in seen.items() if n > 1)
    if dupes:
        report.add(
            f"{label}_duplicate_word_id",
            f"{label} has duplicate word_id(s): " + ", ".join(dupes),
        )
    return ids


def _check_one_to_one_word_ids(
    log_ids: List[str],
    run_ids: List[str],
    report: RenderValidationReport,
) -> None:
    log_set = set(log_ids)
    run_set = set(run_ids)
    only_in_log = sorted(log_set - run_set)
    only_in_run = sorted(run_set - log_set)
    if only_in_log:
        report.add(
            "extra_render_log_row",
            "render_log.yml has word_id(s) absent from "
            f"run_validation.xlsx: {', '.join(only_in_log)}",
        )
    if only_in_run:
        report.add(
            "missing_render_log_row",
            "run_validation.xlsx has word_id(s) absent from "
            f"render_log.yml: {', '.join(only_in_run)}",
        )


def _check_run_validation_status_ok(
    rows: List[Dict[str, object]],
    report: RenderValidationReport,
) -> None:
    """Every confirmed run_validation row must be Status=ok before the
    docx ships. ``render-docx`` already refuses to substitute a non-ok
    row, so a non-ok row reaching this validator means either the
    validation file was hand-edited after the render, or the docx was
    produced by an out-of-band tool — either way, refuse loudly.
    """
    bad: List[Tuple[str, str]] = []
    for r in rows:
        status = str(r.get("Status") or "").strip()
        if status != "ok":
            wid = str(r.get("Word ID") or "").strip() or "<missing word_id>"
            bad.append((wid, status))
    if bad:
        report.add(
            "run_validation_status_not_ok",
            "run_validation.xlsx has non-ok row(s): "
            + ", ".join(f"{wid}={status!r}" for wid, status in bad),
        )


def _check_render_log_status_ok(
    entries: List[Dict],
    report: RenderValidationReport,
) -> None:
    bad: List[Tuple[str, str]] = []
    for e in entries:
        status = str(e.get("status") or "").strip()
        if status != "ok":
            wid = str(e.get("word_id") or "").strip() or "<missing word_id>"
            bad.append((wid, status))
    if bad:
        report.add(
            "render_log_status_not_ok",
            "render_log.yml has non-ok entr(ies): "
            + ", ".join(f"{wid}={status!r}" for wid, status in bad),
        )


def _check_render_log_required_fields(
    entries: List[Dict],
    report: RenderValidationReport,
) -> None:
    """Every rendered entry must carry the four traceability fields a
    reviewer needs to walk a number back to its Excel origin.

    A missing ``generated_value``, ``source_sheet``, ``source_cell``,
    or ``display_text`` would leave a hole in the audit trail — even
    if the docx happens to look right, the next reviewer can't prove
    where the number came from.
    """
    REQUIRED_FIELDS = (
        "generated_value",
        "source_sheet",
        "source_cell",
        "display_text",
    )
    for e in entries:
        wid = str(e.get("word_id") or "").strip() or "<missing word_id>"
        missing: List[str] = []
        for field_name in REQUIRED_FIELDS:
            val = e.get(field_name)
            # ``generated_value`` is numeric — treat the literal ``None``
            # and an empty string as missing. ``0`` and ``0.0`` are
            # legitimate values and must NOT trip the gate.
            if field_name == "generated_value":
                if val is None or (isinstance(val, str) and not val.strip()):
                    missing.append(field_name)
            else:
                if val is None or (isinstance(val, str) and not val.strip()):
                    missing.append(field_name)
        if missing:
            report.add(
                "render_log_missing_field",
                f"{wid}: render_log entry missing field(s): "
                + ", ".join(missing),
            )


def _check_placeholder_occurrences(
    entries: List[Dict],
    report: RenderValidationReport,
) -> None:
    """``placeholder_occurrences`` records how many times the renderer
    substituted a given word_id into the docx. A zero count means
    either the template never referenced this word_id (in which case
    ``render-docx`` would have failed with STATUS_NO_PLACEHOLDER) or
    the log was hand-edited — either way, refuse loudly so a downstream
    consumer can't trust a rendered docx that drops a confirmed metric.
    """
    bad: List[Tuple[str, object]] = []
    for e in entries:
        wid = str(e.get("word_id") or "").strip() or "<missing word_id>"
        count = e.get("placeholder_occurrences")
        if not isinstance(count, int) or isinstance(count, bool) or count < 1:
            bad.append((wid, count))
    if bad:
        report.add(
            "zero_placeholder_occurrences",
            "render_log.yml has entr(ies) with no placeholder occurrences: "
            + ", ".join(f"{wid}={count!r}" for wid, count in bad),
        )


def _check_audit_value_agreement(
    log_entries: List[Dict],
    run_rows: List[Dict[str, object]],
    report: RenderValidationReport,
) -> None:
    """For every word_id present in both artifacts, the log's traceability
    fields must match the run_validation row.

    Without this check, presence alone (``_check_render_log_required_fields``)
    only proves the log carries *some* value — a hand-edited log could
    claim a fabricated ``source_cell`` or invented ``generated_value``
    while run_validation still holds the true value, and the rendered
    docx would look successful while silently misrepresenting where the
    number came from. That's the exact "invented sources/cells/numbers"
    risk CLAUDE.md forbids.

    The compared fields are the ones a reviewer walks back to the Excel
    origin: ``source_sheet``, ``source_cell``, ``raw_excel_value``,
    ``generated_value``, ``raw_token``, and ``unit``. Drift in any of
    them is surfaced per word_id with the specific fields that
    disagreed, so a reviewer can spot whether the log or the validation
    was tampered with.
    """
    runs_by_id: Dict[str, Dict[str, object]] = {}
    for r in run_rows:
        wid = str(r.get("Word ID") or "").strip()
        if wid and wid not in runs_by_id:
            # First occurrence wins; the duplicate check has already
            # surfaced the dupe separately, so don't double-fire here.
            runs_by_id[wid] = r

    PAIRS = (
        # (log_field, run_field, compare_mode)
        ("source_sheet", "Source Sheet", "str"),
        ("source_cell", "Source Cell", "str"),
        ("raw_excel_value", "Raw Excel Value", "float"),
        ("generated_value", "Generated Value", "float"),
        ("raw_token", "Word Raw Token", "str"),
        ("unit", "Word Unit", "str"),
    )

    for e in log_entries:
        wid = str(e.get("word_id") or "").strip()
        if not wid:
            continue
        r = runs_by_id.get(wid)
        if r is None:
            # One-to-one check already flagged the orphan; skip here.
            continue
        drifted: List[str] = []
        for log_field, run_field, mode in PAIRS:
            log_v = e.get(log_field)
            run_v = r.get(run_field)
            if mode == "float":
                if not _floats_agree(log_v, run_v):
                    drifted.append(
                        f"{log_field}={log_v!r} vs run "
                        f"{run_field}={run_v!r}"
                    )
            else:
                if _str_norm(log_v) != _str_norm(run_v):
                    drifted.append(
                        f"{log_field}={log_v!r} vs run "
                        f"{run_field}={run_v!r}"
                    )
        if drifted:
            report.add(
                "audit_value_drift",
                f"{wid}: render_log and run_validation disagree on "
                + "; ".join(drifted),
            )


def _str_norm(v: object) -> str:
    """Treat ``None`` and ``""`` as the same missing value so a log entry
    that omits a string field doesn't have to match against ``"None"``."""
    if v is None:
        return ""
    return str(v).strip()


def _floats_agree(a: object, b: object) -> bool:
    """Return True when two numeric audit-trail values are close enough
    that the only difference is float round-trip noise.

    ``None`` on both sides agrees (both artifacts say "no value");
    ``None`` on exactly one side disagrees (drift). Non-numeric values
    are compared via their stripped string forms — that way an
    accidentally-stringified ``"100"`` in the log still agrees with a
    numeric ``100`` in the validation cell.
    """
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    fa = _to_float_or_none(a)
    fb = _to_float_or_none(b)
    if fa is None or fb is None:
        return _str_norm(a) == _str_norm(b)
    diff = abs(fa - fb)
    if diff <= _FLOAT_ABS_TOL:
        return True
    return diff <= _FLOAT_REL_TOL * max(abs(fa), abs(fb))


def _to_float_or_none(v: object) -> Optional[float]:
    if isinstance(v, bool):
        # bool is a subclass of int — treat TRUE/FALSE as non-numeric
        # so they fall through to string comparison rather than silently
        # coercing to 1/0 and matching a true 1.0 in the other artifact.
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).strip()
    if not s:
        return None
    try:
        return float(s.replace(",", ""))
    except ValueError:
        return None


def _check_display_text_present_in_docx(
    log_entries: List[Dict],
    docx_text: str,
    report: RenderValidationReport,
) -> None:
    """Every log entry's ``display_text`` must appear in the docx at
    least ``placeholder_occurrences`` times.

    Without this check, the validator only knows the audit trail is
    self-consistent — it has no eyes on what the docx actually says.
    A hand-edited docx that swapped a rendered "100元" for "999元"
    would pass every other gate (no leftover placeholders, the log
    and validation agree on what was *supposed* to be rendered)
    while silently misrepresenting the metric. That's the exact
    "looks successful while silently misrepresenting metrics" risk
    CLAUDE.md forbids.

    The check is occurrence-counted and aggregated per unique
    display_text so multiple word_ids that legitimately share the
    same display string (e.g. two metrics both rendered as "100%")
    are tallied together. Under-count fires; over-count is left
    silent because a docx may legitimately contain the display_text
    in narrative prose surrounding the placeholder.

    Substring-overlap safety: a naive ``docx_text.count(display_text)``
    would falsely match a shorter display_text inside a longer one —
    a docx tampered from "100元 1100元" to "999元 1100元" still has
    "100元" as a substring of "1100元", so the missing occurrence
    would slip through. To close that gap, we process display_texts
    longest-first and replace each matched span with a sentinel before
    counting shorter ones, so a shorter display_text can never count
    an occurrence already consumed by a longer overlapping match.
    """
    expected_by_text: Dict[str, int] = {}
    wids_by_text: Dict[str, List[str]] = {}
    for e in log_entries:
        text = e.get("display_text")
        text_str = "" if text is None else str(text)
        if not text_str:
            # Missing display_text is already surfaced by
            # _check_render_log_required_fields; don't double-fire here.
            continue
        occ = e.get("placeholder_occurrences")
        if (
            not isinstance(occ, int)
            or isinstance(occ, bool)
            or occ < 1
        ):
            # Invalid occurrences are surfaced by
            # _check_placeholder_occurrences; don't double-fire here.
            continue
        expected_by_text[text_str] = (
            expected_by_text.get(text_str, 0) + occ
        )
        wid = str(e.get("word_id") or "").strip() or "<missing word_id>"
        wids_by_text.setdefault(text_str, []).append(wid)

    # Private-use Unicode sentinel — chosen so it can't collide with
    # legitimate docx text (Word strips most control chars; private-use
    # codepoints round-trip only when explicitly authored).
    SENTINEL = ""
    remaining = docx_text
    actual_by_text: Dict[str, int] = {}
    # Longest-first: a longer display_text gets first claim on its
    # span, then shorter ones count only what survives. Ties broken by
    # the string itself so the ordering is stable — when two distinct
    # display_texts have the same length, neither can be a substring
    # of the other, so tie-break order doesn't change the count.
    for text_str in sorted(
        expected_by_text.keys(), key=lambda t: (-len(t), t),
    ):
        count = remaining.count(text_str)
        actual_by_text[text_str] = count
        if count > 0:
            # Preserve character width so any later positional checks
            # still see the same offsets.
            remaining = remaining.replace(
                text_str, SENTINEL * len(text_str),
            )

    drifted: List[Tuple[str, int, int, List[str]]] = []
    for text_str, expected in expected_by_text.items():
        actual = actual_by_text.get(text_str, 0)
        if actual < expected:
            drifted.append(
                (text_str, expected, actual, wids_by_text[text_str])
            )

    if drifted:
        details = "; ".join(
            f"display_text={text!r} expected {expected}x but docx has "
            f"{actual}x (word_id(s): {', '.join(wids)})"
            for text, expected, actual, wids in drifted
        )
        report.add(
            "docx_rendered_text_missing",
            "rendered docx is missing expected display text(s): " + details,
        )
