"""Generate a paired synthetic Excel + Word sample for learn-mode testing.

The data is hand-built, not random: tests need stable byte-for-byte outputs
and the matcher needs to be exercised against deliberate edge cases —
ambiguous duplicates (234,567 appears in three places with distinct labels),
loose rounding (1.23 亿元 ↔ 123,456,789), Word numbers with no Excel source
(``0.70 个百分点``, ``15%``), and unit-suffix coverage (万元, 万人, 万单, %).

Real company data must never live here. Per the privacy rules in CLAUDE.md,
the corpus is fake but structurally realistic.
"""
from __future__ import annotations

from pathlib import Path
from typing import Tuple

import docx
import openpyxl

# (label, 2026年4月, 2026年5月, 同比%)
_METRICS_MONTHLY = [
    ("营业收入(元)",   98_765_432.10, 123_456_789.00, 25.00),
    ("净利润(元)",     18_765_432.00,  23_456_789.00, 24.95),
    ("毛利率(%)",            35.50,         36.20,    1.97),
    ("用户总数",        12_345_678,     13_456_789,   8.99),
    ("新增用户",           234_567,        345_678,  47.36),
]

# (label, 第18周, 第19周, 第20周)
_METRICS_WEEKLY = [
    ("日均订单数",     234_567,    245_678,    256_789),
    ("客单价(元)",       156.78,     162.45,     168.90),
    ("转化率(%)",          3.45,       3.67,       3.89),
    ("退款率(%)",          1.23,       1.18,       1.15),
    ("周GMV(万元)",   12_345.67,  13_456.78,  14_567.89),
]

# (channel, 用户数, 订单数, 销售额(元), 占比%)
_METRICS_CHANNELS = [
    ("自然流量", 3_456_789, 234_567, 45_678_901, 32.50),
    ("付费广告", 2_345_678, 198_765, 38_765_432, 27.60),
    ("社交媒体", 1_987_654, 154_321, 30_123_456, 21.45),
    ("应用商店", 1_456_789, 123_456, 24_987_654, 17.80),
    ("其他",       234_567,  12_345,    956_321,  0.65),
]


def generate(out_dir: Path) -> Tuple[Path, Path]:
    """Write ``historical.xlsx`` and ``finished_report.docx`` under ``out_dir``."""
    out_dir.mkdir(parents=True, exist_ok=True)
    xlsx_path = out_dir / "historical.xlsx"
    docx_path = out_dir / "finished_report.docx"
    _write_excel(xlsx_path)
    _write_word(docx_path)
    return xlsx_path, docx_path


def _write_excel(path: Path) -> None:
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "月度核心指标"
    ws1.append(["指标", "2026年4月", "2026年5月", "同比(%)"])
    for row in _METRICS_MONTHLY:
        ws1.append(list(row))

    ws2 = wb.create_sheet("周度运营指标")
    ws2.append(["指标", "第18周", "第19周", "第20周"])
    for row in _METRICS_WEEKLY:
        ws2.append(list(row))

    ws3 = wb.create_sheet("渠道分析")
    ws3.append(["渠道", "用户数", "订单数", "销售额(元)", "占比(%)"])
    for row in _METRICS_CHANNELS:
        ws3.append(list(row))

    wb.save(str(path))


def _write_word(path: Path) -> None:
    doc = docx.Document()
    doc.add_heading("2026年5月经营月度报告", level=1)

    doc.add_heading("一、整体业绩", level=2)
    doc.add_paragraph("本月报告期为2026年5月。")
    doc.add_paragraph(
        "5月营业收入达12,345.68万元，同比增长25.00%；"
        "净利润2,345.68万元，同比增长24.95%。"
    )
    doc.add_paragraph(
        "毛利率提升至36.20%，较上月毛利率35.50%提升0.70个百分点。"
    )
    doc.add_paragraph(
        "本月用户总数达到1,345.68万人，其中新增用户34.57万人。"
    )
    doc.add_paragraph(
        "若以亿元口径展示，本月营业收入约为1.23亿元。"
    )
    doc.add_paragraph(
        "4月营业收入为9,876.54万元，4月净利润为1,876.54万元。"
    )

    doc.add_heading("二、周度运营趋势", level=2)
    doc.add_paragraph("第20周日均订单数达25.68万单。")
    doc.add_paragraph(
        "第19周日均订单数为24.57万单，第18周为23.46万单。"
    )
    doc.add_paragraph(
        "客单价稳步提升：第18周156.78元、第19周162.45元、第20周168.90元。"
    )
    doc.add_paragraph(
        "第20周转化率达3.89%，退款率控制在1.15%以内。"
    )
    doc.add_paragraph("第20周GMV达14,567.89万元。")

    doc.add_heading("三、渠道分析", level=2)
    doc.add_paragraph("下表展示主要获客渠道的表现：")
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "渠道"
    hdr[1].text = "用户数"
    hdr[2].text = "销售额"
    hdr[3].text = "占比"
    rows_data = [
        ("自然流量", "345.68万人", "4,567.89万元", "32.50%"),
        ("付费广告", "234.57万人", "3,876.54万元", "27.60%"),
        ("社交媒体", "198.77万人", "3,012.35万元", "21.45%"),
        ("应用商店", "145.68万人", "2,498.77万元", "17.80%"),
        ("其他",     "23.46万人",  "95.63万元",   "0.65%"),
    ]
    for channel, users, sales, share in rows_data:
        row = table.add_row().cells
        row[0].text = channel
        row[1].text = users
        row[2].text = sales
        row[3].text = share

    doc.add_heading("四、未来展望", level=2)
    # The "15%" below has no source in the workbook — deliberately, so the
    # matcher must surface at least one UNRESOLVED case end-to-end.
    doc.add_paragraph("我们计划在下季度将运营效率提升15%。")

    doc.save(str(path))
