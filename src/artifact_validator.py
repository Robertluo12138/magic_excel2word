"""Read-only consistency check across the four learn-mode artifacts.

``learn`` writes ``mapping_review.xlsx``, ``auto_mapping.yml``,
``converted_template.docx``, and ``confidence_report.md``. Each one is a
different lens on the same set of Word numbers; if they drift apart, the
audit trail silently lies. This module reloads all four from disk and
proves they tell the same story:

* every ``word_id`` is unique within each artifact;
* the XLSX has exactly one review row per YAML mapping (and vice versa);
* shared rows agree on location, raw text, confidence status, review
  status, and placeholder status;
* the only Word IDs that carry a ``{{ word_NNNN }}`` placeholder (in the
  YAML, in the XLSX, *and* in the converted docx) are the ones the YAML
  marks ``placeholder_status: applied`` — and those must be HIGH or
  MEDIUM. LOW, UNRESOLVED, and EXCLUDED tokens must never be templatized;
* if the YAML summary reports any UNRESOLVED or EXCLUDED counts, the
  Markdown confidence report mentions them by name.

The module is read-only: it never mutates an artifact, never calls out to
an LLM, and never invents data. A failure is always a directive to a
human reviewer, not something the tool can paper over.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

import docx
import openpyxl
import yaml

# Word IDs are emitted as ``word_NNNN`` (zero-padded, >=4 digits) by
# template_builder.assign_word_ids. The regex accepts longer runs so we
# don't false-negative if the corpus ever exceeds 9999 tokens.
PLACEHOLDER_RE = re.compile(r"\{\{\s*(word_\d{4,})\s*\}\}")

REQUIRED_ARTIFACTS = (
    "mapping_review.xlsx",
    "auto_mapping.yml",
    "converted_template.docx",
    "confidence_report.md",
)

SAFE_STATUSES = frozenset({"HIGH", "MEDIUM"})
UNSAFE_STATUSES = frozenset({"LOW", "UNRESOLVED", "EXCLUDED"})


@dataclass
class ValidationIssue:
    code: str
    message: str


@dataclass
class ValidationReport:
    issues: List[ValidationIssue] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return not self.issues

    def add(self, code: str, message: str) -> None:
        self.issues.append(ValidationIssue(code=code, message=message))


def validate_artifacts(out_dir: Path) -> ValidationReport:
    """Run every consistency check and return a structured report."""
    report = ValidationReport()
    paths = {name: out_dir / name for name in REQUIRED_ARTIFACTS}
    for name, path in paths.items():
        if not path.exists():
            report.add("missing_artifact", f"required artifact not found: {path}")
    if not report.ok:
        # Returning early here is intentional: every downstream check
        # would crash on the missing file, and the missing-artifact list
        # is already the most useful thing we can hand a reviewer.
        return report

    xlsx_rows = _read_mapping_review(paths["mapping_review.xlsx"], report)
    yaml_doc = _read_auto_mapping(paths["auto_mapping.yml"], report)
    docx_text = _read_docx_text(paths["converted_template.docx"], report)
    md_text = _read_md(paths["confidence_report.md"], report)
    if xlsx_rows is None or yaml_doc is None or docx_text is None or md_text is None:
        return report

    yaml_mappings = yaml_doc.get("mappings", [])
    if not isinstance(yaml_mappings, list):
        report.add("yaml_schema", "auto_mapping.yml has no 'mappings' list")
        return report

    _check_unique_word_ids(xlsx_rows, yaml_mappings, report)
    _check_one_xlsx_row_per_yaml(xlsx_rows, yaml_mappings, report)
    _check_field_agreement(xlsx_rows, yaml_mappings, report)
    _check_placeholder_rules(yaml_mappings, xlsx_rows, docx_text, report)
    _check_report_mentions(yaml_doc, md_text, report)
    return report


def format_validation_summary(report: ValidationReport, out_dir: Path) -> str:
    if report.ok:
        return (
            f"validate-artifacts: OK — {out_dir} is internally consistent "
            f"across {len(REQUIRED_ARTIFACTS)} artifacts."
        )
    lines = [
        f"validate-artifacts: FAILED — {len(report.issues)} issue(s) in {out_dir}",
    ]
    for issue in report.issues:
        lines.append(f"  - [{issue.code}] {issue.message}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Loaders
# ---------------------------------------------------------------------------

def _read_mapping_review(
    path: Path, report: ValidationReport
) -> Optional[List[Dict[str, object]]]:
    try:
        wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    except Exception as exc:  # broad on purpose: openpyxl raises several types
        report.add("xlsx_unreadable", f"cannot open {path.name}: {exc}")
        return None
    try:
        ws = wb.active
        rows_iter = ws.iter_rows(values_only=True)
        try:
            header = list(next(rows_iter))
        except StopIteration:
            report.add("xlsx_empty", f"{path.name} has no header row")
            return None
        rows: List[Dict[str, object]] = []
        for raw in rows_iter:
            if raw is None:
                continue
            if all(cell is None or cell == "" for cell in raw):
                continue
            rows.append({h: v for h, v in zip(header, raw)})
        return rows
    finally:
        wb.close()


def _read_auto_mapping(path: Path, report: ValidationReport) -> Optional[Dict]:
    try:
        data = yaml.safe_load(path.read_text(encoding="utf-8"))
    except Exception as exc:
        report.add("yaml_unreadable", f"cannot parse {path.name}: {exc}")
        return None
    if not isinstance(data, dict):
        report.add("yaml_schema", f"{path.name} top-level is not a mapping")
        return None
    return data


def _read_docx_text(path: Path, report: ValidationReport) -> Optional[str]:
    try:
        doc = docx.Document(str(path))
    except Exception as exc:
        report.add("docx_unreadable", f"cannot open {path.name}: {exc}")
        return None
    parts: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _read_md(path: Path, report: ValidationReport) -> Optional[str]:
    try:
        return path.read_text(encoding="utf-8")
    except Exception as exc:
        report.add("md_unreadable", f"cannot read {path.name}: {exc}")
        return None


# ---------------------------------------------------------------------------
# Checks
# ---------------------------------------------------------------------------

def _check_unique_word_ids(
    xlsx_rows: List[Dict[str, object]],
    yaml_mappings: List[Dict],
    report: ValidationReport,
) -> None:
    xlsx_ids = [r.get("Word ID") for r in xlsx_rows]
    yaml_ids = [m.get("word_id") for m in yaml_mappings]
    for label, ids in (("xlsx", xlsx_ids), ("yaml", yaml_ids)):
        seen: Dict[object, int] = {}
        for wid in ids:
            if wid is None or wid == "":
                report.add(
                    f"{label}_missing_word_id",
                    f"{label} has a row with empty Word ID",
                )
                continue
            seen[wid] = seen.get(wid, 0) + 1
        dupes = sorted(str(k) for k, n in seen.items() if n > 1)
        if dupes:
            report.add(
                f"{label}_duplicate_word_id",
                f"{label} has duplicate Word IDs: {', '.join(dupes)}",
            )


def _check_one_xlsx_row_per_yaml(
    xlsx_rows: List[Dict[str, object]],
    yaml_mappings: List[Dict],
    report: ValidationReport,
) -> None:
    xlsx_ids = {r.get("Word ID") for r in xlsx_rows if r.get("Word ID")}
    yaml_ids = {m.get("word_id") for m in yaml_mappings if m.get("word_id")}
    only_xlsx = sorted(str(x) for x in xlsx_ids - yaml_ids)
    only_yaml = sorted(str(x) for x in yaml_ids - xlsx_ids)
    if only_xlsx:
        report.add(
            "xlsx_orphan_word_ids",
            f"XLSX has Word IDs absent from YAML: {', '.join(only_xlsx)}",
        )
    if only_yaml:
        report.add(
            "yaml_orphan_word_ids",
            f"YAML has Word IDs absent from XLSX: {', '.join(only_yaml)}",
        )
    if len(xlsx_rows) != len(yaml_mappings):
        report.add(
            "row_count_mismatch",
            f"XLSX has {len(xlsx_rows)} review row(s) but YAML has "
            f"{len(yaml_mappings)} mapping entr(ies); learn mode must emit "
            "exactly one XLSX row per YAML mapping",
        )


def _check_field_agreement(
    xlsx_rows: List[Dict[str, object]],
    yaml_mappings: List[Dict],
    report: ValidationReport,
) -> None:
    xlsx_by_id: Dict[str, Dict[str, object]] = {}
    for r in xlsx_rows:
        wid = r.get("Word ID")
        if isinstance(wid, str) and wid:
            xlsx_by_id[wid] = r
    for m in yaml_mappings:
        wid = m.get("word_id")
        if not isinstance(wid, str) or wid not in xlsx_by_id:
            continue
        x = xlsx_by_id[wid]
        # location must match exactly — same join key for table/paragraph
        # offsets, so a drift here would mean the artifacts disagree on
        # where the Word number lives.
        if x.get("Word Location") != m.get("location"):
            report.add(
                "location_mismatch",
                f"{wid}: XLSX location={x.get('Word Location')!r} "
                f"vs YAML location={m.get('location')!r}",
            )
        if str(x.get("Word Raw Token") or "") != str(m.get("raw") or ""):
            report.add(
                "raw_mismatch",
                f"{wid}: XLSX raw={x.get('Word Raw Token')!r} "
                f"vs YAML raw={m.get('raw')!r}",
            )
        # Both columns carry the matcher's confidence label; the YAML
        # exposes it as 'status' (the canonical name) AND 'confidence'.
        if x.get("Confidence") != m.get("status"):
            report.add(
                "status_mismatch",
                f"{wid}: XLSX Confidence={x.get('Confidence')!r} "
                f"vs YAML status={m.get('status')!r}",
            )
        # Review Status and Placeholder Status are the audit-trail columns
        # the reviewer reads first. Tampering with either in the XLSX would
        # let an LOW/UNRESOLVED row impersonate a confirmed mapping (or
        # vice versa) without touching the YAML, so the validator must
        # surface them as distinct issue codes from the confidence drift
        # above — a Codex-style review probe needs to tell *which* column
        # was corrupted.
        if x.get("Review Status") != m.get("review_status"):
            report.add(
                "review_status_mismatch",
                f"{wid}: XLSX Review Status={x.get('Review Status')!r} "
                f"vs YAML review_status={m.get('review_status')!r}",
            )
        if x.get("Placeholder Status") != m.get("placeholder_status"):
            report.add(
                "placeholder_status_mismatch",
                f"{wid}: XLSX Placeholder Status={x.get('Placeholder Status')!r} "
                f"vs YAML placeholder_status={m.get('placeholder_status')!r}",
            )


def _check_placeholder_rules(
    yaml_mappings: List[Dict],
    xlsx_rows: List[Dict[str, object]],
    docx_text: str,
    report: ValidationReport,
) -> None:
    # ``PLACEHOLDER_RE`` accepts every Jinja-legal spacing (``{{ word_0001 }}``,
    # ``{{word_0001}}``, ``{{  word_0001  }}``) so a leak written in a more
    # compact form than the writer's canonical spaced output is still caught.
    # Naive substring matching against ``"{{ wid }}"`` would miss the compact
    # form and silently green-light an unsafe template.
    docx_placeholder_ids = set(PLACEHOLDER_RE.findall(docx_text))
    xlsx_by_id: Dict[str, Dict[str, object]] = {
        r.get("Word ID"): r for r in xlsx_rows
        if isinstance(r.get("Word ID"), str)
    }

    for m in yaml_mappings:
        wid = m.get("word_id")
        if not isinstance(wid, str):
            continue
        status = m.get("status")
        ph_status = m.get("placeholder_status")
        ph_value = m.get("placeholder")
        # The canonical spaced form is the writer's contract — checking YAML
        # and XLSX placeholder fields against it catches both compact-form
        # tampering and accidental whitespace drift.
        canonical_token = "{{ " + wid + " }}"

        if ph_status == "applied":
            # Only HIGH/MEDIUM may carry a placeholder. The matcher and
            # template builder enforce this on write; the validator
            # enforces it on read to catch hand-edited YAML.
            if status not in SAFE_STATUSES:
                report.add(
                    "unsafe_status_with_placeholder",
                    f"{wid}: status={status} but placeholder_status=applied "
                    "(LOW/UNRESOLVED/EXCLUDED must never be templatized)",
                )
            if ph_value != canonical_token:
                report.add(
                    "yaml_placeholder_token_missing",
                    f"{wid}: placeholder_status=applied but placeholder="
                    f"{ph_value!r} (expected {canonical_token!r})",
                )
            if wid not in docx_placeholder_ids:
                report.add(
                    "docx_placeholder_missing",
                    f"{wid}: placeholder_status=applied but no "
                    "{{ ... }} reference appears in converted_template.docx",
                )
            x = xlsx_by_id.get(wid)
            if x is not None and x.get("Placeholder") != canonical_token:
                report.add(
                    "xlsx_placeholder_token_missing",
                    f"{wid}: YAML placeholder_status=applied but XLSX "
                    f"Placeholder={x.get('Placeholder')!r}",
                )
        else:
            # Skipped rows must not carry a placeholder anywhere — that
            # is the audit guarantee for LOW/UNRESOLVED/EXCLUDED tokens.
            if ph_value not in (None, ""):
                report.add(
                    "yaml_placeholder_leak",
                    f"{wid}: placeholder_status={ph_status} but YAML "
                    f"placeholder={ph_value!r} (must be null)",
                )
            if wid in docx_placeholder_ids:
                # The regex catches compact forms too, so this fires whether
                # the leak is written as ``{{ word_NNNN }}`` or ``{{word_NNNN}}``.
                report.add(
                    "docx_placeholder_leak",
                    f"{wid}: placeholder_status={ph_status} but a "
                    "{{ ... }} reference appears in converted_template.docx",
                )
            x = xlsx_by_id.get(wid)
            if x is not None:
                xv = x.get("Placeholder")
                if xv not in (None, ""):
                    report.add(
                        "xlsx_placeholder_leak",
                        f"{wid}: placeholder_status={ph_status} but XLSX "
                        f"Placeholder={xv!r} (must be empty)",
                    )

    # Catch a stray ``{{ word_NNNN }}`` in the docx that has no YAML row
    # at all — that would mean the converted template references a Word
    # ID the audit trail never declared.
    yaml_ids = {m.get("word_id") for m in yaml_mappings if m.get("word_id")}
    orphan_in_docx = sorted(docx_placeholder_ids - yaml_ids)
    if orphan_in_docx:
        report.add(
            "docx_orphan_placeholder",
            "converted_template.docx contains placeholder(s) with no YAML "
            f"row: {', '.join(orphan_in_docx)}",
        )


def _check_report_mentions(
    yaml_doc: Dict, md_text: str, report: ValidationReport
) -> None:
    counts = ((yaml_doc.get("summary") or {}).get("by_confidence")) or {}
    for label in ("UNRESOLVED", "EXCLUDED"):
        count = counts.get(label, 0) or 0
        if count > 0 and label not in md_text:
            report.add(
                "report_missing_section",
                f"YAML summary reports {label}={count} but "
                f"confidence_report.md never mentions {label!r}",
            )
