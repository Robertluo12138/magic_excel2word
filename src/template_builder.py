"""Convert learn-mode matches into reviewable mapping artifacts.

Produces two files under the chosen output directory:

* ``auto_mapping.yml`` — one entry per visible Word numeric token. Every
  match record from the matcher appears here, including UNRESOLVED, LOW,
  and EXCLUDED rows. The entry carries ``word_id``, ``location``, the raw
  Word substring, the surrounding context, the confidence/review status,
  the recommended Excel source (when the matcher found one), and the
  transform metadata that explains how the value was matched.
* ``converted_template.docx`` — a copy of the original Word report in
  which only HIGH/MEDIUM tokens have been replaced by stable
  ``{{ word_NNNN }}`` placeholders. LOW, UNRESOLVED, and EXCLUDED values
  are deliberately left visible so a reviewer can still audit them; if
  any safe replacement could not be applied (offset mismatch, raw text
  drift) it is recorded in the YAML's ``placeholder_status`` field
  instead of being silently dropped.

This module is part of the learn-mode review loop, **not** a production
renderer. The placeholders are a contract a future deterministic renderer
must honor against a human-confirmed ``confirmed_mapping.yml``.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Tuple

import docx
import yaml

from .value_matcher import WordMatch

# HIGH and MEDIUM are the only confidence levels whose tokens are safe to
# templatize automatically. LOW carries no strong context evidence, and
# UNRESOLVED/EXCLUDED tokens have no confirmed source at all — leaving
# them visible is the audit guarantee.
SAFE_CONFIDENCES = frozenset({"HIGH", "MEDIUM"})


@dataclass
class TemplateArtifacts:
    yaml_path: Path
    docx_path: Path
    placeholders_applied: int
    placeholder_status: Dict[int, str]


def assign_word_ids(matches: List[WordMatch]) -> List[str]:
    """Zero-padded reading-order IDs (``word_0001`` … ``word_NNNN``).

    Stable across runs given the same input pair, because the underlying
    match ordering is deterministic.
    """
    return [f"word_{i + 1:04d}" for i in range(len(matches))]


def write_template_artifacts(
    matches: List[WordMatch],
    source_docx: Path,
    out_dir: Path,
) -> TemplateArtifacts:
    """Write both artifacts and return their paths plus per-match status."""
    out_dir.mkdir(parents=True, exist_ok=True)
    word_ids = assign_word_ids(matches)
    docx_out = out_dir / "converted_template.docx"
    yaml_out = out_dir / "auto_mapping.yml"

    placeholder_status = _write_converted_template(
        matches=matches,
        word_ids=word_ids,
        source_docx=source_docx,
        out_path=docx_out,
    )
    _write_auto_mapping_yaml(
        matches=matches,
        word_ids=word_ids,
        placeholder_status=placeholder_status,
        out_path=yaml_out,
    )
    return TemplateArtifacts(
        yaml_path=yaml_out,
        docx_path=docx_out,
        placeholders_applied=sum(1 for s in placeholder_status.values() if s == "applied"),
        placeholder_status=placeholder_status,
    )


# ---------------------------------------------------------------------------
# Converted .docx
# ---------------------------------------------------------------------------

def _write_converted_template(
    matches: List[WordMatch],
    word_ids: List[str],
    source_docx: Path,
    out_path: Path,
) -> Dict[int, str]:
    """Re-open the source ``.docx`` and substitute eligible tokens.

    Returns a status dict keyed by match index whose values are one of
    ``applied``, ``skipped_low_confidence``, ``skipped_unresolved``,
    ``skipped_excluded``, ``skipped_offset_out_of_range``, or
    ``skipped_raw_mismatch``.
    """
    doc = docx.Document(str(source_docx))

    # Group matches by their location string so we can apply all
    # replacements within a single paragraph/cell in one pass.
    by_location: Dict[str, List[Tuple[int, WordMatch]]] = {}
    for i, m in enumerate(matches):
        by_location.setdefault(m.word_number.location, []).append((i, m))

    placeholder_status: Dict[int, str] = {
        i: _initial_status(m.confidence) for i, m in enumerate(matches)
    }

    for p_idx, para in enumerate(doc.paragraphs):
        items = by_location.get(f"paragraph:{p_idx}", [])
        text = para.text or ""
        new_text = _replace_safe_tokens(text, items, word_ids, placeholder_status)
        if new_text != text:
            # python-docx replaces all runs with a single run containing
            # this text — the paragraph style (e.g. heading level) is
            # preserved, but inline run-level formatting is collapsed.
            # That's an accepted limitation for a learn-mode template
            # artifact; the future deterministic renderer will reapply
            # styling from a confirmed template.
            para.text = new_text

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                items = by_location.get(
                    f"table:{t_idx}/row:{r_idx}/col:{c_idx}", []
                )
                text = cell.text or ""
                new_text = _replace_safe_tokens(text, items, word_ids, placeholder_status)
                if new_text != text:
                    cell.text = new_text

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return placeholder_status


def _replace_safe_tokens(
    text: str,
    items: List[Tuple[int, WordMatch]],
    word_ids: List[str],
    placeholder_status: Dict[int, str],
) -> str:
    """Replace HIGH/MEDIUM tokens with placeholders, right-to-left.

    Right-to-left order keeps earlier offsets valid as the string mutates.
    Anything that fails an offset/raw safety check is left visible and the
    status dict is updated so the YAML reflects the skip.
    """
    eligible = [
        (i, m) for (i, m) in items if placeholder_status.get(i) == "eligible"
    ]
    eligible.sort(key=lambda x: x[1].word_number.offset, reverse=True)

    new_text = text
    for i, m in eligible:
        wn = m.word_number
        start = wn.offset
        end = start + len(wn.raw)
        if start < 0 or end > len(new_text):
            placeholder_status[i] = "skipped_offset_out_of_range"
            continue
        if new_text[start:end] != wn.raw:
            # Defensive: the paragraph text we just read disagrees with
            # what the profiler saw. Surface this so a reviewer can
            # investigate instead of letting us paste a placeholder over
            # the wrong characters.
            placeholder_status[i] = "skipped_raw_mismatch"
            continue
        placeholder = "{{ " + word_ids[i] + " }}"
        new_text = new_text[:start] + placeholder + new_text[end:]
        placeholder_status[i] = "applied"
    return new_text


def _initial_status(confidence: str) -> str:
    if confidence == "HIGH" or confidence == "MEDIUM":
        return "eligible"
    if confidence == "LOW":
        return "skipped_low_confidence"
    if confidence == "UNRESOLVED":
        return "skipped_unresolved"
    if confidence == "EXCLUDED":
        return "skipped_excluded"
    return "skipped_unknown"


# ---------------------------------------------------------------------------
# auto_mapping.yml
# ---------------------------------------------------------------------------

def _write_auto_mapping_yaml(
    matches: List[WordMatch],
    word_ids: List[str],
    placeholder_status: Dict[int, str],
    out_path: Path,
) -> Path:
    """Serialize every match (HIGH/MEDIUM/LOW/UNRESOLVED/EXCLUDED) to YAML."""
    entries = []
    for i, m in enumerate(matches):
        wn = m.word_number
        # ``chosen`` is the matcher's pick; for LOW it's still set so we
        # can show the best guess. For UNRESOLVED/EXCLUDED both ``chosen``
        # and ``candidates`` are empty and recommended_source stays null.
        top = m.chosen or (m.candidates[0] if m.candidates else None)
        status_val = placeholder_status.get(i, "skipped_unknown")

        entry: Dict = {
            "word_id": word_ids[i],
            "location": wn.location,
            "raw": wn.raw,
            "value": wn.value,
            "unit": wn.unit or "",
            "sign": wn.sign,
            "context": {
                "snippet": wn.snippet,
                "label_context": list(wn.label_context),
            },
            "status": m.confidence,
            "confidence": m.confidence,
            "review_status": _review_status(m.confidence, status_val),
            "note": m.note,
            "placeholder": (
                "{{ " + word_ids[i] + " }}" if status_val == "applied" else None
            ),
            "placeholder_status": status_val,
            "recommended_source": None,
            "transform": None,
            "alternatives": _alternatives_payload(m, top),
        }
        if top is not None:
            entry["recommended_source"] = {
                "sheet": top.cell.sheet,
                "address": top.cell.address,
                "value": top.cell.numeric_value,
                "row_context": list(top.cell.row_context[:8]),
                "column_context": list(top.cell.column_context[:8]),
            }
            entry["transform"] = {
                "interpretation": top.interpretation,
                "value_score": round(top.value_score, 3),
                "context_score": round(top.context_score, 3),
                "overlap_tokens": list(top.overlap_tokens),
            }
        entries.append(entry)

    doc_root = {
        "schema_version": 1,
        "summary": {
            "total": len(matches),
            "by_confidence": _count_by_conf(matches),
            "placeholders_applied": sum(
                1 for s in placeholder_status.values() if s == "applied"
            ),
        },
        "mappings": entries,
    }

    out_path.parent.mkdir(parents=True, exist_ok=True)
    yaml_text = yaml.safe_dump(
        doc_root,
        allow_unicode=True,
        sort_keys=False,
        width=1000,
    )
    out_path.write_text(yaml_text, encoding="utf-8")
    return out_path


def _alternatives_payload(m: WordMatch, top) -> List[Dict]:
    """Serialize every runner-up the matcher kept.

    The matcher caps ``candidates`` at the top 10 by combined score and
    always sorts so the chosen pick is at index 0. We surface everything
    after that — including ties at the top — so a reviewer auditing an
    ambiguous MEDIUM (or wondering why a LOW landed there) can see what
    else the matcher considered without re-running the pipeline.
    """
    cands = list(m.candidates) if m.candidates else []
    # Drop the candidate that's already exposed as recommended_source.
    if top is not None and cands and cands[0] is top:
        rest = cands[1:]
    elif top is not None:
        rest = [c for c in cands if c is not top]
    else:
        rest = cands
    return [
        {
            "sheet": c.cell.sheet,
            "address": c.cell.address,
            "value": c.cell.numeric_value,
            "row_context": list(c.cell.row_context[:8]),
            "column_context": list(c.cell.column_context[:8]),
            "interpretation": c.interpretation,
            "value_score": round(c.value_score, 3),
            "context_score": round(c.context_score, 3),
            "overlap_tokens": list(c.overlap_tokens),
        }
        for c in rest
    ]


def _count_by_conf(matches: List[WordMatch]) -> Dict[str, int]:
    counts = {"HIGH": 0, "MEDIUM": 0, "LOW": 0, "UNRESOLVED": 0, "EXCLUDED": 0}
    for m in matches:
        counts[m.confidence] = counts.get(m.confidence, 0) + 1
    return counts


def _review_status(confidence: str, placeholder_status: str) -> str:
    """Single review-status enum derived from confidence + template outcome.

    ``HIGH``/``MEDIUM`` with an applied placeholder → ``pending_review``
    (auto-mapped, awaiting human sign-off). The same confidence levels
    with a skipped placeholder fall back to ``needs_review`` so the
    reviewer is alerted to the offset/raw drift. ``LOW`` is always
    ``needs_review``; ``UNRESOLVED`` is ``needs_source``; ``EXCLUDED``
    is ``audited_excluded``.
    """
    if confidence in SAFE_CONFIDENCES:
        return "pending_review" if placeholder_status == "applied" else "needs_review"
    if confidence == "LOW":
        return "needs_review"
    if confidence == "UNRESOLVED":
        return "needs_source"
    if confidence == "EXCLUDED":
        return "audited_excluded"
    return "needs_review"
