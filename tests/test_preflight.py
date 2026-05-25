"""Focused tests for the privacy preflight advisory.

The guard is intentionally narrow: it never reads file contents, never
blocks execution, and never alters matching/confirmation/render/validate
logic. These tests pin four guarantees:

  1. ``check_paths`` flags inputs/outputs that resolve under
     ``<repo>/samples/`` and ignores paths anywhere else.
  2. ``format_advisory``/``emit_advisory`` produce a stderr-shaped
     block that names every offending flag and points at the pilot doc.
  3. The CLI surfaces the advisory but still exits with the same code
     it would have without the advisory — i.e. matching/render logic
     is untouched.
  4. Every pilot-sequence subcommand (``learn``, ``validate-artifacts``,
     ``confirm-mapping``, ``run-preview``, ``render-docx``,
     ``validate-render``) checks **every** CLI path flag the pilot doc
     §5 promises it does — no quiet gaps.
"""
from __future__ import annotations

import io
import shutil
from pathlib import Path
from typing import Tuple

import docx
import openpyxl

from src.main import main as cli_main
from src.preflight import (
    check_paths,
    emit_advisory,
    format_advisory,
)
from src.synthetic_generator import generate


REPO_ROOT = Path(__file__).resolve().parents[1]
SAMPLES_DIR = REPO_ROOT / "samples"


# ---------------------------------------------------------------------------
# check_paths — pure detection
# ---------------------------------------------------------------------------

def test_check_paths_returns_empty_when_paths_are_outside_samples(tmp_path: Path):
    warnings = check_paths(
        [
            ("--excel", tmp_path / "pilot" / "historical.xlsx"),
            ("--word", tmp_path / "pilot" / "report.docx"),
            ("--out", tmp_path / "pilot" / "output"),
        ]
    )
    assert warnings == []


def test_check_paths_flags_paths_inside_samples_directory():
    """A path resolving under ``<repo>/samples/`` must be flagged, even
    when the file does not yet exist. ``Path.resolve()`` handles missing
    targets in Python 3.6+, so this works for both inputs (which exist)
    and outputs (which may not)."""
    warnings = check_paths(
        [
            ("--excel", SAMPLES_DIR / "synthetic" / "historical.xlsx"),
            ("--word", SAMPLES_DIR / "synthetic" / "finished_report.docx"),
        ]
    )
    assert len(warnings) == 2
    assert all("samples/" in w for w in warnings)
    assert any("--excel" in w for w in warnings)
    assert any("--word" in w for w in warnings)


def test_check_paths_flags_output_paths_too():
    """Output paths matter as much as inputs: pointing ``--out`` at
    ``samples/`` would mix generated artifacts into the synthetic
    fixture folder."""
    warnings = check_paths(
        [("--out", SAMPLES_DIR / "synthetic" / "pilot_output")]
    )
    assert len(warnings) == 1
    assert "--out" in warnings[0]


def test_check_paths_ignores_none_entries(tmp_path: Path):
    """Callers should be able to pass optional CLI args straight through;
    ``None`` must not crash and must not produce a warning."""
    warnings = check_paths(
        [
            ("--excel", None),
            ("--out", tmp_path / "elsewhere"),
        ]
    )
    assert warnings == []


def test_check_paths_handles_relative_paths(monkeypatch, tmp_path: Path):
    """A relative path like ``samples/synthetic/x.xlsx`` invoked from
    the repo root must resolve to the same place as the absolute form,
    so the guard fires regardless of cwd."""
    monkeypatch.chdir(REPO_ROOT)
    warnings = check_paths([("--excel", Path("samples/synthetic/historical.xlsx"))])
    assert len(warnings) == 1
    assert "samples/" in warnings[0]


# ---------------------------------------------------------------------------
# format_advisory / emit_advisory — presentation
# ---------------------------------------------------------------------------

def test_format_advisory_is_empty_when_no_warnings():
    assert format_advisory([]) == ""


def test_format_advisory_lists_warnings_and_points_at_pilot_doc():
    text = format_advisory(["--excel samples/foo.xlsx is inside ..."])
    assert "PRIVACY PREFLIGHT ADVISORY" in text
    assert "--excel samples/foo.xlsx" in text
    assert "OUTSIDE" in text
    assert "docs/real_file_pilot.md" in text


def test_emit_advisory_writes_to_stream_and_returns_true():
    stream = io.StringIO()
    fired = emit_advisory(
        [("--excel", SAMPLES_DIR / "synthetic" / "historical.xlsx")],
        stream,
    )
    assert fired is True
    output = stream.getvalue()
    assert "PRIVACY PREFLIGHT ADVISORY" in output
    assert "samples/" in output


def test_emit_advisory_is_silent_on_clean_paths(tmp_path: Path):
    stream = io.StringIO()
    fired = emit_advisory(
        [("--excel", tmp_path / "pilot.xlsx")],
        stream,
    )
    assert fired is False
    assert stream.getvalue() == ""


# ---------------------------------------------------------------------------
# CLI integration — advisory is informational, never blocks
# ---------------------------------------------------------------------------

def test_learn_cli_emits_advisory_when_inputs_live_under_samples(
    tmp_path: Path, capsys
):
    """When the operator runs ``learn`` against paths inside the repo's
    ``samples/`` folder, the advisory must appear in stderr. The
    underlying learn run must still execute and exit with its normal
    code — the guard only informs, never blocks."""
    # Build the synthetic fixture inside the repo's samples/ so the
    # resolved paths land under <repo>/samples/ and trigger the guard.
    # We use a session-scoped subdirectory so a partial test run can be
    # cleaned by hand if needed; it's git-ignored under samples/synthetic.
    fixture_dir = SAMPLES_DIR / "synthetic" / "_preflight_cli_check"
    try:
        generate(fixture_dir)
        out = tmp_path / "output"

        rc = cli_main(
            [
                "learn",
                "--excel", str(fixture_dir / "historical.xlsx"),
                "--word", str(fixture_dir / "finished_report.docx"),
                "--out", str(out),
            ]
        )

        captured = capsys.readouterr()
        # The synthetic corpus has deliberate UNRESOLVED rows, so default
        # mode returns 0 with a warning. The advisory must NOT change
        # that — same exit code as without the advisory.
        assert rc == 0
        assert "PRIVACY PREFLIGHT ADVISORY" in captured.err
        assert "--excel" in captured.err
        assert "--word" in captured.err
        # Underlying learn pipeline still ran and wrote its artifacts.
        assert (out / "mapping_review.xlsx").exists()
        assert (out / "auto_mapping.yml").exists()
    finally:
        # Clean the helper fixture so we don't leave files inside the
        # tracked repo tree.
        import shutil
        if fixture_dir.exists():
            shutil.rmtree(fixture_dir)


def test_learn_cli_does_not_emit_advisory_for_outside_paths(
    tmp_path: Path, capsys
):
    """Mirror of the previous test: when inputs live outside the repo
    (the normal pilot path), no advisory fires."""
    samples = tmp_path / "pilot_inputs"
    generate(samples)
    out = tmp_path / "pilot_output"

    rc = cli_main(
        [
            "learn",
            "--excel", str(samples / "historical.xlsx"),
            "--word", str(samples / "finished_report.docx"),
            "--out", str(out),
        ]
    )

    captured = capsys.readouterr()
    assert rc == 0
    assert "PRIVACY PREFLIGHT ADVISORY" not in captured.err
    assert (out / "mapping_review.xlsx").exists()


# ---------------------------------------------------------------------------
# Pilot doc §5 contract — every documented flag on every documented
# command must actually be checked. The risk this pins is the original
# review finding: the doc claims a command checks "any of the paths you
# passed", but the wiring only listed a subset. A quiet gap there would
# let a real-file pilot path slip through unwarned.
# ---------------------------------------------------------------------------


def _build_clean_pair(samples_dir: Path) -> Tuple[Path, Path]:
    """Tiny clean Excel + Word pair with zero eligible UNRESOLVED/LOW rows.

    Mirrors ``tests/test_acceptance_smoke.py::_build_clean_pair`` so the
    confirm-mapping gate can reach ``complete: true`` without the
    ``--allow-incomplete`` escape hatch — that lets us drive the full
    pilot sequence end-to-end and assert the advisory fires at every
    stage with every documented flag listed.
    """
    samples_dir.mkdir(parents=True, exist_ok=True)
    xlsx = samples_dir / "clean.xlsx"
    docx_path = samples_dir / "clean.docx"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "月度核心指标"
    ws.append(["指标", "2026年5月"])
    ws.append(["营业收入(元)", 123_456_789.00])
    ws.append(["毛利率(%)", 36.20])
    ws.append(["用户总数(人)", 13_456_789])
    wb.save(str(xlsx))

    doc = docx.Document()
    doc.add_heading("2026年5月经营月度报告", level=1)
    doc.add_paragraph("本月营业收入达12,345.68万元。")
    doc.add_paragraph("毛利率达36.20%。")
    doc.add_paragraph("本月用户总数达到1,345.68万人。")
    doc.save(str(docx_path))
    return xlsx, docx_path


def _confirm_every_high_medium(review_xlsx: Path) -> None:
    wb = openpyxl.load_workbook(str(review_xlsx))
    ws = wb.active
    header = [c.value for c in ws[1]]
    conf_col = header.index("Confidence") + 1
    decision_col = header.index("Reviewer Decision") + 1
    for r in range(2, ws.max_row + 1):
        if ws.cell(row=r, column=conf_col).value in ("HIGH", "MEDIUM"):
            ws.cell(row=r, column=decision_col).value = "confirm"
    wb.save(str(review_xlsx))


def test_advisory_fires_even_when_existence_check_returns_rc_2(
    tmp_path: Path, capsys
):
    """The advisory must fire on the typo'd-path → rc=2 path too.

    If the advisory ran after the existence check, an operator who
    typo'd ``samples/synthetic/missing.xlsx`` would get an unhelpful
    ``error: excel file not found`` and never see that their intended
    path was inside the repo. The fix is to run the advisory first;
    this test pins that ordering so a future reorder cannot regress it.

    Uses ``learn`` as the representative; the same ordering applies to
    every instrumented command.
    """
    nonexistent_under_samples = (
        SAMPLES_DIR / "synthetic" / "_preflight_rc2_check_does_not_exist.xlsx"
    )
    # Sanity guard — if a previous run somehow left this fixture in
    # place, the test premise (a missing path) is no longer being
    # exercised. The same path is used in the assertion below.
    assert not nonexistent_under_samples.exists()

    out = tmp_path / "output"
    rc = cli_main([
        "learn",
        "--excel", str(nonexistent_under_samples),
        "--word", str(tmp_path / "also_missing.docx"),
        "--out", str(out),
    ])
    captured = capsys.readouterr()

    # Exit code unchanged: existence check still owns rc=2.
    assert rc == 2
    assert "excel file not found" in captured.err

    # And the advisory still fires for the samples/-internal --excel.
    assert "PRIVACY PREFLIGHT ADVISORY" in captured.err
    assert "--excel" in captured.err


def test_every_pilot_command_advisory_covers_every_documented_flag(
    capsys,
):
    """Drive the full pilot sequence with **every** path under
    ``samples/`` and assert each command's advisory mentions **every**
    flag the pilot doc §5 promises it checks.

    Pinned flags (matches ``docs/real_file_pilot.md`` §5):

    * ``learn``               — ``--excel``, ``--word``, ``--out``
    * ``validate-artifacts``  — ``--out``
    * ``confirm-mapping``     — ``--auto``, ``--review``, ``--out``
    * ``run-preview``         — ``--excel``, ``--confirmed``, ``--out``
    * ``render-docx``         — ``--template``, ``--run-validation``, ``--out``
    * ``validate-render``     — ``--docx``, ``--render-log``, ``--run-validation``

    The fixture lives under ``samples/synthetic/_preflight_full_check``
    (git-ignored by the existing ``samples/synthetic/`` rule) and is
    removed in the ``finally`` so a partial run never leaves files inside
    the repo tree.
    """
    fixture_dir = SAMPLES_DIR / "synthetic" / "_preflight_full_check"
    try:
        inputs = fixture_dir / "inputs"
        out_dir = fixture_dir / "output"
        preview_dir = out_dir / "run_preview"
        xlsx, word = _build_clean_pair(inputs)

        # 1) learn — must list --excel, --word, --out
        rc = cli_main([
            "learn",
            "--excel", str(xlsx),
            "--word", str(word),
            "--out", str(out_dir),
        ])
        assert rc == 0
        err = capsys.readouterr().err
        assert "PRIVACY PREFLIGHT ADVISORY" in err
        for flag in ("--excel", "--word", "--out"):
            assert flag in err, f"learn advisory missing {flag}"

        # 2) validate-artifacts — must list --out
        rc = cli_main(["validate-artifacts", "--out", str(out_dir)])
        assert rc == 0
        err = capsys.readouterr().err
        assert "PRIVACY PREFLIGHT ADVISORY" in err
        assert "--out" in err

        # 3) confirm-mapping — must list --auto, --review, --out
        _confirm_every_high_medium(out_dir / "mapping_review.xlsx")
        confirmed_yaml = out_dir / "confirmed_mapping.yml"
        rc = cli_main([
            "confirm-mapping",
            "--auto", str(out_dir / "auto_mapping.yml"),
            "--review", str(out_dir / "mapping_review.xlsx"),
            "--out", str(confirmed_yaml),
        ])
        assert rc == 0
        err = capsys.readouterr().err
        assert "PRIVACY PREFLIGHT ADVISORY" in err
        for flag in ("--auto", "--review", "--out"):
            assert flag in err, f"confirm-mapping advisory missing {flag}"

        # 4) run-preview — must list --excel, --confirmed, --out
        rc = cli_main([
            "run-preview",
            "--excel", str(xlsx),
            "--confirmed", str(confirmed_yaml),
            "--out", str(preview_dir),
        ])
        assert rc == 0
        err = capsys.readouterr().err
        assert "PRIVACY PREFLIGHT ADVISORY" in err
        for flag in ("--excel", "--confirmed", "--out"):
            assert flag in err, f"run-preview advisory missing {flag}"

        # 5) render-docx — must list --template, --run-validation, --out
        new_report = out_dir / "new_report.docx"
        rc = cli_main([
            "render-docx",
            "--template", str(out_dir / "converted_template.docx"),
            "--run-validation", str(preview_dir / "run_validation.xlsx"),
            "--out", str(new_report),
        ])
        assert rc == 0
        err = capsys.readouterr().err
        assert "PRIVACY PREFLIGHT ADVISORY" in err
        for flag in ("--template", "--run-validation", "--out"):
            assert flag in err, f"render-docx advisory missing {flag}"

        # 6) validate-render — must list --docx, --render-log, --run-validation
        rc = cli_main([
            "validate-render",
            "--docx", str(new_report),
            "--render-log", str(out_dir / "render_log.yml"),
            "--run-validation", str(preview_dir / "run_validation.xlsx"),
        ])
        assert rc == 0
        err = capsys.readouterr().err
        assert "PRIVACY PREFLIGHT ADVISORY" in err
        for flag in ("--docx", "--render-log", "--run-validation"):
            assert flag in err, f"validate-render advisory missing {flag}"
    finally:
        if fixture_dir.exists():
            shutil.rmtree(fixture_dir)
