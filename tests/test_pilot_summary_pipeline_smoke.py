"""Integration smoke: ``pilot-summary`` over real pipeline artifacts.

``tests/test_pilot_summary.py`` builds artifacts by hand with sentinel
strings to pin the redaction contract — fast and exhaustive on the
contract surface, but it doesn't prove the summarizer can actually read
what the real pipeline emits. ``tests/test_acceptance_smoke.py`` drives
the six-stage pipeline to a clean rendered docx but never invokes
``pilot-summary`` on the result.

This file fills the gap with **one** narrow smoke: build a clean
synthetic Excel + Word pair, drive the full pipeline through the CLI
(``learn --strict`` → ``validate-artifacts`` → ``confirm-mapping`` →
``run-preview`` → ``render-docx`` → ``validate-render``), then run
``pilot-summary --out`` against the resulting output directory.

What it proves:

* ``pilot-summary`` exits ``0`` on a real, clean pipeline run.
* The printed summary surfaces useful aggregate per-stage information
  so an operator can see at a glance that every stage ran.
* The printed summary does NOT leak any real content from the run:
  raw Word display tokens, raw Word snippets, source sheet/cell
  details, individual ``word_id`` values, generated numeric values,
  or full filesystem paths must not appear.

Synthetic-only, repo-safe: no real data, no LLM, no GUI, no network,
no Microsoft Office automation. The clean fixture below mirrors the one
in ``test_acceptance_smoke.py`` — kept inline rather than imported so a
refactor of either test file does not silently couple the two.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Tuple

import docx
import openpyxl

from src.main import main as cli_main


def _build_clean_pair(samples_dir: Path) -> Tuple[Path, Path]:
    """Write a paired Excel + Word with zero eligible UNRESOLVED/LOW numbers.

    Mirrors the fixture in ``tests/test_acceptance_smoke.py`` so the
    pipeline reaches a clean ``complete: true`` without
    ``--allow-incomplete``. Three HIGH/MEDIUM metrics (revenue in 万元,
    margin %, users in 万人) plus the two EXCLUDED date markers from the
    heading.
    """
    samples_dir.mkdir(parents=True, exist_ok=True)
    xlsx = samples_dir / "clean.xlsx"
    docx_path = samples_dir / "clean.docx"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "月度核心指标"
    ws.append(["指标", "2026年5月"])
    ws.append(["营业收入(元)", 123_456_789.00])
    ws.append(["毛利率(%)", 36.20])
    ws.append(["用户总数(人)", 13_456_789])
    wb.save(str(xlsx))

    doc = docx.Document()
    doc.add_heading("2026年5月经营月度报告", level=1)
    doc.add_paragraph("本月营业收入达12,345.68万元。")
    doc.add_paragraph("毛利率达36.20%。")
    doc.add_paragraph("本月用户总数达到1,345.68万人。")
    doc.save(str(docx_path))
    return xlsx, docx_path


def _fill_reviewer_decisions(review_xlsx: Path) -> None:
    """Mark every HIGH/MEDIUM row in mapping_review.xlsx as ``confirm``."""
    wb = openpyxl.load_workbook(str(review_xlsx))
    ws = wb.active
    header = [c.value for c in ws[1]]
    conf_col = header.index("Confidence") + 1
    decision_col = header.index("Reviewer Decision") + 1
    for r in range(2, ws.max_row + 1):
        if ws.cell(row=r, column=conf_col).value in ("HIGH", "MEDIUM"):
            ws.cell(row=r, column=decision_col).value = "confirm"
    wb.save(str(review_xlsx))


def test_pilot_summary_on_real_pipeline_artifacts(tmp_path: Path, capsys):
    samples = tmp_path / "samples"
    out = tmp_path / "out"
    xlsx, word = _build_clean_pair(samples)

    # Drive the full pipeline through the CLI surface so the artifacts
    # under test come from the real deterministic pipeline, not hand-built
    # YAML/XLSX with the right shape.
    assert cli_main([
        "learn", "--strict",
        "--excel", str(xlsx), "--word", str(word), "--out", str(out),
    ]) == 0
    assert cli_main(["validate-artifacts", "--out", str(out)]) == 0
    _fill_reviewer_decisions(out / "mapping_review.xlsx")
    confirmed = out / "confirmed_mapping.yml"
    assert cli_main([
        "confirm-mapping",
        "--auto", str(out / "auto_mapping.yml"),
        "--review", str(out / "mapping_review.xlsx"),
        "--out", str(confirmed),
    ]) == 0
    assert cli_main([
        "run-preview",
        "--excel", str(xlsx),
        "--confirmed", str(confirmed),
        "--out", str(out),
    ]) == 0
    rendered = out / "new_report.docx"
    assert cli_main([
        "render-docx",
        "--template", str(out / "converted_template.docx"),
        "--run-validation", str(out / "run_validation.xlsx"),
        "--out", str(rendered),
    ]) == 0
    assert cli_main([
        "validate-render",
        "--docx", str(rendered),
        "--render-log", str(out / "render_log.yml"),
        "--run-validation", str(out / "run_validation.xlsx"),
    ]) == 0

    # Discard captured stdout/stderr from the prior pipeline stages so
    # only the pilot-summary output is asserted against.
    capsys.readouterr()

    rc = cli_main(["pilot-summary", "--out", str(out)])
    captured = capsys.readouterr()
    assert rc == 0, captured.err
    text = captured.out

    # --- Positive shape: every stage was reached and aggregate info shows.
    # The header reduces ``--out`` to its basename for safe sharing.
    assert "pilot-summary for: out" in text
    assert "[learn] [present]" in text
    assert "[confirm-mapping] [present]" in text
    assert "[run-preview] [present]" in text
    assert "[render-docx] [present]" in text
    # Aggregate learn-stage counts from auto_mapping.yml's summary block.
    assert "total Word numbers" in text
    assert "HIGH" in text
    assert "EXCLUDED" in text
    # Aggregate confirm-stage flags lifted from confirmed_mapping.yml.
    assert "summary.complete" in text and "true" in text
    assert "summary.allow_incomplete" in text and "false" in text
    # Aggregate run-preview status counts (status enums are policy strings,
    # not raw business content — these are explicitly allowed to surface).
    assert "  ok" in text
    # Aggregate render-stage counts from render_log.yml's summary block.
    assert "log rows" in text
    assert "total placeholder replacements" in text
    assert "rendered docx files" in text

    # --- Negative redaction: not a single piece of real content from the
    # run may surface in the pilot-summary output. Each block below pins
    # one category from the docstring contract in src/pilot_summary.py.

    # (a) Raw metric display values (the final docx contains these — the
    # summary must NOT echo them).
    for display in ("12,345.68万元", "36.20%", "1,345.68万人"):
        assert display not in text, (
            f"display value {display!r} leaked into pilot-summary output"
        )

    # (b) Raw Word snippets (paragraph text and heading) from the
    # historical report.
    for snippet in (
        "本月营业收入达",
        "毛利率达",
        "本月用户总数达到",
        "经营月度报告",
    ):
        assert snippet not in text, (
            f"Word snippet {snippet!r} leaked into pilot-summary output"
        )

    # (c) Source sheet/cell details for confirmed mappings.
    assert "月度核心指标" not in text, "Excel sheet name leaked"
    # B2/B3/B4 are the natural source-cell addresses for the three
    # metrics in the clean fixture; none should reach the formatter.
    for cell in ("B2", "B3", "B4"):
        assert cell not in text, f"Excel cell address {cell!r} leaked"

    # (d) Individual ``word_id`` values. The pipeline assigns stable
    # ``word_NNNN`` identifiers throughout the audit artifacts; the
    # summary may print the literal label ``word_ids`` (e.g. "distinct
    # placeholder word_ids") but never an individual ``word_0001``.
    assert re.search(r"word_\d{3,}", text) is None, (
        "individual word_id leaked into pilot-summary output"
    )

    # (e) Generated and raw numeric values. ``run_validation.xlsx``'s
    # ``Generated Value`` column and ``render_log.yml``'s
    # ``generated_value`` / ``raw_excel_value`` fields hold the actual
    # business numbers; none of those forms (rounded display value,
    # unrounded generated value, raw Excel integer) may surface.
    for numeric in (
        "12345.68",     # display-rounded revenue without comma
        "12345.6789",   # generated revenue after 万元→base_unit
        "123456789",    # raw Excel revenue integer
        "1345.68",      # display-rounded users without comma
        "1345.6789",    # generated users after 万人→base_unit
        "13456789",     # raw Excel users integer
        "36.20",        # display and Excel margin
    ):
        assert numeric not in text, (
            f"numeric value {numeric!r} leaked into pilot-summary output"
        )

    # (f) Full filesystem paths. The operator's parent tree (home dir,
    # project codename, customer label) must reduce to the basename of
    # ``--out`` in the header, and no other absolute path may appear.
    assert str(out) not in text, "full --out path leaked"
    assert str(tmp_path) not in text, "tmp_path leaked"
    assert str(samples) not in text, "synthetic samples path leaked"
    assert str(xlsx) not in text, "input Excel path leaked"
    assert str(word) not in text, "input Word path leaked"
