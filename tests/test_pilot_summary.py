"""Focused tests for the read-only redacted pilot summary command.

The summary is the only artifact a real-file pilot operator might
paste back into a chat, ticket, or commit — so the redaction contract
is the load-bearing claim. Each test below either:

  * asserts the redaction contract directly by planting unique
    sentinel strings in every content field of every artifact and
    asserting NONE of them surface in the formatter output, OR
  * asserts the missing-artifact and exit-code behavior the pilot
    doc §2c promises.

The tests do NOT process or generate any real-data file. Artifacts are
constructed in ``tmp_path`` with the minimum schema each
summarizer-stage reads.
"""
from __future__ import annotations

import shutil
from pathlib import Path
from typing import Dict, Optional

import docx
import openpyxl
import yaml

from src.main import main as cli_main
from src.pilot_summary import (
    PilotSummary,
    StageStatus,
    format_summary,
    summarize_pilot,
)


REPO_ROOT = Path(__file__).resolve().parents[1]
SAMPLES_DIR = REPO_ROOT / "samples"


# ---------------------------------------------------------------------------
# Sentinels — unique strings the summarizer must NEVER echo.
# Placed in every content field of every artifact so a redaction
# regression in ANY stage surfaces as a single failing assertion.
# ---------------------------------------------------------------------------

_S_RAW_TOKEN = "ZXSENTINEL_RAW_TOKEN_ZX"
_S_SHEET = "ZXSENTINEL_SHEET_ZX"
_S_CELL = "ZXSENTINEL_CELL_ZX"
_S_NOTE = "ZXSENTINEL_NOTE_ZX"
_S_DETAIL = "ZXSENTINEL_DETAIL_ZX"
_S_WORDID = "wordzxsentinel0042zx"  # YAML-safe but unmistakable
_S_PATH = "/ZXSENTINEL_PATH_ZX/secret"
_S_LOCATION = "ZXSENTINEL_LOCATION_ZX"
_S_VALUE_INT = 8888888888  # int part is what would appear in str(float)
_S_VALUE = float(_S_VALUE_INT) + 0.77

_SENTINELS = [
    _S_RAW_TOKEN, _S_SHEET, _S_CELL, _S_NOTE, _S_DETAIL,
    _S_WORDID, _S_PATH, _S_LOCATION, str(_S_VALUE_INT),
]


# ---------------------------------------------------------------------------
# Artifact builders — produce a minimal-schema file with sentinels in
# every content field but valid aggregate counts in the summary blocks.
# ---------------------------------------------------------------------------

def _write_auto_mapping(
    path: Path,
    *,
    total: int = 10,
    by_conf: Optional[Dict[str, int]] = None,
    placeholders_applied: Optional[int] = None,
) -> None:
    by_conf = by_conf or {
        "HIGH": 7, "MEDIUM": 2, "LOW": 0, "UNRESOLVED": 0, "EXCLUDED": 1,
    }
    if placeholders_applied is None:
        placeholders_applied = by_conf["HIGH"] + by_conf["MEDIUM"]
    doc = {
        "schema_version": 1,
        "summary": {
            "total": total,
            "by_confidence": by_conf,
            "placeholders_applied": placeholders_applied,
        },
        "mappings": [
            {
                "word_id": _S_WORDID,
                "location": _S_LOCATION,
                "raw": _S_RAW_TOKEN,
                "value": _S_VALUE,
                "unit": "万元",
                "context": {"snippet": _S_NOTE, "label_context": [_S_NOTE]},
                "status": "HIGH",
                "confidence": "HIGH",
                "review_status": "pending_review",
                "placeholder": "{{ " + _S_WORDID + " }}",
                "placeholder_status": "applied",
                "recommended_source": {
                    "sheet": _S_SHEET,
                    "address": _S_CELL,
                    "value": _S_VALUE,
                },
                "transform": {"interpretation": "万元→base_unit"},
                "alternatives": [],
                "note": _S_NOTE,
            },
        ],
    }
    path.write_text(
        yaml.safe_dump(doc, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )


def _write_mapping_review(path: Path, *, n: int = 10) -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append([
        "Word ID", "Word Location", "Word Snippet", "Word Raw Token",
        "Word Value", "Word Unit", "Confidence", "Review Status",
        "Placeholder Status", "Placeholder", "Top Excel Sheet",
        "Top Excel Cell", "Top Excel Value",
        "Reviewer Decision", "Reviewer Notes",
    ])
    for i in range(n):
        ws.append([
            _S_WORDID + f"_row{i}",
            _S_LOCATION, _S_NOTE, _S_RAW_TOKEN, _S_VALUE,
            "万元", "HIGH", "pending_review",
            "applied", "{{ " + _S_WORDID + " }}",
            _S_SHEET, _S_CELL, _S_VALUE,
            "confirm", _S_NOTE,
        ])
    wb.save(str(path))


def _write_confidence_report(path: Path) -> None:
    path.write_text(
        f"# learn-mode\n\nUNRESOLVED\n{_S_NOTE}\n",
        encoding="utf-8",
    )


def _write_converted_template(path: Path) -> None:
    d = docx.Document()
    d.add_paragraph(f"placeholder body {_S_RAW_TOKEN}")
    d.save(str(path))


def _write_confirmed_mapping(
    path: Path,
    *,
    complete: bool = True,
    allow_incomplete: bool = False,
    confirmed: int = 9,
    review_required: int = 0,
    audit_only: int = 1,
) -> None:
    doc = {
        "schema_version": 1,
        "source_artifacts": {
            "auto_mapping": _S_PATH + "/auto_mapping.yml",
            "mapping_review": _S_PATH + "/mapping_review.xlsx",
        },
        "summary": {
            "total_word_numbers": confirmed + review_required + audit_only,
            "confirmed": confirmed,
            "review_required": review_required,
            "audit_only_excluded": audit_only,
            "allow_incomplete": allow_incomplete,
            "complete": complete,
        },
        "confirmed_mappings": [
            {
                "word_id": _S_WORDID,
                "raw": _S_RAW_TOKEN,
                "value": _S_VALUE,
                "confirmed_source": {
                    "sheet": _S_SHEET,
                    "address": _S_CELL,
                    "value": _S_VALUE,
                },
                "reviewer_notes": _S_NOTE,
            },
        ],
        "review_required": (
            [{"word_id": _S_WORDID, "raw": _S_RAW_TOKEN, "reviewer_notes": _S_NOTE}]
            * max(0, review_required)
        ),
        "audit_only_excluded": [],
    }
    path.write_text(
        yaml.safe_dump(doc, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )


def _write_run_validation(
    path: Path,
    *,
    ok_count: int = 9,
    failed: Optional[Dict[str, int]] = None,
) -> None:
    failed = failed or {}
    path.parent.mkdir(parents=True, exist_ok=True)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append([
        "Word ID", "Word Location", "Word Context", "Word Raw Token",
        "Word Unit", "Source Sheet", "Source Cell", "Raw Excel Value",
        "Generated Value", "Transform Interpretation", "Confidence",
        "Status", "Detail",
    ])
    for i in range(ok_count):
        ws.append([
            _S_WORDID + f"_ok{i}",
            _S_LOCATION, _S_NOTE, _S_RAW_TOKEN, "万元",
            _S_SHEET, _S_CELL, _S_VALUE, _S_VALUE,
            "万元→base_unit", "HIGH", "ok", _S_DETAIL,
        ])
    for status, count in failed.items():
        for i in range(count):
            ws.append([
                _S_WORDID + f"_{status}{i}",
                _S_LOCATION, _S_NOTE, _S_RAW_TOKEN, "万元",
                _S_SHEET, _S_CELL, _S_VALUE, _S_VALUE,
                "万元→base_unit", "HIGH", status, _S_DETAIL,
            ])
    wb.save(str(path))


def _write_render_log(
    path: Path,
    *,
    total: int = 9,
    ok: int = 9,
    failed: int = 0,
    total_replacements: int = 12,
) -> None:
    doc = {
        "schema_version": 1,
        "inputs": {
            "template": _S_PATH + "/converted_template.docx",
            "run_validation": _S_PATH + "/run_validation.xlsx",
        },
        "summary": {
            "total_rows": total,
            "ok": ok,
            "failed": failed,
            "total_replacements": total_replacements,
            "distinct_placeholder_word_ids": total,
        },
        "out_docx": _S_PATH + "/new_report.docx",
        "replacements": [
            {
                "word_id": _S_WORDID,
                "source_sheet": _S_SHEET,
                "source_cell": _S_CELL,
                "raw_excel_value": _S_VALUE,
                "generated_value": _S_VALUE,
                "raw_token": _S_RAW_TOKEN,
                "unit": "万元",
                "display_text": _S_RAW_TOKEN,
                "placeholder_occurrences": 1,
                "status": "ok",
                "detail": _S_DETAIL,
            },
        ],
    }
    path.write_text(
        yaml.safe_dump(doc, allow_unicode=True, sort_keys=False),
        encoding="utf-8",
    )


def _write_rendered_docx(path: Path) -> None:
    d = docx.Document()
    d.add_paragraph(f"rendered: {_S_RAW_TOKEN}")
    d.save(str(path))


def _build_full_pilot(pilot_out: Path) -> None:
    pilot_out.mkdir(parents=True, exist_ok=True)
    _write_auto_mapping(pilot_out / "auto_mapping.yml")
    _write_mapping_review(pilot_out / "mapping_review.xlsx")
    _write_confidence_report(pilot_out / "confidence_report.md")
    _write_converted_template(pilot_out / "converted_template.docx")
    _write_confirmed_mapping(pilot_out / "confirmed_mapping.yml")
    _write_run_validation(pilot_out / "run_preview" / "run_validation.xlsx")
    _write_render_log(pilot_out / "render_log.yml")
    _write_rendered_docx(pilot_out / "new_report.docx")


def _assert_no_sentinels(text: str) -> None:
    for needle in _SENTINELS:
        assert needle not in text, (
            f"redaction failed: pilot-summary printed sentinel {needle!r}"
        )


# ---------------------------------------------------------------------------
# Missing-artifact and CLI exit-code behavior
# ---------------------------------------------------------------------------

def test_cli_missing_out_dir_returns_2(tmp_path: Path, capsys):
    rc = cli_main(["pilot-summary", "--out", str(tmp_path / "does_not_exist")])
    assert rc == 2
    err = capsys.readouterr().err
    assert "not found" in err


def test_cli_out_is_a_file_returns_2(tmp_path: Path, capsys):
    bad = tmp_path / "actually_a_file.txt"
    bad.write_text("hi", encoding="utf-8")
    rc = cli_main(["pilot-summary", "--out", str(bad)])
    assert rc == 2
    err = capsys.readouterr().err
    assert "must be a directory" in err


def test_summarize_pilot_fatal_when_auto_mapping_absent(tmp_path: Path):
    out = tmp_path / "empty_pilot"
    out.mkdir()
    summary = summarize_pilot(out)
    assert summary.ok is False
    assert len(summary.fatal_errors) == 1
    msg = summary.fatal_errors[0]
    assert "auto_mapping.yml" in msg
    assert "empty_pilot" in msg
    # Full path must NOT appear in the fatal message; only the basename.
    assert str(out) not in msg
    # Useful next-action hint.
    assert "run `learn`" in msg


def test_cli_missing_auto_mapping_returns_11(tmp_path: Path, capsys):
    out = tmp_path / "empty_pilot"
    out.mkdir()
    rc = cli_main(["pilot-summary", "--out", str(out)])
    assert rc == 11
    err = capsys.readouterr().err
    assert "auto_mapping.yml" in err
    assert "run `learn`" in err


def test_partial_pilot_marks_missing_stages_not_yet_run(tmp_path: Path):
    out = tmp_path / "partial_pilot"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    summary = summarize_pilot(out)
    assert summary.fatal_errors == []
    assert summary.learn.present is True
    assert summary.confirm.present is False
    assert summary.run_preview.present is False
    assert summary.render.present is False
    text = format_summary(summary)
    assert "[learn] [present]" in text
    assert "[confirm-mapping] [not yet run]" in text
    assert "[run-preview] [not yet run]" in text
    assert "[render-docx] [not yet run]" in text


def test_learn_missing_companion_artifacts_are_flagged(tmp_path: Path):
    """auto_mapping.yml alone is enough to summarize, but the formatter
    must still call out which learn-mode companions are missing."""
    out = tmp_path / "lonely_auto_only"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    summary = summarize_pilot(out)
    text = format_summary(summary)
    assert "mapping_review.xlsx" in text and "MISSING" in text
    assert "confidence_report.md" in text
    assert "converted_template.docx" in text
    # And every other stage is marked not-yet-run.
    assert "[confirm-mapping] [not yet run]" in text


# ---------------------------------------------------------------------------
# Redaction contract — the load-bearing claim of the new command
# ---------------------------------------------------------------------------

def test_full_pilot_summary_never_prints_sentinels(tmp_path: Path):
    out = tmp_path / "full_pilot_redacted"
    _build_full_pilot(out)
    summary = summarize_pilot(out)
    text = format_summary(summary)
    _assert_no_sentinels(text)
    # Positive shape assertions so a vacuously-redacted (empty) output
    # cannot trivially pass the test.
    assert "pilot-summary for:" in text
    assert "[learn] [present]" in text
    assert "[confirm-mapping] [present]" in text
    assert "[run-preview] [present]" in text
    assert "[render-docx] [present]" in text


def test_full_pilot_summary_never_prints_full_out_path(tmp_path: Path):
    out = tmp_path / "full_pilot_with_path"
    _build_full_pilot(out)
    summary = summarize_pilot(out)
    text = format_summary(summary)
    assert str(out) not in text
    assert str(out.parent) not in text


def test_basename_used_in_header_not_parent_tree(tmp_path: Path):
    """The header reduces ``--out`` to its basename so the operator's
    parent tree (home dir, project codename, customer label) never
    leaks into a pasted summary."""
    sensitive_parent = "customer_alpha_2026Q2_secret"
    out = tmp_path / sensitive_parent / "output"
    _build_full_pilot(out)
    summary = summarize_pilot(out)
    text = format_summary(summary)
    assert "output" in text
    assert sensitive_parent not in text
    assert str(tmp_path) not in text


def test_pilot_summary_with_run_preview_failures_keeps_redaction(tmp_path: Path):
    out = tmp_path / "preview_failed_redacted"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    _write_confirmed_mapping(out / "confirmed_mapping.yml")
    _write_run_validation(
        out / "run_preview" / "run_validation.xlsx",
        ok_count=4,
        failed={"missing_sheet": 2, "non_numeric_cell": 1},
    )
    text = format_summary(summarize_pilot(out))
    _assert_no_sentinels(text)
    # Status enum strings are policy-level, NOT content — they must surface.
    assert "missing_sheet" in text
    assert "non_numeric_cell" in text
    assert "ok" in text


def test_pilot_summary_with_render_failure_keeps_redaction(tmp_path: Path):
    out = tmp_path / "render_failed_redacted"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    _write_confirmed_mapping(out / "confirmed_mapping.yml")
    _write_run_validation(out / "run_preview" / "run_validation.xlsx", ok_count=4)
    _write_render_log(
        out / "render_log.yml", total=4, ok=2, failed=2, total_replacements=2,
    )
    _write_rendered_docx(out / "new_report.docx")
    text = format_summary(summarize_pilot(out))
    _assert_no_sentinels(text)
    # The log's failure count must reach the formatter (it's a count, not content).
    assert "failed" in text and "2" in text


# ---------------------------------------------------------------------------
# Counts and per-stage hints are surfaced
# ---------------------------------------------------------------------------

def test_summary_surfaces_learn_counts_and_strict_hint(tmp_path: Path):
    out = tmp_path / "counts_with_strict_blockers"
    out.mkdir()
    _write_auto_mapping(
        out / "auto_mapping.yml",
        total=50,
        by_conf={"HIGH": 30, "MEDIUM": 5, "LOW": 4, "UNRESOLVED": 3, "EXCLUDED": 8},
    )
    text = format_summary(summarize_pilot(out))
    assert "total Word numbers" in text and "50" in text
    assert "HIGH" in text and "30" in text
    assert "UNRESOLVED" in text and "3" in text
    assert "EXCLUDED" in text and "8" in text
    assert "learn --strict" in text


def test_summary_confirm_complete_hint_points_to_run_preview(tmp_path: Path):
    out = tmp_path / "confirmed_clean"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    _write_confirmed_mapping(out / "confirmed_mapping.yml", complete=True)
    text = format_summary(summarize_pilot(out))
    assert "summary.complete" in text and "true" in text
    assert "summary.allow_incomplete" in text and "false" in text
    assert "run-preview" in text


def test_summary_flags_allow_incomplete_as_unsafe_for_pilot(tmp_path: Path):
    out = tmp_path / "exploratory"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    _write_confirmed_mapping(
        out / "confirmed_mapping.yml",
        complete=False,
        allow_incomplete=True,
        confirmed=0,
        review_required=0,
        audit_only=1,
    )
    text = format_summary(summarize_pilot(out))
    assert "allow_incomplete" in text and "true" in text
    assert "must NOT" in text or "exploratory" in text.lower()


def test_summary_surfaces_run_validation_status_counts(tmp_path: Path):
    out = tmp_path / "run_preview_mixed"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    _write_confirmed_mapping(out / "confirmed_mapping.yml")
    _write_run_validation(
        out / "run_preview" / "run_validation.xlsx",
        ok_count=4,
        failed={"missing_cell": 2, "transform_unknown": 1},
    )
    text = format_summary(summarize_pilot(out))
    assert "ok" in text
    assert "missing_cell" in text
    assert "transform_unknown" in text


def test_summary_locates_run_validation_at_top_level_too(tmp_path: Path):
    """An operator who points ``run-preview --out`` directly at the
    pilot output (instead of $PILOT/output/run_preview/) drops the
    file at the top of --out. pilot-summary must still find it."""
    out = tmp_path / "flat_layout"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    _write_confirmed_mapping(out / "confirmed_mapping.yml")
    _write_run_validation(out / "run_validation.xlsx", ok_count=3)
    summary = summarize_pilot(out)
    assert summary.run_preview.present is True


def test_summary_surfaces_render_counts_when_log_present(tmp_path: Path):
    out = tmp_path / "rendered_clean"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    _write_confirmed_mapping(out / "confirmed_mapping.yml")
    _write_run_validation(out / "run_preview" / "run_validation.xlsx", ok_count=4)
    _write_render_log(
        out / "render_log.yml", total=4, ok=4, total_replacements=7,
    )
    _write_rendered_docx(out / "new_report.docx")
    text = format_summary(summarize_pilot(out))
    assert "log rows" in text
    assert "7" in text  # total_replacements
    assert "rendered docx files" in text
    assert "validate-render" in text


# ---------------------------------------------------------------------------
# Read-only contract — summarizer must never mutate any artifact
# ---------------------------------------------------------------------------

def test_summarize_pilot_never_mutates_any_artifact(tmp_path: Path):
    out = tmp_path / "immutable_pilot"
    _build_full_pilot(out)
    files = [p for p in out.rglob("*") if p.is_file()]
    sizes_before = {p: p.stat().st_size for p in files}
    mtimes_before = {p: p.stat().st_mtime_ns for p in files}

    summarize_pilot(out)
    # And again via the CLI to cover the wiring path.
    cli_main(["pilot-summary", "--out", str(out)])

    for p, size in sizes_before.items():
        assert p.stat().st_size == size, f"{p.name} size changed"
    for p, mt in mtimes_before.items():
        assert p.stat().st_mtime_ns == mt, f"{p.name} mtime changed"


def test_companions_detected_by_stat_only_not_parsed(tmp_path: Path):
    """Pin the load-surface claim: ``mapping_review.xlsx``,
    ``converted_template.docx``, ``confidence_report.md``, and any
    rendered ``.docx`` are detected by ``Path.exists`` / ``Path.stat``
    / directory iteration only. The summarizer never opens them, so
    a corrupt file in any of those slots must NOT break the summary
    or surface a ``cannot parse`` / ``cannot open`` issue.

    A regression here (e.g. someone adding ``docx.Document(p)`` to the
    learn-stage walker to count paragraphs) would let a torn-write
    docx crash the summary or — worse — load full Word body text into
    memory, broadening the load surface beyond what the docstring
    claims.
    """
    out = tmp_path / "corrupt_companions"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    (out / "mapping_review.xlsx").write_bytes(b"definitely not a real xlsx")
    (out / "converted_template.docx").write_bytes(b"definitely not a real docx")
    (out / "confidence_report.md").write_bytes(b"\xff\xfe random bytes")
    (out / "new_report.docx").write_bytes(b"also not a real docx")

    summary = summarize_pilot(out)
    text = format_summary(summary)

    # Every corrupt companion is reported as PRESENT with a byte size.
    for name in (
        "mapping_review.xlsx",
        "converted_template.docx",
        "confidence_report.md",
    ):
        assert name in text and "bytes" in text
    assert "rendered docx files" in text and "1" in text

    # And — the load-surface guarantee — no parse / open issue surfaced,
    # because the summarizer never opens any of these files.
    assert "cannot parse" not in text, text
    assert "cannot open" not in text, text
    assert summary.ok is True


def test_summarize_handles_corrupt_run_validation_xlsx(tmp_path: Path):
    """``run_validation.xlsx`` IS opened (openpyxl) to count its
    ``Status`` column — so a torn-write or non-xlsx file there must
    surface as an ``cannot read`` issue on the stage, not propagate
    the openpyxl exception out of ``summarize_pilot``."""
    out = tmp_path / "corrupt_run_validation"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    _write_confirmed_mapping(out / "confirmed_mapping.yml")
    (out / "run_preview").mkdir()
    (out / "run_preview" / "run_validation.xlsx").write_bytes(b"not a real xlsx")

    summary = summarize_pilot(out)
    text = format_summary(summary)
    assert summary.run_preview.present is True
    assert "run_validation.xlsx" in text
    assert "cannot read" in text
    # And nothing leaks from the failed read (no sentinels were planted
    # in this file, but we also haven't pulled the openpyxl exception
    # message — which can echo file paths — into the formatter).
    assert str(out) not in text


def test_summarize_pilot_with_corrupt_yaml_does_not_raise(tmp_path: Path):
    """A torn write (e.g. operator killed the process mid-confirm)
    leaves a half-written YAML. The summarizer surfaces the file as
    'cannot parse' and continues; it must never propagate the
    underlying exception out of the function."""
    out = tmp_path / "corrupt"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    (out / "confirmed_mapping.yml").write_text(
        "not: valid: yaml: at: all: }: }: }:",
        encoding="utf-8",
    )
    summary = summarize_pilot(out)
    text = format_summary(summary)
    # Stage is still listed, with the parse failure called out.
    assert "[confirm-mapping]" in text
    assert "cannot parse" in text


def test_format_summary_output_ends_in_newline(tmp_path: Path):
    """Stable trailing newline keeps copy-paste tidy in chat clients
    and avoids 'no newline at end of file' warnings if the output is
    redirected to a text file by an operator."""
    out = tmp_path / "trailing_newline"
    out.mkdir()
    _write_auto_mapping(out / "auto_mapping.yml")
    text = format_summary(summarize_pilot(out))
    assert text.endswith("\n")


# ---------------------------------------------------------------------------
# Privacy advisory wires up
# ---------------------------------------------------------------------------

def test_pilot_summary_emits_privacy_advisory_when_out_under_samples(capsys):
    """If a typo'd --out points inside the repo's samples/, the same
    privacy advisory the rest of the CLI emits must surface here, so
    the operator is warned before pasting the summary anywhere."""
    fixture = SAMPLES_DIR / "synthetic" / "_pilot_summary_advisory_check"
    try:
        _build_full_pilot(fixture)
        rc = cli_main(["pilot-summary", "--out", str(fixture)])
        captured = capsys.readouterr()
        assert rc == 0
        assert "PRIVACY PREFLIGHT ADVISORY" in captured.err
        assert "--out" in captured.err
        # Underlying summary still printed to stdout.
        assert "pilot-summary for:" in captured.out
    finally:
        if fixture.exists():
            shutil.rmtree(fixture)
