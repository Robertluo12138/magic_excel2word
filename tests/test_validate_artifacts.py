"""Tests for the ``validate-artifacts`` consistency gate.

The validator's contract: re-read the four files learn-mode wrote and
prove they tell the same story about every Word number. Two flavours of
coverage here:

* a passing path on a fresh synthetic learn run (sanity-checks that the
  emitter and the validator agree end-to-end);
* a series of *targeted corruptions*. Each one tweaks exactly one row or
  one mention in a single artifact so the failure cause is unambiguous —
  if a future refactor introduces a different inconsistency, the test
  output should point straight at it instead of swallowing it.
"""
from __future__ import annotations

from pathlib import Path
from typing import Tuple

import docx
import openpyxl
import yaml

from src.artifact_validator import validate_artifacts
from src.main import main as cli_main
from src.synthetic_generator import generate


# ---------------------------------------------------------------------------
# Helpers — run learn-mode once into a tmp dir and return the artifact paths.
# ---------------------------------------------------------------------------

def _run_learn(tmp_path: Path) -> Tuple[Path, Path, Path, Path, Path]:
    samples = tmp_path / "samples"
    out = tmp_path / "out"
    generate(samples)
    rc = cli_main([
        "learn",
        "--excel", str(samples / "historical.xlsx"),
        "--word", str(samples / "finished_report.docx"),
        "--out", str(out),
    ])
    assert rc == 0
    return (
        out,
        out / "mapping_review.xlsx",
        out / "auto_mapping.yml",
        out / "converted_template.docx",
        out / "confidence_report.md",
    )


def _load_yaml(path: Path) -> dict:
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def _dump_yaml(data: dict, path: Path) -> None:
    path.write_text(
        yaml.safe_dump(data, allow_unicode=True, sort_keys=False, width=1000),
        encoding="utf-8",
    )


def _codes(report) -> set:
    return {i.code for i in report.issues}


# ---------------------------------------------------------------------------
# Happy path
# ---------------------------------------------------------------------------

def test_validate_passes_on_fresh_synthetic_learn(tmp_path: Path):
    out, *_ = _run_learn(tmp_path)
    report = validate_artifacts(out)
    assert report.ok, (
        "fresh synthetic learn output must be internally consistent; got "
        f"issues: {[(i.code, i.message) for i in report.issues]}"
    )


def test_cli_validate_artifacts_passes_on_fresh_learn(tmp_path: Path, capsys):
    out, *_ = _run_learn(tmp_path)
    rc = cli_main(["validate-artifacts", "--out", str(out)])
    assert rc == 0
    assert "OK" in capsys.readouterr().out


def test_cli_validate_artifacts_returns_2_for_missing_dir(tmp_path: Path, capsys):
    rc = cli_main(["validate-artifacts", "--out", str(tmp_path / "no_such_dir")])
    assert rc == 2
    assert "not found" in capsys.readouterr().err


# ---------------------------------------------------------------------------
# Missing artifacts
# ---------------------------------------------------------------------------

def test_missing_artifact_is_flagged(tmp_path: Path):
    out, _, yaml_path, _, _ = _run_learn(tmp_path)
    yaml_path.unlink()
    report = validate_artifacts(out)
    assert not report.ok
    assert "missing_artifact" in _codes(report)


# ---------------------------------------------------------------------------
# Word ID uniqueness + 1:1 correspondence
# ---------------------------------------------------------------------------

def test_duplicate_word_id_in_xlsx_is_flagged(tmp_path: Path):
    out, xlsx_path, *_ = _run_learn(tmp_path)
    wb = openpyxl.load_workbook(str(xlsx_path))
    ws = wb.active
    # row 1 = header. Force row 3's Word ID to clone row 2's.
    ws.cell(row=3, column=1).value = ws.cell(row=2, column=1).value
    wb.save(str(xlsx_path))
    report = validate_artifacts(out)
    assert not report.ok
    assert "xlsx_duplicate_word_id" in _codes(report)


def test_duplicate_word_id_in_yaml_is_flagged(tmp_path: Path):
    out, _, yaml_path, _, _ = _run_learn(tmp_path)
    data = _load_yaml(yaml_path)
    data["mappings"][1]["word_id"] = data["mappings"][0]["word_id"]
    _dump_yaml(data, yaml_path)
    report = validate_artifacts(out)
    assert not report.ok
    assert "yaml_duplicate_word_id" in _codes(report)


def test_extra_xlsx_row_breaks_one_to_one(tmp_path: Path):
    out, xlsx_path, *_ = _run_learn(tmp_path)
    wb = openpyxl.load_workbook(str(xlsx_path))
    ws = wb.active
    # Append a row with a Word ID that won't exist in YAML.
    n_cols = ws.max_column
    extra = ["word_9999"] + [""] * (n_cols - 1)
    ws.append(extra)
    wb.save(str(xlsx_path))
    report = validate_artifacts(out)
    assert not report.ok
    codes = _codes(report)
    assert "xlsx_orphan_word_ids" in codes
    assert "row_count_mismatch" in codes


def test_dropped_yaml_mapping_breaks_one_to_one(tmp_path: Path):
    out, _, yaml_path, _, _ = _run_learn(tmp_path)
    data = _load_yaml(yaml_path)
    data["mappings"].pop(0)
    _dump_yaml(data, yaml_path)
    report = validate_artifacts(out)
    assert not report.ok
    codes = _codes(report)
    assert "xlsx_orphan_word_ids" in codes
    assert "row_count_mismatch" in codes


# ---------------------------------------------------------------------------
# Field agreement
# ---------------------------------------------------------------------------

def test_location_mismatch_between_xlsx_and_yaml(tmp_path: Path):
    out, xlsx_path, *_ = _run_learn(tmp_path)
    wb = openpyxl.load_workbook(str(xlsx_path))
    ws = wb.active
    # Column 2 is Word Location (the schema is Word ID, Word Location, ...).
    ws.cell(row=2, column=2).value = "paragraph:99999"
    wb.save(str(xlsx_path))
    report = validate_artifacts(out)
    assert "location_mismatch" in _codes(report)


def test_status_mismatch_between_xlsx_and_yaml(tmp_path: Path):
    out, xlsx_path, *_ = _run_learn(tmp_path)
    wb = openpyxl.load_workbook(str(xlsx_path))
    ws = wb.active
    # Column 7 is Confidence.
    header = [c.value for c in ws[1]]
    conf_col = header.index("Confidence") + 1
    ws.cell(row=2, column=conf_col).value = "LOW"
    wb.save(str(xlsx_path))
    report = validate_artifacts(out)
    assert "status_mismatch" in _codes(report)


def test_review_status_mismatch_between_xlsx_and_yaml(tmp_path: Path):
    """A tampered ``Review Status`` cell in the XLSX must be caught as a
    *distinct* issue from confidence drift. Without this check a reviewer
    could be tricked into signing off a row whose XLSX label says
    ``pending_review`` while the YAML still records ``needs_source``."""
    out, xlsx_path, *_ = _run_learn(tmp_path)
    wb = openpyxl.load_workbook(str(xlsx_path))
    ws = wb.active
    header = [c.value for c in ws[1]]
    col = header.index("Review Status") + 1
    original = ws.cell(row=2, column=col).value
    # Force a value the YAML's review_status will not match. Picking the
    # explicit opposite of whatever the canonical value is avoids a
    # tautological "blank vs blank" collision.
    ws.cell(row=2, column=col).value = (
        "needs_source" if original != "needs_source" else "pending_review"
    )
    wb.save(str(xlsx_path))
    report = validate_artifacts(out)
    codes = _codes(report)
    assert "review_status_mismatch" in codes, (
        "tampered Review Status must surface as its own issue code; "
        f"got: {[(i.code, i.message) for i in report.issues]}"
    )
    # Distinct from confidence drift — the test guards against a future
    # refactor collapsing the two checks into one.
    assert "status_mismatch" not in codes


def test_placeholder_status_mismatch_between_xlsx_and_yaml(tmp_path: Path):
    """A tampered ``Placeholder Status`` cell must fail validation with its
    own distinct code, independent of the placeholder-token leak checks.
    This is the audit lever a corrupted XLSX could pull to claim a row was
    ``applied`` while the YAML and docx still show it skipped."""
    out, xlsx_path, *_ = _run_learn(tmp_path)
    wb = openpyxl.load_workbook(str(xlsx_path))
    ws = wb.active
    header = [c.value for c in ws[1]]
    col = header.index("Placeholder Status") + 1
    original = ws.cell(row=2, column=col).value
    ws.cell(row=2, column=col).value = (
        "skipped_unresolved" if original != "skipped_unresolved" else "applied"
    )
    wb.save(str(xlsx_path))
    report = validate_artifacts(out)
    codes = _codes(report)
    assert "placeholder_status_mismatch" in codes, (
        "tampered Placeholder Status must surface as its own issue code; "
        f"got: {[(i.code, i.message) for i in report.issues]}"
    )
    assert "review_status_mismatch" not in codes


# ---------------------------------------------------------------------------
# Placeholder rules
# ---------------------------------------------------------------------------

def test_placeholder_claimed_for_unresolved_is_flagged(tmp_path: Path):
    out, _, yaml_path, _, _ = _run_learn(tmp_path)
    data = _load_yaml(yaml_path)
    bumped = False
    for entry in data["mappings"]:
        if entry["status"] == "UNRESOLVED":
            entry["placeholder_status"] = "applied"
            entry["placeholder"] = "{{ " + entry["word_id"] + " }}"
            bumped = True
            break
    assert bumped, "synthetic corpus should have at least one UNRESOLVED row"
    _dump_yaml(data, yaml_path)
    report = validate_artifacts(out)
    codes = _codes(report)
    # The status is unsafe AND the docx is missing the token: both must surface.
    assert "unsafe_status_with_placeholder" in codes
    assert "docx_placeholder_missing" in codes


def test_placeholder_leak_into_docx_is_flagged(tmp_path: Path):
    """A skipped YAML row whose placeholder snuck into the converted docx must
    be reported — that is the silent-omission risk CLAUDE.md calls out."""
    out, _, yaml_path, docx_path, _ = _run_learn(tmp_path)
    data = _load_yaml(yaml_path)
    target_id = None
    for entry in data["mappings"]:
        if entry["placeholder_status"] != "applied":
            target_id = entry["word_id"]
            break
    assert target_id is not None
    # Manually append a stray placeholder to the converted template.
    doc = docx.Document(str(docx_path))
    doc.add_paragraph("{{ " + target_id + " }}")
    doc.save(str(docx_path))
    report = validate_artifacts(out)
    assert "docx_placeholder_leak" in _codes(report)


def test_placeholder_leak_through_compact_jinja_syntax_is_flagged(tmp_path: Path):
    """Jinja2 accepts both ``{{ word_0001 }}`` and the compact ``{{word_0001}}``.
    If a skipped row leaks through the compact form, naive substring matching
    against the writer's canonical spaced output would silently green-light
    it. The validator must catch it anyway — that was a regression caught
    by stop-time review."""
    out, _, yaml_path, docx_path, _ = _run_learn(tmp_path)
    data = _load_yaml(yaml_path)
    target_id = None
    for entry in data["mappings"]:
        if entry["placeholder_status"] != "applied":
            target_id = entry["word_id"]
            break
    assert target_id is not None
    # Compact, whitespace-stripped form — also valid Jinja syntax.
    doc = docx.Document(str(docx_path))
    doc.add_paragraph("{{" + target_id + "}}")
    doc.save(str(docx_path))
    report = validate_artifacts(out)
    assert "docx_placeholder_leak" in _codes(report), (
        "compact-syntax placeholder must still be detected as a leak; "
        f"got: {[(i.code, i.message) for i in report.issues]}"
    )


def test_orphan_placeholder_in_docx_with_no_yaml_row(tmp_path: Path):
    out, _, _, docx_path, _ = _run_learn(tmp_path)
    doc = docx.Document(str(docx_path))
    doc.add_paragraph("{{ word_9999 }}")
    doc.save(str(docx_path))
    report = validate_artifacts(out)
    assert "docx_orphan_placeholder" in _codes(report)


# ---------------------------------------------------------------------------
# confidence_report.md must reflect the YAML summary's audit buckets.
# ---------------------------------------------------------------------------

def test_report_missing_unresolved_when_count_is_nonzero(tmp_path: Path):
    out, _, yaml_path, _, md_path = _run_learn(tmp_path)
    # Synthetic corpus has UNRESOLVED >= 2; deleting the word from the
    # report must therefore fail validation.
    counts = _load_yaml(yaml_path).get("summary", {}).get("by_confidence", {})
    assert counts.get("UNRESOLVED", 0) > 0
    text = md_path.read_text(encoding="utf-8")
    md_path.write_text(text.replace("UNRESOLVED", "FIXME"), encoding="utf-8")
    report = validate_artifacts(out)
    codes = _codes(report)
    assert "report_missing_section" in codes
    assert any("UNRESOLVED" in i.message for i in report.issues)


def test_report_missing_excluded_when_count_is_nonzero(tmp_path: Path):
    out, _, yaml_path, _, md_path = _run_learn(tmp_path)
    counts = _load_yaml(yaml_path).get("summary", {}).get("by_confidence", {})
    assert counts.get("EXCLUDED", 0) > 0
    text = md_path.read_text(encoding="utf-8")
    md_path.write_text(text.replace("EXCLUDED", "skipped"), encoding="utf-8")
    report = validate_artifacts(out)
    assert "report_missing_section" in _codes(report)
    assert any("EXCLUDED" in i.message for i in report.issues)
