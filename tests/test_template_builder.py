"""Tests for the template_builder learn-mode artifacts.

The contract we protect here:
  * ``auto_mapping.yml`` lists every visible Word number — HIGH, MEDIUM,
    LOW, UNRESOLVED, AND EXCLUDED — with stable, unique ``word_id``s.
  * ``converted_template.docx`` replaces only HIGH/MEDIUM tokens with
    ``{{ word_NNNN }}`` placeholders. LOW, UNRESOLVED, and EXCLUDED
    values remain visible so a reviewer can audit them.
  * Anything that *could not* be replaced safely (offset drift, raw
    mismatch) is recorded in the YAML, not silently dropped.
"""
from __future__ import annotations

from pathlib import Path
from typing import List

import docx
import yaml

from src.excel_profiler import ExcelCell
from src.synthetic_generator import generate
from src.template_builder import (
    SAFE_CONFIDENCES,
    assign_word_ids,
    write_template_artifacts,
)
from src.value_matcher import Candidate, WordMatch, match_word_numbers
from src.excel_profiler import profile_workbook
from src.word_profiler import WordNumber, profile_document


# ---------------------------------------------------------------------------
# Hand-built matches: minimal coverage for status semantics & uniqueness.
# ---------------------------------------------------------------------------

def _wn(location, snippet, raw, value, unit=None, offset=0, label_context=(), excl=None):
    return WordNumber(
        location=location,
        snippet=snippet,
        label_context=list(label_context),
        raw=raw,
        value=value,
        unit=unit,
        sign=1,
        offset=offset,
        exclusion_reason=excl,
    )


def _cell(value, row_ctx=("营业收入",), col_ctx=("2026年5月",), sheet="月度", addr="A1"):
    return ExcelCell(
        sheet=sheet, address=addr, row=1, column=1,
        raw_value=value, numeric_value=float(value),
        row_context=list(row_ctx), column_context=list(col_ctx),
    )


def _candidate(cell, value_score=1.0, context_score=1.0, overlap=("营业收入",), interp="as_written"):
    return Candidate(
        cell=cell,
        value_score=value_score,
        context_score=context_score,
        overlap_tokens=list(overlap),
        interpretation=interp,
    )


def _match(wn, confidence, candidates=(), chosen=None, note=""):
    return WordMatch(
        word_number=wn,
        candidates=list(candidates),
        confidence=confidence,
        chosen=chosen,
        note=note,
    )


def _build_mixed_matches() -> List[WordMatch]:
    cell_high = _cell(100.0, row_ctx=("营业收入",), col_ctx=("2026年5月",), addr="B2")
    cell_med = _cell(200.0, row_ctx=("净利润",), col_ctx=("2026年5月",), addr="B3")
    cell_low = _cell(300.0, row_ctx=("其他",), col_ctx=("列",), addr="B4")
    cand_high = _candidate(cell_high)
    cand_med = _candidate(cell_med, context_score=0.5, interp="as_written")
    cand_low = _candidate(cell_low, value_score=0.7, context_score=0.0, overlap=())

    wn_high = _wn("paragraph:0", "营收100元", "100", 100.0, unit="元", offset=2)
    wn_med = _wn("paragraph:1", "利润200元", "200", 200.0, unit="元", offset=2)
    wn_low = _wn("paragraph:2", "其他300元", "300", 300.0, unit="元", offset=2)
    wn_unres = _wn("paragraph:3", "提升15%", "15%", 15.0, unit="%", offset=2)
    wn_excl = _wn(
        "paragraph:4", "2026年5月", "2026", 2026.0, offset=0,
        excl="date/period marker (followed by '年')",
    )

    return [
        _match(wn_high, "HIGH", [cand_high], cand_high),
        _match(wn_med, "MEDIUM", [cand_med], cand_med, note=""),
        _match(wn_low, "LOW", [cand_low], cand_low, note="value match without strong context overlap"),
        _match(wn_unres, "UNRESOLVED"),
        _match(wn_excl, "EXCLUDED", note="date/period marker (followed by '年')"),
    ]


def _write_minimal_docx(out: Path, paragraphs):
    doc = docx.Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(str(out))


# ---------------------------------------------------------------------------
# assign_word_ids
# ---------------------------------------------------------------------------

def test_assign_word_ids_is_stable_and_zero_padded():
    matches = _build_mixed_matches()
    ids = assign_word_ids(matches)
    assert ids == ["word_0001", "word_0002", "word_0003", "word_0004", "word_0005"]
    # Determinism: same input -> same ids.
    assert assign_word_ids(matches) == ids


def test_safe_confidences_are_exactly_high_and_medium():
    # Pinning this set guards against accidentally promoting LOW into
    # the auto-templating pipeline, which would silently leak unsafe
    # replacements into the converted template.
    assert SAFE_CONFIDENCES == frozenset({"HIGH", "MEDIUM"})


# ---------------------------------------------------------------------------
# YAML contents
# ---------------------------------------------------------------------------

def test_yaml_includes_every_match_with_unique_ids(tmp_path: Path):
    matches = _build_mixed_matches()
    docx_path = tmp_path / "in.docx"
    _write_minimal_docx(docx_path, ["营收100元", "利润200元", "其他300元", "提升15%", "2026年5月"])

    artifacts = write_template_artifacts(matches, docx_path, tmp_path / "out")
    data = yaml.safe_load(artifacts.yaml_path.read_text(encoding="utf-8"))

    assert data["schema_version"] == 1
    assert data["summary"]["total"] == len(matches)
    assert data["summary"]["by_confidence"] == {
        "HIGH": 1, "MEDIUM": 1, "LOW": 1, "UNRESOLVED": 1, "EXCLUDED": 1,
    }

    entries = data["mappings"]
    assert len(entries) == len(matches), "every match must appear in YAML"

    ids = [e["word_id"] for e in entries]
    assert len(set(ids)) == len(ids), "word_ids must be unique"
    assert ids == ["word_0001", "word_0002", "word_0003", "word_0004", "word_0005"]


def test_yaml_carries_required_fields_per_goal(tmp_path: Path):
    """word_id/location/raw/unit/context/status/confidence/recommended source/
    transform metadata/review_status must all be present."""
    matches = _build_mixed_matches()
    docx_path = tmp_path / "in.docx"
    _write_minimal_docx(docx_path, ["营收100元", "利润200元", "其他300元", "提升15%", "2026年5月"])

    artifacts = write_template_artifacts(matches, docx_path, tmp_path / "out")
    entries = yaml.safe_load(artifacts.yaml_path.read_text(encoding="utf-8"))["mappings"]

    required = {
        "word_id", "location", "raw", "unit", "context",
        "status", "confidence", "review_status",
        "recommended_source", "transform", "alternatives",
        "placeholder", "placeholder_status",
    }
    for entry in entries:
        missing = required - entry.keys()
        assert not missing, f"entry {entry.get('word_id')} missing fields {missing}"

    # Recommended source is populated for the records the matcher actually
    # picked a candidate for, and null for UNRESOLVED/EXCLUDED.
    by_id = {e["word_id"]: e for e in entries}
    assert by_id["word_0001"]["recommended_source"]["address"] == "B2"
    assert by_id["word_0001"]["transform"]["interpretation"] == "as_written"
    assert by_id["word_0004"]["recommended_source"] is None  # UNRESOLVED
    assert by_id["word_0005"]["recommended_source"] is None  # EXCLUDED


def test_yaml_preserves_unresolved_and_excluded_rows(tmp_path: Path):
    matches = _build_mixed_matches()
    docx_path = tmp_path / "in.docx"
    _write_minimal_docx(docx_path, ["营收100元", "利润200元", "其他300元", "提升15%", "2026年5月"])

    artifacts = write_template_artifacts(matches, docx_path, tmp_path / "out")
    entries = yaml.safe_load(artifacts.yaml_path.read_text(encoding="utf-8"))["mappings"]
    by_id = {e["word_id"]: e for e in entries}

    unres = by_id["word_0004"]
    assert unres["status"] == "UNRESOLVED"
    assert unres["review_status"] == "needs_source"
    assert unres["placeholder"] is None
    assert unres["placeholder_status"] == "skipped_unresolved"
    assert unres["raw"] == "15%"

    excl = by_id["word_0005"]
    assert excl["status"] == "EXCLUDED"
    assert excl["review_status"] == "audited_excluded"
    assert excl["placeholder"] is None
    assert excl["placeholder_status"] == "skipped_excluded"
    assert excl["raw"] == "2026"


def test_yaml_review_status_for_low_is_needs_review(tmp_path: Path):
    matches = _build_mixed_matches()
    docx_path = tmp_path / "in.docx"
    _write_minimal_docx(docx_path, ["营收100元", "利润200元", "其他300元", "提升15%", "2026年5月"])

    artifacts = write_template_artifacts(matches, docx_path, tmp_path / "out")
    entries = yaml.safe_load(artifacts.yaml_path.read_text(encoding="utf-8"))["mappings"]
    by_id = {e["word_id"]: e for e in entries}

    low = by_id["word_0003"]
    assert low["status"] == "LOW"
    assert low["review_status"] == "needs_review"
    assert low["placeholder"] is None
    assert low["placeholder_status"] == "skipped_low_confidence"


# ---------------------------------------------------------------------------
# Converted template .docx
# ---------------------------------------------------------------------------

def _paragraph_texts(path: Path) -> List[str]:
    doc = docx.Document(str(path))
    return [p.text for p in doc.paragraphs]


def test_high_and_medium_tokens_get_placeholders(tmp_path: Path):
    matches = _build_mixed_matches()
    docx_path = tmp_path / "in.docx"
    _write_minimal_docx(docx_path, ["营收100元", "利润200元", "其他300元", "提升15%", "2026年5月"])

    artifacts = write_template_artifacts(matches, docx_path, tmp_path / "out")
    texts = _paragraph_texts(artifacts.docx_path)

    assert "{{ word_0001 }}" in texts[0]
    assert "100" not in texts[0]
    assert "{{ word_0002 }}" in texts[1]
    assert "200" not in texts[1]
    assert artifacts.placeholders_applied == 2


def test_low_unresolved_excluded_values_remain_visible(tmp_path: Path):
    matches = _build_mixed_matches()
    docx_path = tmp_path / "in.docx"
    _write_minimal_docx(docx_path, ["营收100元", "利润200元", "其他300元", "提升15%", "2026年5月"])

    artifacts = write_template_artifacts(matches, docx_path, tmp_path / "out")
    texts = _paragraph_texts(artifacts.docx_path)

    assert "300" in texts[2], "LOW values must stay visible for audit"
    assert "{{ word_0003 }}" not in texts[2]

    assert "15%" in texts[3], "UNRESOLVED values must stay visible for audit"
    assert "{{ word_0004 }}" not in texts[3]

    assert "2026" in texts[4], "EXCLUDED values must stay visible for audit"
    assert "{{ word_0005 }}" not in texts[4]


def test_placeholders_in_docx_match_yaml_applied_status(tmp_path: Path):
    matches = _build_mixed_matches()
    docx_path = tmp_path / "in.docx"
    _write_minimal_docx(docx_path, ["营收100元", "利润200元", "其他300元", "提升15%", "2026年5月"])

    artifacts = write_template_artifacts(matches, docx_path, tmp_path / "out")
    entries = yaml.safe_load(artifacts.yaml_path.read_text(encoding="utf-8"))["mappings"]
    texts = _paragraph_texts(artifacts.docx_path)
    body = "\n".join(texts)

    applied_ids = {e["word_id"] for e in entries if e["placeholder_status"] == "applied"}
    for word_id in applied_ids:
        token = "{{ " + word_id + " }}"
        assert token in body, f"YAML says {word_id} applied but token missing in docx"

    skipped_ids = {e["word_id"] for e in entries if e["placeholder_status"] != "applied"}
    for word_id in skipped_ids:
        token = "{{ " + word_id + " }}"
        assert token not in body, f"YAML says {word_id} skipped but token leaked into docx"


def test_yaml_surfaces_runner_up_candidates_for_audit(tmp_path: Path):
    """Ambiguous picks must keep every runner-up in YAML — that is the
    audit trail CLAUDE.md mandates for "make ambiguous matches visible"."""
    cell_top = _cell(100.0, row_ctx=("营业收入",), col_ctx=("2026年5月",), sheet="月度", addr="B2")
    cell_alt1 = _cell(100.0, row_ctx=("营业收入",), col_ctx=("2026年5月",), sheet="月度", addr="B3")
    cell_alt2 = _cell(100.0, row_ctx=("其他",), col_ctx=("列",), sheet="渠道", addr="B4")
    cand_top = _candidate(cell_top)
    cand_alt1 = _candidate(cell_alt1, context_score=0.95)
    cand_alt2 = _candidate(cell_alt2, context_score=0.5, overlap=("收入",), interp="as_written")

    wn = _wn("paragraph:0", "营收100元", "100", 100.0, unit="元", offset=2)
    match = _match(
        wn, "MEDIUM",
        candidates=[cand_top, cand_alt1, cand_alt2],
        chosen=cand_top,
        note="ambiguous: multiple cells tie on value+context",
    )

    docx_path = tmp_path / "in.docx"
    _write_minimal_docx(docx_path, ["营收100元"])
    artifacts = write_template_artifacts([match], docx_path, tmp_path / "out")
    entries = yaml.safe_load(artifacts.yaml_path.read_text(encoding="utf-8"))["mappings"]

    [entry] = entries
    assert entry["recommended_source"]["address"] == "B2"
    assert entry["alternatives"], "alternatives must be present for ambiguous picks"
    alt_addresses = [a["address"] for a in entry["alternatives"]]
    assert alt_addresses == ["B3", "B4"]
    # Each alternative must carry enough metadata for a reviewer to pick
    # it as the confirmed source if the recommended one is wrong.
    first_alt = entry["alternatives"][0]
    for field in (
        "sheet", "address", "value", "row_context", "column_context",
        "interpretation", "value_score", "context_score", "overlap_tokens",
    ):
        assert field in first_alt, f"alternative missing {field}"
    # The recommended source must NOT also appear under alternatives —
    # otherwise reviewers would double-count it.
    assert "B2" not in alt_addresses


def test_yaml_alternatives_empty_for_unresolved_and_excluded(tmp_path: Path):
    matches = _build_mixed_matches()
    docx_path = tmp_path / "in.docx"
    _write_minimal_docx(docx_path, ["营收100元", "利润200元", "其他300元", "提升15%", "2026年5月"])

    artifacts = write_template_artifacts(matches, docx_path, tmp_path / "out")
    entries = yaml.safe_load(artifacts.yaml_path.read_text(encoding="utf-8"))["mappings"]
    by_id = {e["word_id"]: e for e in entries}
    assert by_id["word_0004"]["alternatives"] == []  # UNRESOLVED
    assert by_id["word_0005"]["alternatives"] == []  # EXCLUDED


def test_offset_or_raw_mismatch_is_recorded_not_dropped(tmp_path: Path):
    # Build a HIGH match whose offset points at the wrong substring on the
    # actual docx (we deliberately ship a paragraph that no longer matches
    # what the profiler saw). The status must be recorded, not silently
    # ignored, and no placeholder may be written to the .docx.
    cell = _cell(100.0)
    cand = _candidate(cell)
    bogus = WordMatch(
        word_number=_wn("paragraph:0", "actual paragraph 100", "999", 999.0, unit=None, offset=10),
        candidates=[cand],
        confidence="HIGH",
        chosen=cand,
    )
    docx_path = tmp_path / "in.docx"
    _write_minimal_docx(docx_path, ["actual paragraph 100"])

    artifacts = write_template_artifacts([bogus], docx_path, tmp_path / "out")
    entries = yaml.safe_load(artifacts.yaml_path.read_text(encoding="utf-8"))["mappings"]
    assert entries[0]["placeholder_status"] == "skipped_raw_mismatch"
    assert entries[0]["placeholder"] is None
    assert entries[0]["review_status"] == "needs_review"
    body = "\n".join(_paragraph_texts(artifacts.docx_path))
    assert "{{ word_0001 }}" not in body
    assert artifacts.placeholders_applied == 0


# ---------------------------------------------------------------------------
# Synthetic learn smoke — end-to-end via real profilers + matcher
# ---------------------------------------------------------------------------

def test_synthetic_learn_creates_template_artifacts(tmp_path: Path):
    samples = tmp_path / "samples"
    out = tmp_path / "out"
    xlsx, docx_path = generate(samples)
    cells = profile_workbook(xlsx)
    word_numbers = profile_document(docx_path)
    matches = match_word_numbers(word_numbers, cells)

    artifacts = write_template_artifacts(matches, docx_path, out)
    assert artifacts.yaml_path.exists()
    assert artifacts.docx_path.exists()

    data = yaml.safe_load(artifacts.yaml_path.read_text(encoding="utf-8"))
    # Each visible Word numeric token must have exactly one YAML entry.
    assert data["summary"]["total"] == len(matches) == len(data["mappings"])
    # Synthetic corpus has at least the deliberate UNRESOLVED ("15%",
    # "0.70个百分点") and date markers — both must surface in YAML.
    by_status = {}
    for entry in data["mappings"]:
        by_status.setdefault(entry["status"], 0)
        by_status[entry["status"]] += 1
    assert by_status.get("UNRESOLVED", 0) >= 2
    assert by_status.get("EXCLUDED", 0) >= 10

    # Placeholders applied should equal the count of "applied" rows AND
    # equal the count of HIGH/MEDIUM whose offset/raw checks passed
    # (everything in this corpus does, but we don't depend on that).
    applied_in_yaml = sum(
        1 for e in data["mappings"] if e["placeholder_status"] == "applied"
    )
    assert artifacts.placeholders_applied == applied_in_yaml
    assert applied_in_yaml >= 5, (
        "synthetic corpus has several HIGH/MEDIUM mappings; "
        f"got only {applied_in_yaml}"
    )

    # The UNRESOLVED "15%" must still be readable in the template — this
    # is the audit guarantee.
    body = "\n".join(p.text for p in docx.Document(str(artifacts.docx_path)).paragraphs)
    assert "15%" in body
    assert "2026" in body  # the year marker is EXCLUDED, stays visible
