"""Full end-to-end acceptance smoke for the deterministic Excel→Word pipeline.

This smoke proves the six-stage learn-mode + run-mode pipeline holds together
on a small *clean* synthetic case — one with no eligible UNRESOLVED/LOW Word
numbers — so the pipeline can be driven to completion without the
``--allow-incomplete`` escape hatch and the final ``.docx`` carries no
surviving placeholders:

    generate-synthetic (clean) →
    learn → validate-artifacts →
    confirm-mapping (no --allow-incomplete) →
    run-preview → render-docx → validate-render

Two happy-path surfaces and one negative path:

* :func:`test_full_acceptance_smoke_python_api` exercises the modules
  directly (``profile_workbook``, ``confirm_mappings``, ``render_docx``,
  ``validate_render``, etc.) so a refactor of the CLI cannot mask a
  regression in the underlying pipeline.
* :func:`test_full_acceptance_smoke_cli` drives the same case through
  ``cli_main([...])`` so the argument parser, exit codes, and stdout/
  stderr paths stay green end-to-end.
* :func:`test_full_pipeline_halts_on_post_render_tamper` reuses the
  successful CLI run and then corrupts ``run_validation.xlsx``.
  ``validate-render`` must refuse to bless the rendered docx (exit 10).
  This is the silent-omission risk CLAUDE.md flags: a downstream consumer
  that only checks for file existence must never silently trust a
  rendered docx whose audit row was flipped to a non-ok status after
  the fact.

Synthetic-only and repo-safe: no real data, no LLM, no GUI, no
cloud/network, no Microsoft Office automation. The clean fixture is
built inline because the default ``generate-synthetic`` corpus
deliberately ships 2 UNRESOLVED rows (``0.70个百分点``, ``15%``) that
would block the no-``--allow-incomplete`` ``confirm-mapping`` gate this
smoke must pass.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Tuple

import docx
import openpyxl
import yaml

from src.artifact_validator import validate_artifacts
from src.excel_profiler import profile_workbook
from src.main import main as cli_main
from src.mapping_confirmer import confirm_mappings, write_confirmed_yaml
from src.mapping_reviewer import write_confidence_report, write_mapping_review
from src.render_validator import validate_render
from src.renderer import render_docx
from src.run_preview import run_preview, write_run_validation
from src.template_builder import assign_word_ids, write_template_artifacts
from src.validator import summarize
from src.value_matcher import match_word_numbers
from src.word_profiler import profile_document

# Any ``{{ word_NNNN }}`` (spaced or compact) in the final docx is a
# silent-omission failure — the rendered report would carry an unresolved
# template token where a confirmed metric should be.
_PLACEHOLDER_RE = re.compile(r"\{\{\s*word_\d+\s*\}\}")


# ---------------------------------------------------------------------------
# Clean synthetic builder — small enough to grok at a glance, large enough
# to exercise the pipeline's audit invariants (multi-unit transforms,
# EXCLUDED date markers, run-mode generated-value formatting).
# ---------------------------------------------------------------------------

def _build_clean_pair(samples_dir: Path) -> Tuple[Path, Path]:
    """Write a paired Excel + Word with zero eligible UNRESOLVED/LOW numbers.

    Shape:
    * Excel ``月度核心指标``: three labeled rows (营业收入, 毛利率, 用户总数).
    * Word: a date heading (yields 2 EXCLUDED markers: ``2026`` followed by
      ``年`` and ``5`` followed by ``月``) plus three numeric metrics that
      each match exactly one Excel cell at HIGH confidence under a distinct
      unit interpretation: ``万元→base_unit``, ``as_written`` (%), and
      ``万人→base_unit``.

    Why this shape: it exercises the three most common business-number
    transforms while keeping the matcher unambiguous, so any pipeline-stage
    regression points at the stage rather than at a flaky corpus.
    """
    samples_dir.mkdir(parents=True, exist_ok=True)
    xlsx = samples_dir / "clean.xlsx"
    docx_path = samples_dir / "clean.docx"

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "月度核心指标"
    ws.append(["指标", "2026年5月"])
    ws.append(["营业收入(元)", 123_456_789.00])
    ws.append(["毛利率(%)", 36.20])
    ws.append(["用户总数(人)", 13_456_789])
    wb.save(str(xlsx))

    doc = docx.Document()
    doc.add_heading("2026年5月经营月度报告", level=1)
    doc.add_paragraph("本月营业收入达12,345.68万元。")
    doc.add_paragraph("毛利率达36.20%。")
    doc.add_paragraph("本月用户总数达到1,345.68万人。")
    doc.save(str(docx_path))
    return xlsx, docx_path


def _fill_reviewer_decisions(review_xlsx: Path) -> None:
    """Mark every HIGH/MEDIUM row in mapping_review.xlsx as ``confirm``.

    The clean corpus has no LOW/UNRESOLVED rows — confirming every
    HIGH/MEDIUM is sufficient to drive the ``confirm-mapping`` gate to
    ``complete: true`` without ``--allow-incomplete``.
    """
    wb = openpyxl.load_workbook(str(review_xlsx))
    ws = wb.active
    header = [c.value for c in ws[1]]
    conf_col = header.index("Confidence") + 1
    decision_col = header.index("Reviewer Decision") + 1
    for r in range(2, ws.max_row + 1):
        if ws.cell(row=r, column=conf_col).value in ("HIGH", "MEDIUM"):
            ws.cell(row=r, column=decision_col).value = "confirm"
    wb.save(str(review_xlsx))


def _docx_text(path: Path) -> str:
    """Return the full visible text of a .docx as a newline-joined string."""
    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Happy path #1 — Python API surface
# ---------------------------------------------------------------------------

def test_full_acceptance_smoke_python_api(tmp_path: Path):
    samples = tmp_path / "samples"
    out = tmp_path / "out"
    out.mkdir(parents=True, exist_ok=True)
    xlsx, word = _build_clean_pair(samples)

    # 1) learn — profile, match, write the four review artifacts.
    excel_cells = profile_workbook(xlsx)
    word_numbers = profile_document(word)
    matches = match_word_numbers(word_numbers, excel_cells)
    summary = summarize(matches)
    # Pre-flight: a clean corpus has zero eligible UNRESOLVED/LOW rows —
    # otherwise this smoke is no longer testing the happy path.
    assert summary.by_confidence["UNRESOLVED"] == 0
    assert summary.by_confidence["LOW"] == 0
    assert summary.by_confidence["HIGH"] + summary.by_confidence["MEDIUM"] == 3

    template_artifacts = write_template_artifacts(matches, word, out)
    word_ids = assign_word_ids(matches)
    review_path = write_mapping_review(
        matches, out / "mapping_review.xlsx",
        word_ids, template_artifacts.placeholder_status,
    )
    write_confidence_report(matches, summary, out / "confidence_report.md")
    for name in (
        "mapping_review.xlsx", "auto_mapping.yml",
        "converted_template.docx", "confidence_report.md",
    ):
        assert (out / name).exists(), f"learn missed {name}"

    # 2) validate-artifacts — the four learn artifacts must agree.
    artifact_report = validate_artifacts(out)
    assert artifact_report.ok, [
        (i.code, i.message) for i in artifact_report.issues
    ]

    # 3) reviewer decisions: HIGH/MEDIUM → "confirm".
    _fill_reviewer_decisions(review_path)

    # 4) confirm-mapping — no --allow-incomplete; must reach complete.
    confirm_report = confirm_mappings(out / "auto_mapping.yml", review_path)
    assert confirm_report.fatal_errors == []
    assert confirm_report.review_required == []
    confirmed_path = out / "confirmed_mapping.yml"
    write_confirmed_yaml(
        confirm_report, out / "auto_mapping.yml", review_path, confirmed_path,
        allow_incomplete=False, total_word_numbers=summary.total,
    )
    confirmed_doc = yaml.safe_load(confirmed_path.read_text(encoding="utf-8"))
    assert confirmed_doc["summary"]["complete"] is True
    assert confirmed_doc["summary"]["confirmed"] == 3
    assert confirmed_doc["summary"]["review_required"] == 0

    # 5) run-preview — every row must resolve cleanly. The smoke reuses the
    # same workbook as a stand-in "new-period" file; production would point
    # at a fresh xlsx, but the contract under test is the pipeline shape.
    preview_report = run_preview(xlsx, confirmed_path)
    assert preview_report.fatal_errors == []
    assert all(r.status == "ok" for r in preview_report.rows), [
        (r.word_id, r.status, r.detail) for r in preview_report.rows
    ]
    run_validation_path = write_run_validation(preview_report, out)
    assert run_validation_path.exists()

    # 6) render-docx — every placeholder substituted, render_log emitted.
    rendered = out / "new_report.docx"
    render_report = render_docx(
        out / "converted_template.docx", run_validation_path, rendered,
    )
    assert render_report.ok, (
        f"render failed; fatal={render_report.fatal_errors} "
        f"failures={[(r.word_id, r.status, r.detail) for r in render_report.failures]}"
    )
    assert rendered.exists()
    assert (out / "render_log.yml").exists()

    # 7) validate-render — the three rendered-output artifacts must agree.
    render_validation = validate_render(
        rendered, out / "render_log.yml", run_validation_path,
    )
    assert render_validation.ok, [
        (i.code, i.message) for i in render_validation.issues
    ]

    # Final invariant: the rendered docx contains no surviving placeholders.
    text = _docx_text(rendered)
    assert not _PLACEHOLDER_RE.search(text), (
        f"placeholders survived render — silent-omission risk; got:\n{text}"
    )
    # Sanity: every formatted display string is in the docx, matching the
    # historical raw token shape applied to the new generated value.
    for expected in ("12,345.68万元", "36.20%", "1,345.68万人"):
        assert expected in text, f"expected display text {expected!r} missing"


# ---------------------------------------------------------------------------
# Happy path #2 — CLI surface (cli_main argv → exit codes)
# ---------------------------------------------------------------------------

def test_full_acceptance_smoke_cli(tmp_path: Path):
    samples = tmp_path / "samples"
    out = tmp_path / "out"
    xlsx, word = _build_clean_pair(samples)

    # 1) learn --strict — the strict gate must pass on a clean corpus.
    assert cli_main([
        "learn", "--strict",
        "--excel", str(xlsx),
        "--word", str(word),
        "--out", str(out),
    ]) == 0
    for name in (
        "mapping_review.xlsx", "auto_mapping.yml",
        "converted_template.docx", "confidence_report.md",
    ):
        assert (out / name).exists(), f"learn missed {name}"

    # 2) validate-artifacts.
    assert cli_main(["validate-artifacts", "--out", str(out)]) == 0

    # 3) reviewer decisions: HIGH/MEDIUM → "confirm".
    _fill_reviewer_decisions(out / "mapping_review.xlsx")

    # 4) confirm-mapping — no --allow-incomplete; must exit 0.
    confirmed = out / "confirmed_mapping.yml"
    assert cli_main([
        "confirm-mapping",
        "--auto", str(out / "auto_mapping.yml"),
        "--review", str(out / "mapping_review.xlsx"),
        "--out", str(confirmed),
    ]) == 0
    confirmed_doc = yaml.safe_load(confirmed.read_text(encoding="utf-8"))
    assert confirmed_doc["summary"]["complete"] is True
    assert confirmed_doc["summary"]["confirmed"] == 3

    # 5) run-preview.
    assert cli_main([
        "run-preview",
        "--excel", str(xlsx),
        "--confirmed", str(confirmed),
        "--out", str(out),
    ]) == 0
    assert (out / "run_validation.xlsx").exists()

    # 6) render-docx.
    rendered = out / "new_report.docx"
    assert cli_main([
        "render-docx",
        "--template", str(out / "converted_template.docx"),
        "--run-validation", str(out / "run_validation.xlsx"),
        "--out", str(rendered),
    ]) == 0
    assert rendered.exists()
    assert (out / "render_log.yml").exists()

    # 7) validate-render.
    assert cli_main([
        "validate-render",
        "--docx", str(rendered),
        "--render-log", str(out / "render_log.yml"),
        "--run-validation", str(out / "run_validation.xlsx"),
    ]) == 0

    # Final invariant: no placeholders survived the render.
    text = _docx_text(rendered)
    assert not _PLACEHOLDER_RE.search(text), (
        f"placeholders survived render — silent-omission risk; got:\n{text}"
    )
    for expected in ("12,345.68万元", "36.20%", "1,345.68万人"):
        assert expected in text, f"expected display text {expected!r} missing"


# ---------------------------------------------------------------------------
# Negative path — post-render tamper of run_validation.xlsx must halt the
# pipeline at validate-render. Without this guard a downstream consumer that
# only checks for ``new_report.docx`` would silently keep using a docx whose
# audit trail no longer matches.
# ---------------------------------------------------------------------------

def test_full_pipeline_halts_on_post_render_tamper(tmp_path: Path):
    samples = tmp_path / "samples"
    out = tmp_path / "out"
    xlsx, word = _build_clean_pair(samples)

    # Reach a successful validate-render first — anything earlier means we
    # haven't shown the *post*-render tamper case.
    assert cli_main([
        "learn", "--strict",
        "--excel", str(xlsx), "--word", str(word), "--out", str(out),
    ]) == 0
    assert cli_main(["validate-artifacts", "--out", str(out)]) == 0
    _fill_reviewer_decisions(out / "mapping_review.xlsx")
    confirmed = out / "confirmed_mapping.yml"
    assert cli_main([
        "confirm-mapping",
        "--auto", str(out / "auto_mapping.yml"),
        "--review", str(out / "mapping_review.xlsx"),
        "--out", str(confirmed),
    ]) == 0
    assert cli_main([
        "run-preview",
        "--excel", str(xlsx),
        "--confirmed", str(confirmed),
        "--out", str(out),
    ]) == 0
    rendered = out / "new_report.docx"
    assert cli_main([
        "render-docx",
        "--template", str(out / "converted_template.docx"),
        "--run-validation", str(out / "run_validation.xlsx"),
        "--out", str(rendered),
    ]) == 0
    # Sanity baseline: validate-render passes before the tamper.
    assert cli_main([
        "validate-render",
        "--docx", str(rendered),
        "--render-log", str(out / "render_log.yml"),
        "--run-validation", str(out / "run_validation.xlsx"),
    ]) == 0

    # The tamper: flip one row's Status from "ok" to a fabricated value.
    # The docx and render_log still claim success — exactly the drift the
    # final gate exists to catch.
    val = out / "run_validation.xlsx"
    wb = openpyxl.load_workbook(str(val))
    ws = wb.active
    headers = [c.value for c in ws[1]]
    status_col = headers.index("Status") + 1
    ws.cell(row=2, column=status_col, value="tampered_after_render")
    wb.save(str(val))

    # validate-render must halt the pipeline with the dedicated exit code.
    assert cli_main([
        "validate-render",
        "--docx", str(rendered),
        "--render-log", str(out / "render_log.yml"),
        "--run-validation", str(val),
    ]) == 10
    # The rendered docx still exists on disk — the gate refuses to bless
    # it, but render-docx never mutated its own outputs. A downstream
    # consumer that respected the gate's exit code is safe; a consumer
    # that just checked for file existence would have used a stale docx.
    assert rendered.exists()
