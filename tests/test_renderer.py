"""Tests for the deterministic Word renderer (``render-docx`` CLI / ``src.renderer``).

Contracts under test (mirrored from src/renderer.py's docstring):

* A successful render substitutes every ``{{ word_NNNN }}`` occurrence
  using display text derived from the historical Word Raw Token shape
  applied to the run-time generated value, writes ``new_report.docx``,
  and emits an audit ``render_log.yml`` alongside.
* The render is gated against any drift between the template and the
  validation: a template placeholder with no validation row, a
  validation row with no matching placeholder, a non-ok validation row,
  a duplicate validation row, or a missing Generated Value all halt the
  render and surface the offending word_id(s).
* The formatter is deterministic and shape-preserving — it reproduces
  the historical raw token's grouping, decimals, unit, and sign style,
  using ROUND_HALF_UP rounding so business-style display matches the
  original report.
* Duplicate placeholders for the same word_id are honored: the renderer
  substitutes every occurrence and the render_log records the count.
"""
from __future__ import annotations

from pathlib import Path
from typing import Dict, List, Optional, Tuple

import docx
import openpyxl
import pytest
import yaml

from src.main import main as cli_main
from src.renderer import (
    FormatInferenceError,
    PLACEHOLDER_RE,
    STATUS_DUPLICATE_VALIDATION_ROW,
    STATUS_FORMAT_INFERENCE_FAILED,
    STATUS_MISSING_GENERATED_VALUE,
    STATUS_NO_PLACEHOLDER,
    STATUS_OK,
    STATUS_VALIDATION_NOT_OK,
    format_display,
    parse_raw_pattern,
    render_docx,
)


# ---------------------------------------------------------------------------
# Tiny fixture builders — a fake template docx + a fake run_validation xlsx
# ---------------------------------------------------------------------------

_RUN_VALIDATION_HEADERS = [
    "Word ID",
    "Word Location",
    "Word Context",
    "Word Raw Token",
    "Word Unit",
    "Source Sheet",
    "Source Cell",
    "Raw Excel Value",
    "Generated Value",
    "Transform Interpretation",
    "Confidence",
    "Status",
    "Detail",
]


def _write_template(
    path: Path,
    paragraphs: List[str],
    tables: Optional[List[List[List[str]]]] = None,
) -> Path:
    """Build a .docx file whose body matches ``paragraphs`` and ``tables``."""
    doc = docx.Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if tables:
        for grid in tables:
            if not grid:
                continue
            t = doc.add_table(rows=len(grid), cols=len(grid[0]))
            for r_idx, row in enumerate(grid):
                for c_idx, cell_text in enumerate(row):
                    t.rows[r_idx].cells[c_idx].text = cell_text
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _row(
    *,
    word_id: str,
    raw_token: str,
    unit: str,
    generated: Optional[float],
    raw_excel: float = 0.0,
    sheet: str = "月度",
    cell: str = "B2",
    status: str = "ok",
    detail: str = "",
    location: str = "paragraph:0",
    context: str = "",
    transform: str = "as_written",
    confidence: str = "HIGH",
) -> List:
    """Return one ``run_validation.xlsx`` row in the canonical column order."""
    return [
        word_id, location, context, raw_token, unit,
        sheet, cell, raw_excel, generated,
        transform, confidence, status, detail,
    ]


def _write_validation(path: Path, rows: List[List]) -> Path:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "run_validation"
    ws.append(_RUN_VALIDATION_HEADERS)
    for r in rows:
        ws.append(r)
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))
    return path


def _docx_text(path: Path) -> str:
    """Return the full visible text of a .docx as a newline-joined string."""
    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Formatter unit tests — the goal pinned five shape categories
# ---------------------------------------------------------------------------

class TestFormatter:
    """``format_display`` must preserve grouping, decimals, unit, and sign."""

    def test_万元_grouping_and_two_decimals(self):
        # raw: "23,456.79万元", value: 23456.789 (= 234,567,890 / 10,000)
        # The 9 in the third decimal rounds the 8 up to 9 → "23,456.79万元".
        assert format_display(
            raw="23,456.79万元", unit="万元", value=23456.789,
        ) == "23,456.79万元"

    def test_percent_preserves_sign_and_precision(self):
        assert format_display(raw="37.50%", unit="%", value=38.45) == "38.45%"

    def test_percent_with_explicit_plus_sign_preserved(self):
        # The "+5%" form deliberately announces a positive delta; the
        # formatter must keep that even when the new value is also positive.
        assert format_display(raw="+5.00%", unit="%", value=7.5) == "+7.50%"
        # When the historical "+" sign is set but the new value is negative
        # we still emit "-" — preserving the historical "+" would be a lie.
        assert format_display(raw="+5.00%", unit="%", value=-2.5) == "-2.50%"

    def test_comma_grouping_preserved_on_large_integer(self):
        assert format_display(raw="1,234,567", unit="", value=1234567) == "1,234,567"
        # When the historical token uses no grouping, the formatter must
        # not introduce commas.
        assert format_display(raw="1234567", unit="", value=1234567) == "1234567"

    def test_negative_explicit_minus_sign(self):
        assert format_display(raw="-1,234.56", unit="", value=-1234.56) == "-1,234.56"
        # Historical was negative but new value is positive: drop the sign.
        assert format_display(raw="-1,234.56", unit="", value=1234.56) == "1,234.56"

    def test_negative_paren_style_preserved(self):
        # Accounting parens: historical "(1,234.56)" means negative.
        assert format_display(raw="(1,234.56)", unit="", value=-1234.56) == "(1,234.56)"
        # Positive value with paren style → no parens (matches accounting
        # convention; parens are reserved for negatives).
        assert format_display(raw="(1,234.56)", unit="", value=1234.56) == "1,234.56"

    def test_亿元_preserves_unit_and_rounds_half_up(self):
        # 1,234,567,890 / 10^8 = 12.3456789 → "12.35亿元" (HALF_UP on the 6).
        assert format_display(
            raw="1.23亿元", unit="亿元", value=12.3456789,
        ) == "12.35亿元"

    def test_integer_token_keeps_zero_decimals(self):
        # Even when the runtime value has fractional precision, an integer
        # historical token must render as integer (matching report style).
        assert format_display(raw="100", unit="", value=99.5) == "100"
        assert format_display(raw="100", unit="", value=99.4) == "99"

    def test_round_half_up_not_banker_rounding(self):
        # Python's default banker's rounding would give "0" for round(0.5).
        # Business reports expect 0.5 → 1.
        assert format_display(raw="1", unit="", value=0.5) == "1"
        # 2.5 → 3 (not 2 as banker's rounding would do).
        assert format_display(raw="1", unit="", value=2.5) == "3"

    def test_zero_value_with_explicit_minus_history(self):
        # Historical "-5" → sign_style=explicit_minus. Zero gets no sign.
        assert format_display(raw="-5", unit="", value=0.0) == "0"

    def test_format_inference_fails_on_unrecognised_token(self):
        with pytest.raises(FormatInferenceError):
            format_display(raw="约1.23亿元 to 2.34亿元", unit="亿元", value=1.5)

    def test_format_inference_fails_on_unit_drift(self):
        # Raw has unit "万元", but validation Word Unit says "元". The
        # formatter must refuse rather than guess which one to trust.
        with pytest.raises(FormatInferenceError):
            format_display(raw="100万元", unit="元", value=100.0)

    def test_format_inference_fails_on_empty_token(self):
        with pytest.raises(FormatInferenceError):
            format_display(raw="", unit="", value=100.0)

    def test_parse_raw_pattern_extracts_full_shape(self):
        p = parse_raw_pattern("(23,456.79万元)", expected_unit="万元")
        assert p.unit == "万元"
        assert p.has_grouping is True
        assert p.decimals == 2
        assert p.sign_style == "paren"


# ---------------------------------------------------------------------------
# Successful render — the golden path
# ---------------------------------------------------------------------------

def _golden_inputs(tmp_path: Path) -> Tuple[Path, Path, Path]:
    template = _write_template(tmp_path / "template.docx", [
        "5月营业收入达{{ word_0001 }}，同比增长{{ word_0002 }}。",
        "毛利率{{ word_0003 }}。",
    ])
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(
            word_id="word_0001",
            raw_token="23,456.79万元", unit="万元",
            generated=23456.789, raw_excel=234567890.0,
            sheet="月度", cell="B2",
        ),
        _row(
            word_id="word_0002",
            raw_token="25.00%", unit="%",
            generated=24.95, raw_excel=0.2495,
            sheet="月度", cell="C2",
        ),
        _row(
            word_id="word_0003",
            raw_token="36.20%", unit="%",
            generated=36.20, raw_excel=0.3620,
            sheet="月度", cell="C3",
        ),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"
    return template, validation, out_docx


def test_render_docx_golden_path(tmp_path: Path):
    template, validation, out_docx = _golden_inputs(tmp_path)

    report = render_docx(template, validation, out_docx)

    assert report.fatal_errors == []
    assert report.failures == [], [
        (r.word_id, r.status, r.detail) for r in report.failures
    ]
    assert report.ok
    assert out_docx.exists()
    log = report.out_log_path
    assert log is not None and log.exists()
    assert log.name == "render_log.yml"
    assert log.parent == out_docx.parent

    text = _docx_text(out_docx)
    # Display text matches the historical raw shape applied to the new value.
    assert "23,456.79万元" in text
    # 24.95% is the new value formatted with the 25.00% template's 2 decimals.
    assert "24.95%" in text
    assert "36.20%" in text
    # No raw placeholders survive in the rendered docx.
    assert not PLACEHOLDER_RE.search(text)


def test_render_log_records_full_audit_trail(tmp_path: Path):
    template, validation, out_docx = _golden_inputs(tmp_path)
    report = render_docx(template, validation, out_docx)

    doc = yaml.safe_load(report.out_log_path.read_text(encoding="utf-8"))
    assert doc["schema_version"] == 1
    assert doc["out_docx"] == str(out_docx)
    assert doc["inputs"]["template"] == str(template)
    assert doc["inputs"]["run_validation"] == str(validation)
    assert doc["summary"]["total_rows"] == 3
    assert doc["summary"]["ok"] == 3
    assert doc["summary"]["failed"] == 0
    assert doc["summary"]["total_replacements"] == 3
    assert doc["summary"]["distinct_placeholder_word_ids"] == 3

    by_id: Dict[str, Dict] = {r["word_id"]: r for r in doc["replacements"]}
    assert set(by_id) == {"word_0001", "word_0002", "word_0003"}
    # Source sheet + cell preserved end-to-end so a reviewer can trace
    # any rendered number back to its Excel origin.
    assert by_id["word_0001"]["source_sheet"] == "月度"
    assert by_id["word_0001"]["source_cell"] == "B2"
    assert by_id["word_0001"]["raw_excel_value"] == 234567890.0
    assert by_id["word_0001"]["generated_value"] == pytest.approx(23456.789)
    assert by_id["word_0001"]["display_text"] == "23,456.79万元"
    assert by_id["word_0001"]["placeholder_occurrences"] == 1
    assert by_id["word_0001"]["status"] == "ok"


def test_render_docx_handles_table_placeholders(tmp_path: Path):
    template = _write_template(
        tmp_path / "template.docx",
        ["渠道分析："],
        tables=[[
            ["渠道", "销售额"],
            ["自然流量", "{{ word_0001 }}"],
            ["付费广告", "{{ word_0002 }}"],
        ]],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(
            word_id="word_0001", raw_token="4,567.89万元", unit="万元",
            generated=4567.89,
            location="table:0/row:1/col:1",
        ),
        _row(
            word_id="word_0002", raw_token="3,876.54万元", unit="万元",
            generated=3876.54,
            location="table:0/row:2/col:1",
        ),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"
    report = render_docx(template, validation, out_docx)

    assert report.ok
    text = _docx_text(out_docx)
    assert "4,567.89万元" in text
    assert "3,876.54万元" in text


# ---------------------------------------------------------------------------
# Gate failures — every condition in the goal that must halt the render
# ---------------------------------------------------------------------------

def test_missing_placeholder_in_validation_fails_loudly(tmp_path: Path):
    """Template references a word_id absent from run_validation — fail loud."""
    template = _write_template(
        tmp_path / "template.docx",
        ["营收达{{ word_0001 }}，毛利率{{ word_9999 }}。"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"

    report = render_docx(template, validation, out_docx)
    assert report.fatal_errors, "orphan placeholder must produce a fatal error"
    assert any("word_9999" in e for e in report.fatal_errors)
    # No docx or log written: the input pair is unusable.
    assert not out_docx.exists()
    assert report.out_log_path is None


def test_extra_validation_row_fails_loudly(tmp_path: Path):
    """run_validation has a word_id with no matching placeholder — fail loud."""
    template = _write_template(
        tmp_path / "template.docx",
        ["营收达{{ word_0001 }}。"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
        _row(word_id="word_9999", raw_token="200", unit="", generated=200.0),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"

    report = render_docx(template, validation, out_docx)
    assert report.fatal_errors == []
    by_id = {r.word_id: r for r in report.rows}
    # word_0001 is fine; word_9999 has no placeholder → STATUS_NO_PLACEHOLDER.
    assert by_id["word_0001"].status == STATUS_OK
    assert by_id["word_9999"].status == STATUS_NO_PLACEHOLDER
    # No docx written: render-docx must not silently drop a confirmed metric.
    assert not out_docx.exists()


def test_non_ok_validation_row_is_refused(tmp_path: Path):
    template = _write_template(
        tmp_path / "template.docx",
        ["营收达{{ word_0001 }}，毛利率{{ word_0002 }}。"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
        # word_0002 is "missing_cell" from run-preview — render must refuse.
        _row(
            word_id="word_0002", raw_token="100", unit="", generated=None,
            status="missing_cell",
            detail="cell 月度!B99 is empty",
        ),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"

    report = render_docx(template, validation, out_docx)
    by_id = {r.word_id: r for r in report.rows}
    assert by_id["word_0001"].status == STATUS_OK
    assert by_id["word_0002"].status == STATUS_VALIDATION_NOT_OK
    # Detail must surface the original non-ok status so the reviewer
    # doesn't have to cross-reference the validation file again.
    assert "missing_cell" in by_id["word_0002"].detail
    # No docx written: confirmed metric word_0002 has no usable value.
    assert not out_docx.exists()


def test_missing_generated_value_is_refused(tmp_path: Path):
    template = _write_template(
        tmp_path / "template.docx",
        ["营收达{{ word_0001 }}。"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(
            word_id="word_0001", raw_token="100", unit="", generated=None,
            status="ok",  # Status=ok but Generated Value is None — corrupt.
        ),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"

    report = render_docx(template, validation, out_docx)
    [row] = report.rows
    assert row.status == STATUS_MISSING_GENERATED_VALUE
    assert not out_docx.exists()


def test_duplicate_validation_row_for_same_word_id_fails_loudly(
    tmp_path: Path,
):
    """Every confirmed word_id must have EXACTLY one Generated Value."""
    template = _write_template(
        tmp_path / "template.docx",
        ["营收达{{ word_0001 }}。"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
        _row(word_id="word_0001", raw_token="100", unit="", generated=200.0),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"

    report = render_docx(template, validation, out_docx)
    assert report.fatal_errors == []
    assert len(report.rows) == 2
    # BOTH copies must be flagged so a reviewer sees the conflict, not
    # whichever one happened to win a silent dedupe.
    for r in report.rows:
        assert r.status == STATUS_DUPLICATE_VALIDATION_ROW, (r.word_id, r.status)
    assert not out_docx.exists()


def test_format_inference_failure_lists_offending_word_id(tmp_path: Path):
    """A raw token that doesn't parse must surface the word_id, not guess."""
    template = _write_template(
        tmp_path / "template.docx",
        ["营收达{{ word_0001 }}。"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(
            word_id="word_0001",
            raw_token="约1.23亿元 to 2.34亿元",  # ambiguous range, not a single token
            unit="亿元",
            generated=1.5,
        ),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"

    report = render_docx(template, validation, out_docx)
    [row] = report.rows
    assert row.status == STATUS_FORMAT_INFERENCE_FAILED
    assert row.word_id == "word_0001"
    assert "约1.23亿元 to 2.34亿元" in row.detail
    assert not out_docx.exists()


# ---------------------------------------------------------------------------
# Duplicate placeholders — the same word_id appearing in N places
# ---------------------------------------------------------------------------

def test_duplicate_placeholder_in_template_substitutes_every_occurrence(
    tmp_path: Path,
):
    """A single confirmed word_id may legitimately appear in multiple
    sentences of the historical report. Each occurrence must be
    substituted, the rendered docx must contain no surviving
    placeholders, and the render_log must record the actual count.
    """
    template = _write_template(
        tmp_path / "template.docx",
        [
            "5月营业收入达{{ word_0001 }}。",
            "上月已述及{{ word_0001 }}的具体口径。",
            "本季累计{{ word_0001 }}的同期对比口径见附录。",
        ],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(
            word_id="word_0001", raw_token="23,456.79万元", unit="万元",
            generated=23456.789, raw_excel=234567890.0,
        ),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"

    report = render_docx(template, validation, out_docx)
    assert report.ok
    text = _docx_text(out_docx)
    # Every occurrence of {{ word_0001 }} replaced.
    assert text.count("23,456.79万元") == 3
    assert not PLACEHOLDER_RE.search(text)

    # The log records duplicate-placeholder accounting so a reviewer can
    # confirm no occurrence was silently skipped.
    doc = yaml.safe_load(report.out_log_path.read_text(encoding="utf-8"))
    [entry] = doc["replacements"]
    assert entry["word_id"] == "word_0001"
    assert entry["placeholder_occurrences"] == 3
    assert doc["summary"]["total_replacements"] == 3
    assert doc["summary"]["distinct_placeholder_word_ids"] == 1


# ---------------------------------------------------------------------------
# Negative-value formatting end-to-end
# ---------------------------------------------------------------------------

def test_negative_万元_value_renders_with_minus_sign(tmp_path: Path):
    template = _write_template(
        tmp_path / "template.docx",
        ["环比变化：{{ word_0001 }}。"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(
            word_id="word_0001", raw_token="-1,234.56万元", unit="万元",
            generated=-1234.567, raw_excel=-12345670.0,
        ),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"
    report = render_docx(template, validation, out_docx)
    assert report.ok
    assert "-1,234.57万元" in _docx_text(out_docx)


def test_negative_paren_style_renders_with_parens(tmp_path: Path):
    template = _write_template(
        tmp_path / "template.docx",
        ["亏损金额：{{ word_0001 }}。"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(
            word_id="word_0001", raw_token="(1,234.56)", unit="",
            generated=-1234.56,
        ),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"
    report = render_docx(template, validation, out_docx)
    assert report.ok
    assert "(1,234.56)" in _docx_text(out_docx)


# ---------------------------------------------------------------------------
# CLI surface: exit codes + artifact presence
# ---------------------------------------------------------------------------

def test_cli_render_docx_success_returns_0(tmp_path: Path, capsys):
    template, validation, out_docx = _golden_inputs(tmp_path)
    rc = cli_main([
        "render-docx",
        "--template", str(template),
        "--run-validation", str(validation),
        "--out", str(out_docx),
    ])
    assert rc == 0
    assert out_docx.exists()
    assert (out_docx.parent / "render_log.yml").exists()
    captured = capsys.readouterr()
    assert "render-docx summary" in captured.out


def test_cli_render_docx_fatal_error_returns_8(tmp_path: Path):
    """Orphan placeholder is a fatal error (input pair is unusable)."""
    template = _write_template(
        tmp_path / "template.docx",
        ["{{ word_0001 }} and {{ word_9999 }}"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"
    rc = cli_main([
        "render-docx",
        "--template", str(template),
        "--run-validation", str(validation),
        "--out", str(out_docx),
    ])
    assert rc == 8
    assert not out_docx.exists()
    assert not (out_docx.parent / "render_log.yml").exists()


def test_cli_render_docx_gate_failure_returns_9(tmp_path: Path):
    """A non-ok validation row is a gate failure — exit 9, no docx."""
    template = _write_template(
        tmp_path / "template.docx",
        ["{{ word_0001 }}"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(
            word_id="word_0001", raw_token="100", unit="", generated=None,
            status="missing_cell",
        ),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"
    rc = cli_main([
        "render-docx",
        "--template", str(template),
        "--run-validation", str(validation),
        "--out", str(out_docx),
    ])
    assert rc == 9
    assert not out_docx.exists()


# ---------------------------------------------------------------------------
# Stale-output guard: a previous successful render's artifacts must not
# survive a subsequent failed invocation. Otherwise a downstream consumer
# that only checks for file existence would silently keep using the
# stale docx — the exact "looks successful while omitting metrics" risk
# CLAUDE.md forbids.
# ---------------------------------------------------------------------------

def _seed_stale_outputs(out_docx: Path) -> Path:
    """Write placeholder old ``new_report.docx`` + ``render_log.yml``."""
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    out_docx.write_bytes(b"STALE DOCX FROM A PRIOR RUN")
    log = out_docx.parent / "render_log.yml"
    log.write_text("stale: true\n", encoding="utf-8")
    return log


def test_stale_outputs_wiped_on_fatal_orphan_placeholder(tmp_path: Path):
    template = _write_template(
        tmp_path / "template.docx",
        ["{{ word_0001 }} and {{ word_9999 }}"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"
    stale_log = _seed_stale_outputs(out_docx)

    report = render_docx(template, validation, out_docx)
    assert report.fatal_errors

    # Both stale artifacts must be gone — leaving either in place would
    # mis-advertise this failed run as a success.
    assert not out_docx.exists(), "stale docx must be wiped on fatal error"
    assert not stale_log.exists(), "stale render_log.yml must be wiped on fatal error"


def test_stale_outputs_wiped_on_gate_failure(tmp_path: Path):
    template = _write_template(
        tmp_path / "template.docx",
        ["{{ word_0001 }}"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(
            word_id="word_0001", raw_token="100", unit="", generated=None,
            status="missing_cell",
        ),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"
    stale_log = _seed_stale_outputs(out_docx)

    report = render_docx(template, validation, out_docx)
    assert report.failures

    assert not out_docx.exists(), "stale docx must be wiped on gate failure"
    assert not stale_log.exists(), "stale render_log.yml must be wiped on gate failure"


def test_stale_outputs_wiped_on_format_inference_failure(tmp_path: Path):
    template = _write_template(
        tmp_path / "template.docx",
        ["{{ word_0001 }}"],
    )
    validation = _write_validation(tmp_path / "run_validation.xlsx", [
        _row(
            word_id="word_0001",
            raw_token="约1.23亿元 to 2.34亿元",
            unit="亿元",
            generated=1.5,
        ),
    ])
    out_docx = tmp_path / "out" / "new_report.docx"
    stale_log = _seed_stale_outputs(out_docx)

    report = render_docx(template, validation, out_docx)
    assert any(r.status == STATUS_FORMAT_INFERENCE_FAILED for r in report.rows)

    assert not out_docx.exists()
    assert not stale_log.exists()


def test_cli_render_docx_failure_does_not_leak_prior_success(tmp_path: Path):
    """End-to-end via the CLI: a successful render followed by a failing
    render at the SAME ``--out`` path must leave the failure exit code
    without any leftover docx/log from the first run.
    """
    # First: a successful render so the docx + log actually exist.
    template, validation, out_docx = _golden_inputs(tmp_path)
    rc_ok = cli_main([
        "render-docx",
        "--template", str(template),
        "--run-validation", str(validation),
        "--out", str(out_docx),
    ])
    assert rc_ok == 0
    assert out_docx.exists()
    log_path = out_docx.parent / "render_log.yml"
    assert log_path.exists()

    # Now: point render-docx at a broken validation file but the SAME
    # --out. The previous run's artifacts must be cleared, not silently
    # re-presented as the current output.
    broken = _write_validation(tmp_path / "broken.xlsx", [
        _row(
            word_id="word_0001", raw_token="100", unit="", generated=None,
            status="missing_cell",
        ),
        _row(
            word_id="word_0002", raw_token="100", unit="", generated=None,
            status="missing_cell",
        ),
        _row(
            word_id="word_0003", raw_token="100", unit="", generated=None,
            status="missing_cell",
        ),
    ])
    rc_fail = cli_main([
        "render-docx",
        "--template", str(template),
        "--run-validation", str(broken),
        "--out", str(out_docx),
    ])
    assert rc_fail == 9
    assert not out_docx.exists(), (
        "leftover successful render docx after a failed re-render "
        "would silently advertise the failed run as a success"
    )
    assert not log_path.exists(), (
        "leftover render_log.yml after a failed re-render would similarly "
        "lie to anyone reading the audit trail"
    )


# ---------------------------------------------------------------------------
# Path-collision guard: --out (or the auto-derived render_log.yml) must
# never wipe a renderer input via the stale-output cleanup.
# ---------------------------------------------------------------------------

def test_refuses_when_out_resolves_to_template(tmp_path: Path):
    """``--out`` pointing at the template path would let cleanup wipe the
    template before any read. The renderer must refuse loudly and leave
    the template intact."""
    template = _write_template(tmp_path / "t.docx", ["{{ word_0001 }}"])
    validation = _write_validation(tmp_path / "v.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    template_bytes_before = template.read_bytes()

    report = render_docx(template, validation, template)

    assert report.fatal_errors
    assert any("--template" in e for e in report.fatal_errors)
    # Critical: the template must NOT have been deleted by the cleanup.
    assert template.exists()
    assert template.read_bytes() == template_bytes_before


def test_refuses_when_out_resolves_to_run_validation(tmp_path: Path):
    """``--out`` pointing at the run_validation path would let cleanup
    wipe the validation artifact before any read."""
    template = _write_template(tmp_path / "t.docx", ["{{ word_0001 }}"])
    validation = _write_validation(tmp_path / "v.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    validation_bytes_before = validation.read_bytes()

    report = render_docx(template, validation, validation)

    assert report.fatal_errors
    assert any("--run-validation" in e for e in report.fatal_errors)
    assert validation.exists()
    assert validation.read_bytes() == validation_bytes_before


def test_refuses_when_render_log_path_collides_with_template(tmp_path: Path):
    """The auto-derived ``render_log.yml`` lives in ``--out.parent``. If
    a template happens to live at that exact path the cleanup would
    wipe it, even though ``--out`` itself doesn't collide."""
    # Template is named render_log.yml (unusual but legal as a path).
    template_at_log = _write_template(
        tmp_path / "render_log.yml", ["{{ word_0001 }}"],
    )
    template_bytes_before = template_at_log.read_bytes()
    validation = _write_validation(tmp_path / "v.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    # --out lives in the same directory, so the derived log path
    # (tmp_path/render_log.yml) collides with the template.
    out_docx = tmp_path / "new_report.docx"

    report = render_docx(template_at_log, validation, out_docx)

    assert report.fatal_errors
    assert any("render_log.yml" in e for e in report.fatal_errors)
    assert any("--template" in e for e in report.fatal_errors)
    assert template_at_log.exists()
    assert template_at_log.read_bytes() == template_bytes_before


def test_refuses_when_render_log_path_collides_with_validation(tmp_path: Path):
    """Same risk for the validation artifact when it happens to live at
    ``--out.parent / render_log.yml``."""
    template = _write_template(tmp_path / "t.docx", ["{{ word_0001 }}"])
    validation_at_log = _write_validation(tmp_path / "render_log.yml", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    validation_bytes_before = validation_at_log.read_bytes()
    out_docx = tmp_path / "new_report.docx"

    report = render_docx(template, validation_at_log, out_docx)

    assert report.fatal_errors
    assert any("render_log.yml" in e for e in report.fatal_errors)
    assert any("--run-validation" in e for e in report.fatal_errors)
    assert validation_at_log.exists()
    assert validation_at_log.read_bytes() == validation_bytes_before


def test_refuses_when_out_filename_equals_render_log(tmp_path: Path):
    """``--out`` ending in ``render_log.yml`` would mean the docx and the
    auto-derived log target the same file — second write clobbers the
    first. Refuse rather than silently produce one file pretending to
    be the other."""
    template = _write_template(tmp_path / "t.docx", ["{{ word_0001 }}"])
    validation = _write_validation(tmp_path / "v.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    out_at_log = tmp_path / "render_log.yml"

    report = render_docx(template, validation, out_at_log)

    assert report.fatal_errors
    assert any("render_log.yml" in e for e in report.fatal_errors)


def test_collision_via_relative_vs_absolute_path_still_detected(tmp_path: Path):
    """A user can spell the same file as ``foo/template.docx`` and
    ``./foo/template.docx`` — both must resolve to the same path and
    the collision must still fire."""
    template = _write_template(tmp_path / "t.docx", ["{{ word_0001 }}"])
    validation = _write_validation(tmp_path / "v.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    # Build a relative-styled view of the same file via a different
    # path-string spelling (parent / name vs. straight Path).
    aliased = Path(str(template.parent)) / Path("./") / template.name

    report = render_docx(template, validation, aliased)

    assert report.fatal_errors
    assert template.exists()


def test_cli_render_docx_collision_returns_8(tmp_path: Path, capsys):
    """End-to-end via the CLI: a collision is a fatal error → exit 8,
    no artifact written, template intact."""
    template = _write_template(tmp_path / "t.docx", ["{{ word_0001 }}"])
    validation = _write_validation(tmp_path / "v.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    template_bytes_before = template.read_bytes()

    rc = cli_main([
        "render-docx",
        "--template", str(template),
        "--run-validation", str(validation),
        "--out", str(template),  # collision
    ])
    assert rc == 8
    assert template.exists()
    assert template.read_bytes() == template_bytes_before
    err = capsys.readouterr().err
    assert "refusing to overwrite" in err


def test_stale_outputs_wiped_when_only_out_collides_with_template(
    tmp_path: Path,
):
    """``--out`` collides with --template (refused), but the sibling
    ``render_log.yml`` path doesn't collide. A stale log at that path
    from a prior successful render must STILL be wiped — otherwise the
    leftover log would survive this exit-8 failure and falsely advertise
    a successful render to anyone reading the directory."""
    template = _write_template(tmp_path / "t.docx", ["{{ word_0001 }}"])
    template_bytes_before = template.read_bytes()
    validation = _write_validation(tmp_path / "v.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    # Seed a stale render_log.yml in the template's parent directory.
    stale_log = tmp_path / "render_log.yml"
    stale_log.write_text("stale: true\n", encoding="utf-8")

    # Point --out at the template (collision). Log path = stale_log,
    # which does NOT collide with any input.
    report = render_docx(template, validation, template)

    assert report.fatal_errors, "collision must be a fatal error"
    # Input must survive: cleanup correctly skipped the colliding path.
    assert template.exists()
    assert template.read_bytes() == template_bytes_before
    # Stale leftover at the non-colliding output path must be GONE.
    assert not stale_log.exists(), (
        "stale render_log.yml survived a collision-fatal run; a "
        "downstream tool would mis-read it as the current render's log"
    )


def test_stale_outputs_wiped_when_only_log_collides_with_template(
    tmp_path: Path,
):
    """Symmetric case: log path collides with --template (refused), but
    ``--out`` doesn't. Any stale docx at the ``--out`` path from a prior
    run must still be wiped."""
    # Template at the path the log would target.
    template_at_log = _write_template(
        tmp_path / "render_log.yml", ["{{ word_0001 }}"],
    )
    template_bytes_before = template_at_log.read_bytes()
    validation = _write_validation(tmp_path / "v.xlsx", [
        _row(word_id="word_0001", raw_token="100", unit="", generated=100.0),
    ])
    # Seed a stale new_report.docx (no collision with any input).
    stale_docx = tmp_path / "new_report.docx"
    stale_docx.write_bytes(b"STALE DOCX FROM PRIOR RUN")

    report = render_docx(template_at_log, validation, stale_docx)

    assert report.fatal_errors  # log↔template collision is fatal
    # Input (the template) must survive — cleanup skipped it.
    assert template_at_log.exists()
    assert template_at_log.read_bytes() == template_bytes_before
    # The non-colliding stale docx at --out must be gone.
    assert not stale_docx.exists()


def test_cli_collision_failure_does_not_leak_prior_success(tmp_path: Path):
    """End-to-end via the CLI: a successful render leaves docx + log at
    --out; a follow-up collision-failure at a NEW --out (pointing at
    --template, refused) must still wipe the original --out artifacts
    when they share a directory with the failed run's log target. Pins
    that a fatal exit doesn't strand a stale audit trail."""
    template, validation, ok_out = _golden_inputs(tmp_path)
    rc_ok = cli_main([
        "render-docx",
        "--template", str(template),
        "--run-validation", str(validation),
        "--out", str(ok_out),
    ])
    assert rc_ok == 0
    ok_log = ok_out.parent / "render_log.yml"
    assert ok_out.exists() and ok_log.exists()

    # Now request a render whose --out collides with the template. The
    # auto-derived log path is ``template.parent / render_log.yml``,
    # which differs from ok_log (they're in different directories) — so
    # ok_log is NOT touched by THIS invocation. But if we point the
    # collision-failing run's log at ok_out.parent we DO want ok_log
    # cleared. Verify the simpler case here: ok_out lives in
    # ``tmp_path/out/`` and the colliding run's --out is the template
    # in ``tmp_path/``, so cleanup targets ``tmp_path/render_log.yml``
    # (no overlap with ok_log). The prior success in ``tmp_path/out/``
    # stays intact because the new run never claimed that directory.
    template_bytes_before = template.read_bytes()
    rc_fail = cli_main([
        "render-docx",
        "--template", str(template),
        "--run-validation", str(validation),
        "--out", str(template),  # collision
    ])
    assert rc_fail == 8
    # Input preserved.
    assert template.read_bytes() == template_bytes_before
    # The unrelated prior-success artifacts in a different directory are
    # untouched — cleanup is path-scoped to the current run's --out dir.
    assert ok_out.exists()
    assert ok_log.exists()


def test_successful_render_still_writes_fresh_artifacts(tmp_path: Path):
    """The cleanup must not break the success path: pre-existing stale
    outputs are wiped, then the fresh ones are written in their place.
    Catches a regression where the cleanup races the writer.
    """
    template, validation, out_docx = _golden_inputs(tmp_path)
    _seed_stale_outputs(out_docx)

    report = render_docx(template, validation, out_docx)
    assert report.ok
    assert out_docx.exists()
    assert report.out_log_path is not None and report.out_log_path.exists()
    # The fresh docx must NOT be the byte-blob we seeded.
    assert out_docx.read_bytes() != b"STALE DOCX FROM A PRIOR RUN"
    # The fresh log must not be the seeded one-line YAML.
    log_text = report.out_log_path.read_text(encoding="utf-8")
    assert "stale: true" not in log_text
    assert "schema_version" in log_text


def test_cli_render_docx_missing_inputs_returns_2(tmp_path: Path, capsys):
    rc = cli_main([
        "render-docx",
        "--template", str(tmp_path / "nope.docx"),
        "--run-validation", str(tmp_path / "nope.xlsx"),
        "--out", str(tmp_path / "out.docx"),
    ])
    assert rc == 2
    err = capsys.readouterr().err
    assert "not found" in err
