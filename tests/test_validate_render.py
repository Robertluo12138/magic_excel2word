"""Tests for the ``validate-render`` final-artifact cross-check.

The validator's contract: re-read the three files render-docx + run-preview
together write (``new_report.docx``, ``render_log.yml``, and
``run_validation.xlsx``) and prove they tell the same story about every
rendered ``word_id``. Two flavors of coverage here:

* A passing path that runs ``render-docx`` end-to-end on a tiny synthetic
  template + validation, so the emitter and the validator agree.
* A series of *targeted corruptions* — each one tweaks exactly one row
  or one cell in a single artifact so the failure cause is unambiguous.
  If a future refactor introduces a different inconsistency, the test
  output should point straight at it rather than swallowing it.

No real data, no LLM, no GUI, no Microsoft Word automation.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import List, Optional, Tuple

import docx
import openpyxl
import pytest
import yaml

from src.main import main as cli_main
from src.render_validator import (
    PLACEHOLDER_RE,
    validate_render,
)
from src.renderer import render_docx


# ---------------------------------------------------------------------------
# Fixture helpers — mirror the shapes ``render-docx`` writes in real runs.
# ---------------------------------------------------------------------------

_RUN_VALIDATION_HEADERS = [
    "Word ID",
    "Word Location",
    "Word Context",
    "Word Raw Token",
    "Word Unit",
    "Source Sheet",
    "Source Cell",
    "Raw Excel Value",
    "Generated Value",
    "Transform Interpretation",
    "Confidence",
    "Status",
    "Detail",
]


def _write_template(path: Path, paragraphs: List[str]) -> Path:
    doc = docx.Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _row(
    *,
    word_id: str,
    raw_token: str,
    unit: str,
    generated: Optional[float],
    raw_excel: float = 0.0,
    sheet: str = "月度",
    cell: str = "B2",
    status: str = "ok",
    detail: str = "",
    location: str = "paragraph:0",
    context: str = "",
    transform: str = "as_written",
    confidence: str = "HIGH",
) -> List:
    """Return one ``run_validation.xlsx`` row in the canonical column order."""
    return [
        word_id, location, context, raw_token, unit,
        sheet, cell, raw_excel, generated,
        transform, confidence, status, detail,
    ]


def _write_validation(path: Path, rows: List[List]) -> Path:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "run_validation"
    ws.append(_RUN_VALIDATION_HEADERS)
    for r in rows:
        ws.append(r)
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))
    return path


def _load_yaml(path: Path) -> dict:
    return yaml.safe_load(path.read_text(encoding="utf-8"))


def _dump_yaml(data: dict, path: Path) -> None:
    path.write_text(
        yaml.safe_dump(data, allow_unicode=True, sort_keys=False, width=1000),
        encoding="utf-8",
    )


def _codes(report) -> set:
    return {i.code for i in report.issues}


def _golden_render(
    tmp_path: Path,
) -> Tuple[Path, Path, Path]:
    """Run render-docx end-to-end against a tiny synthetic pair.

    Returns ``(docx, render_log, run_validation)`` — the three inputs
    ``validate-render`` consumes. Every test below starts here and may
    then mutate one artifact to exercise a specific failure.
    """
    template = _write_template(
        tmp_path / "template.docx",
        [
            "5月营业收入达{{ word_0001 }}，同比增长{{ word_0002 }}。",
            "毛利率{{ word_0003 }}。",
        ],
    )
    validation = _write_validation(
        tmp_path / "run_validation.xlsx",
        [
            _row(
                word_id="word_0001",
                raw_token="23,456.79万元", unit="万元",
                generated=23456.789, raw_excel=234567890.0,
                sheet="月度", cell="B2",
            ),
            _row(
                word_id="word_0002",
                raw_token="25.00%", unit="%",
                generated=24.95, raw_excel=0.2495,
                sheet="月度", cell="C2",
            ),
            _row(
                word_id="word_0003",
                raw_token="36.20%", unit="%",
                generated=36.20, raw_excel=0.3620,
                sheet="月度", cell="C3",
            ),
        ],
    )
    out_docx = tmp_path / "out" / "new_report.docx"
    report = render_docx(template, validation, out_docx)
    assert report.ok, (
        "golden render must succeed before any validate-render coverage; got "
        f"failures={[(r.word_id, r.status, r.detail) for r in report.failures]} "
        f"fatal={report.fatal_errors}"
    )
    log_path = out_docx.parent / "render_log.yml"
    assert log_path.exists()
    return out_docx, log_path, validation


# ---------------------------------------------------------------------------
# Happy path
# ---------------------------------------------------------------------------

def test_validate_render_passes_on_clean_render(tmp_path: Path):
    docx_path, log_path, validation = _golden_render(tmp_path)
    report = validate_render(docx_path, log_path, validation)
    assert report.ok, (
        "clean render must validate; got issues: "
        f"{[(i.code, i.message) for i in report.issues]}"
    )


def test_cli_validate_render_passes_on_clean_render(
    tmp_path: Path, capsys
):
    docx_path, log_path, validation = _golden_render(tmp_path)
    rc = cli_main([
        "validate-render",
        "--docx", str(docx_path),
        "--render-log", str(log_path),
        "--run-validation", str(validation),
    ])
    assert rc == 0
    assert "OK" in capsys.readouterr().out


def test_cli_validate_render_returns_2_for_missing_inputs(
    tmp_path: Path, capsys
):
    rc = cli_main([
        "validate-render",
        "--docx", str(tmp_path / "nope.docx"),
        "--render-log", str(tmp_path / "nope.yml"),
        "--run-validation", str(tmp_path / "nope.xlsx"),
    ])
    assert rc == 2
    assert "not found" in capsys.readouterr().err


# ---------------------------------------------------------------------------
# Leftover placeholder — the docx still carries an un-substituted token.
# ---------------------------------------------------------------------------

def test_leftover_placeholder_in_docx_is_flagged(tmp_path: Path):
    docx_path, log_path, validation = _golden_render(tmp_path)
    # Inject a stray placeholder into the rendered docx. This simulates
    # the worst case: render-docx claimed success but a token survived.
    doc = docx.Document(str(docx_path))
    doc.paragraphs[0].text = (
        doc.paragraphs[0].text + " {{ word_4242 }}"
    )
    doc.save(str(docx_path))

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "leftover_placeholder" in _codes(report)
    msg = next(
        i.message for i in report.issues if i.code == "leftover_placeholder"
    )
    assert "word_4242" in msg


def test_cli_validate_render_returns_10_on_leftover_placeholder(
    tmp_path: Path, capsys
):
    docx_path, log_path, validation = _golden_render(tmp_path)
    doc = docx.Document(str(docx_path))
    doc.paragraphs[0].text = (
        doc.paragraphs[0].text + " {{ word_4242 }}"
    )
    doc.save(str(docx_path))

    rc = cli_main([
        "validate-render",
        "--docx", str(docx_path),
        "--render-log", str(log_path),
        "--run-validation", str(validation),
    ])
    assert rc == 10
    err = capsys.readouterr().err
    assert "FAILED" in err
    assert "leftover_placeholder" in err


# ---------------------------------------------------------------------------
# Missing log row — run_validation declares a word_id the log forgot.
# ---------------------------------------------------------------------------

def test_missing_log_row_for_run_validation_word_id(tmp_path: Path):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    # Drop word_0003 from the log even though run_validation still lists
    # it (and the docx still rendered it).
    log["replacements"] = [
        e for e in log["replacements"] if e.get("word_id") != "word_0003"
    ]
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "missing_render_log_row" in _codes(report)
    msg = next(
        i.message for i in report.issues
        if i.code == "missing_render_log_row"
    )
    assert "word_0003" in msg


# ---------------------------------------------------------------------------
# Extra log row — the log claims a word_id run_validation never confirmed.
# ---------------------------------------------------------------------------

def test_extra_log_row_with_no_run_validation_match(tmp_path: Path):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    log["replacements"].append({
        "word_id": "word_9999",
        "location": "paragraph:99",
        "source_sheet": "月度",
        "source_cell": "Z99",
        "raw_excel_value": 1.0,
        "generated_value": 1.0,
        "raw_token": "1",
        "unit": "",
        "display_text": "1",
        "placeholder_occurrences": 1,
        "status": "ok",
        "detail": "",
    })
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "extra_render_log_row" in _codes(report)
    msg = next(
        i.message for i in report.issues if i.code == "extra_render_log_row"
    )
    assert "word_9999" in msg


# ---------------------------------------------------------------------------
# Duplicate word_id — either artifact.
# ---------------------------------------------------------------------------

def test_duplicate_word_id_in_render_log_is_flagged(tmp_path: Path):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    # Duplicate word_0001 in the log: same payload, just appended again.
    dup = dict(next(e for e in log["replacements"] if e["word_id"] == "word_0001"))
    log["replacements"].append(dup)
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "render_log_duplicate_word_id" in _codes(report)
    msg = next(
        i.message for i in report.issues
        if i.code == "render_log_duplicate_word_id"
    )
    assert "word_0001" in msg


def test_duplicate_word_id_in_run_validation_is_flagged(tmp_path: Path):
    docx_path, log_path, validation = _golden_render(tmp_path)
    # Append a duplicate row for word_0001 to the run_validation xlsx.
    wb = openpyxl.load_workbook(str(validation))
    ws = wb.active
    ws.append(_row(
        word_id="word_0001",
        raw_token="23,456.79万元", unit="万元",
        generated=23456.789, raw_excel=234567890.0,
        sheet="月度", cell="B2",
    ))
    wb.save(str(validation))

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "run_validation_duplicate_word_id" in _codes(report)


# ---------------------------------------------------------------------------
# Non-ok run_validation row — refuse to bless the rendered docx.
# ---------------------------------------------------------------------------

def test_non_ok_run_validation_row_is_refused(tmp_path: Path):
    docx_path, log_path, validation = _golden_render(tmp_path)
    # Hand-edit one run_validation row to a non-ok status. The docx and
    # the log still claim success — exactly the drift the validator
    # exists to catch.
    wb = openpyxl.load_workbook(str(validation))
    ws = wb.active
    # Find the Status column index by header.
    headers = [c.value for c in ws[1]]
    status_col = headers.index("Status") + 1
    wid_col = headers.index("Word ID") + 1
    for row_idx in range(2, ws.max_row + 1):
        if ws.cell(row=row_idx, column=wid_col).value == "word_0002":
            ws.cell(row=row_idx, column=status_col, value="missing_cell")
            break
    wb.save(str(validation))

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "run_validation_status_not_ok" in _codes(report)
    msg = next(
        i.message for i in report.issues
        if i.code == "run_validation_status_not_ok"
    )
    assert "word_0002" in msg
    assert "missing_cell" in msg


def test_non_ok_render_log_row_is_refused(tmp_path: Path):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0002":
            e["status"] = "format_inference_failed"
            break
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "render_log_status_not_ok" in _codes(report)


# ---------------------------------------------------------------------------
# Zero placeholder_occurrences — a rendered metric was silently skipped.
# ---------------------------------------------------------------------------

def test_zero_placeholder_occurrences_is_flagged(tmp_path: Path):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0001":
            e["placeholder_occurrences"] = 0
            break
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "zero_placeholder_occurrences" in _codes(report)
    msg = next(
        i.message for i in report.issues
        if i.code == "zero_placeholder_occurrences"
    )
    assert "word_0001" in msg


# ---------------------------------------------------------------------------
# Missing source / display fields — audit-trail hole even if docx looks OK.
# ---------------------------------------------------------------------------

@pytest.mark.parametrize(
    "field_name",
    ["generated_value", "source_sheet", "source_cell", "display_text"],
)
def test_missing_required_field_in_render_log_is_flagged(
    tmp_path: Path, field_name: str
):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0001":
            e[field_name] = None
            break
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    codes = _codes(report)
    assert "render_log_missing_field" in codes
    msg = next(
        i.message for i in report.issues
        if i.code == "render_log_missing_field"
    )
    assert "word_0001" in msg
    assert field_name in msg


def test_generated_value_zero_does_not_trip_missing_field_gate(
    tmp_path: Path,
):
    """``0``/``0.0`` are legitimate generated values — the gate must
    only fire on the literal ``None``/empty string, not on a falsy
    numeric zero."""
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0001":
            e["generated_value"] = 0
            break
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert "render_log_missing_field" not in _codes(report)


# ---------------------------------------------------------------------------
# CLI exit-code surface — make sure the non-ok variants land on 10.
# ---------------------------------------------------------------------------

def test_cli_validate_render_returns_10_on_missing_field(
    tmp_path: Path, capsys
):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0001":
            e["source_sheet"] = None
            break
    _dump_yaml(log, log_path)

    rc = cli_main([
        "validate-render",
        "--docx", str(docx_path),
        "--render-log", str(log_path),
        "--run-validation", str(validation),
    ])
    assert rc == 10
    err = capsys.readouterr().err
    assert "render_log_missing_field" in err


# ---------------------------------------------------------------------------
# Audit-value drift between render_log.yml and run_validation.xlsx —
# *presence* of the source/value fields isn't enough; a hand-edited log
# could claim an invented source or generated value while run_validation
# still holds the true one. That's the exact "no invented sources, cells,
# or numbers" risk CLAUDE.md forbids, so the gate must catch it.
# ---------------------------------------------------------------------------

def test_source_cell_drift_between_log_and_validation_is_flagged(
    tmp_path: Path,
):
    """Hand-edit the log's source_cell while leaving run_validation as
    the true source — the validator must refuse to call the audit
    consistent."""
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0001":
            e["source_cell"] = "Z99"  # fabricated cell address
            break
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "audit_value_drift" in _codes(report)
    msg = next(
        i.message for i in report.issues if i.code == "audit_value_drift"
    )
    assert "word_0001" in msg
    assert "source_cell" in msg
    assert "Z99" in msg


def test_source_sheet_drift_between_log_and_validation_is_flagged(
    tmp_path: Path,
):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0002":
            e["source_sheet"] = "伪造表"  # fabricated sheet name
            break
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "audit_value_drift" in _codes(report)
    msg = next(
        i.message for i in report.issues if i.code == "audit_value_drift"
    )
    assert "word_0002" in msg
    assert "source_sheet" in msg


def test_generated_value_drift_between_log_and_validation_is_flagged(
    tmp_path: Path,
):
    """A log that claims a different ``generated_value`` than what
    run-preview confirmed is exactly the "invented number" failure mode
    the validator exists to catch."""
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0001":
            e["generated_value"] = 999.0  # fabricated value
            break
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "audit_value_drift" in _codes(report)
    msg = next(
        i.message for i in report.issues if i.code == "audit_value_drift"
    )
    assert "generated_value" in msg
    assert "999" in msg


def test_raw_excel_value_drift_between_log_and_validation_is_flagged(
    tmp_path: Path,
):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0002":
            e["raw_excel_value"] = 0.5  # was 0.2495
            break
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "audit_value_drift" in _codes(report)
    msg = next(
        i.message for i in report.issues if i.code == "audit_value_drift"
    )
    assert "raw_excel_value" in msg


def test_raw_token_drift_between_log_and_validation_is_flagged(
    tmp_path: Path,
):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0001":
            e["raw_token"] = "99,999.99万元"  # was 23,456.79万元
            break
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "audit_value_drift" in _codes(report)
    msg = next(
        i.message for i in report.issues if i.code == "audit_value_drift"
    )
    assert "raw_token" in msg


def test_unit_drift_between_log_and_validation_is_flagged(tmp_path: Path):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0002":
            e["unit"] = "‰"  # was %
            break
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "audit_value_drift" in _codes(report)
    msg = next(
        i.message for i in report.issues if i.code == "audit_value_drift"
    )
    assert "unit" in msg


def test_float_roundtrip_noise_does_not_trip_drift_gate(tmp_path: Path):
    """A legitimate float-representation difference (e.g. ``0.1 + 0.2``
    vs ``0.3``) must NOT trip the drift check — otherwise the gate is
    useless against true tampering because reviewers will start
    ignoring it. The tolerance is tight enough that any hand-edit a
    reviewer might make (changing a digit) still trips."""
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0001":
            # Add a tiny perturbation well under _FLOAT_ABS_TOL.
            e["generated_value"] = float(e["generated_value"]) + 1e-12
            break
    _dump_yaml(log, log_path)

    report = validate_render(docx_path, log_path, validation)
    # Specifically: the drift code must NOT fire on noise this small.
    assert "audit_value_drift" not in _codes(report), (
        f"float noise tripped the drift gate; issues: "
        f"{[(i.code, i.message) for i in report.issues]}"
    )


def test_cli_validate_render_returns_10_on_audit_value_drift(
    tmp_path: Path, capsys
):
    docx_path, log_path, validation = _golden_render(tmp_path)
    log = _load_yaml(log_path)
    for e in log["replacements"]:
        if e["word_id"] == "word_0001":
            e["source_cell"] = "Z99"
            e["generated_value"] = 999.0
            break
    _dump_yaml(log, log_path)

    rc = cli_main([
        "validate-render",
        "--docx", str(docx_path),
        "--render-log", str(log_path),
        "--run-validation", str(validation),
    ])
    assert rc == 10
    err = capsys.readouterr().err
    assert "audit_value_drift" in err
    assert "word_0001" in err


# ---------------------------------------------------------------------------
# Docx-text drift — the log and validation can agree on what was
# *supposed* to be rendered, yet the docx itself can be hand-edited
# after the fact. The validator must read the docx and refuse to bless
# a number that doesn't actually appear there.
# ---------------------------------------------------------------------------

def _replace_in_docx(path: Path, old: str, new: str) -> None:
    """Replace ``old`` with ``new`` in every paragraph + cell of the docx."""
    d = docx.Document(str(path))
    for p in d.paragraphs:
        if old in p.text:
            p.text = p.text.replace(old, new)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                if old in cell.text:
                    cell.text = cell.text.replace(old, new)
    d.save(str(path))


def test_hand_edited_docx_number_is_flagged(tmp_path: Path):
    """The most dangerous failure mode: log + validation still agree
    on the audited value, but a reviewer (or a downstream tool) has
    swapped the rendered number for something else. The gate must
    catch it — otherwise the docx silently lies about the metric."""
    docx_path, log_path, validation = _golden_render(tmp_path)
    # Swap "23,456.79万元" for a wholly fabricated "99,999.99万元". The
    # log still claims display_text="23,456.79万元" — no other gate
    # would notice.
    _replace_in_docx(docx_path, "23,456.79万元", "99,999.99万元")

    report = validate_render(docx_path, log_path, validation)
    assert not report.ok
    assert "docx_rendered_text_missing" in _codes(report)
    msg = next(
        i.message for i in report.issues
        if i.code == "docx_rendered_text_missing"
    )
    assert "23,456.79万元" in msg
    assert "word_0001" in msg


def test_one_of_multiple_duplicate_occurrences_changed_is_flagged(
    tmp_path: Path,
):
    """Build a template where one word_id is rendered three times; then
    hand-edit ONE of the three occurrences. ``placeholder_occurrences``
    in the log will be 3 but the docx will only contain the display
    text twice — drift must fire even though the other two are
    correct."""
    template = _write_template(
        tmp_path / "template.docx",
        [
            "5月营业收入达{{ word_0001 }}。",
            "上月已述及{{ word_0001 }}的具体口径。",
            "本季累计{{ word_0001 }}的同期对比口径见附录。",
        ],
    )
    validation = _write_validation(
        tmp_path / "run_validation.xlsx",
        [
            _row(
                word_id="word_0001",
                raw_token="23,456.79万元", unit="万元",
                generated=23456.789, raw_excel=234567890.0,
            ),
        ],
    )
    out_docx = tmp_path / "out" / "new_report.docx"
    report = render_docx(template, validation, out_docx)
    assert report.ok
    log_path = out_docx.parent / "render_log.yml"

    # Tamper: change ONE of the three rendered occurrences. Read the
    # current text and rewrite the second paragraph so the others stay
    # intact.
    d = docx.Document(str(out_docx))
    d.paragraphs[1].text = "上月已述及99,999.99万元的具体口径。"
    d.save(str(out_docx))

    report = validate_render(out_docx, log_path, validation)
    assert not report.ok
    assert "docx_rendered_text_missing" in _codes(report)
    msg = next(
        i.message for i in report.issues
        if i.code == "docx_rendered_text_missing"
    )
    assert "expected 3x" in msg
    assert "docx has 2x" in msg


def test_shared_display_text_across_word_ids_does_not_false_positive(
    tmp_path: Path,
):
    """Two word_ids can legitimately render to the same display text
    (e.g. two unrelated metrics that both happen to be ``100元``).
    The check must aggregate per display_text so the gate passes."""
    template = _write_template(
        tmp_path / "template.docx",
        [
            "A组营业收入达{{ word_0001 }}。",
            "B组营业收入达{{ word_0002 }}。",
        ],
    )
    validation = _write_validation(
        tmp_path / "run_validation.xlsx",
        [
            _row(
                word_id="word_0001",
                raw_token="100元", unit="元",
                generated=100.0, raw_excel=100.0,
                sheet="月度", cell="B2",
            ),
            _row(
                word_id="word_0002",
                raw_token="100元", unit="元",
                generated=100.0, raw_excel=100.0,
                sheet="月度", cell="B3",
            ),
        ],
    )
    out_docx = tmp_path / "out" / "new_report.docx"
    report = render_docx(template, validation, out_docx)
    assert report.ok
    log_path = out_docx.parent / "render_log.yml"

    vr = validate_render(out_docx, log_path, validation)
    # Critical: the shared display_text must not trip the drift gate.
    assert "docx_rendered_text_missing" not in _codes(vr), (
        f"shared display_text falsely tripped drift gate; issues: "
        f"{[(i.code, i.message) for i in vr.issues]}"
    )


def test_extra_text_in_docx_does_not_false_positive(tmp_path: Path):
    """A docx may contain narrative text that happens to include the
    display_text again (e.g. a footnote restating the figure). Extra
    occurrences must NOT trip the gate — only under-count does, since
    over-count cannot prove tampering on its own."""
    docx_path, log_path, validation = _golden_render(tmp_path)
    # Append a paragraph that mentions one of the display texts again.
    d = docx.Document(str(docx_path))
    d.add_paragraph("脚注：上述23,456.79万元口径与去年一致。")
    d.save(str(docx_path))

    report = validate_render(docx_path, log_path, validation)
    assert "docx_rendered_text_missing" not in _codes(report)


def test_cli_validate_render_returns_10_on_docx_text_drift(
    tmp_path: Path, capsys
):
    docx_path, log_path, validation = _golden_render(tmp_path)
    _replace_in_docx(docx_path, "24.95%", "99.99%")

    rc = cli_main([
        "validate-render",
        "--docx", str(docx_path),
        "--render-log", str(log_path),
        "--run-validation", str(validation),
    ])
    assert rc == 10
    err = capsys.readouterr().err
    assert "docx_rendered_text_missing" in err
    assert "24.95%" in err


# ---------------------------------------------------------------------------
# Substring-overlap tampering — the subtlest greenlight risk. Two
# display_texts where one is a substring of the other: a docx tampered
# to remove the smaller's occurrence still has the smaller as substring
# of the larger, so a naive ``count()`` would pass. The validator must
# process longest-first so the larger consumes its span before the
# smaller is counted.
# ---------------------------------------------------------------------------

def test_substring_overlap_in_untampered_docx_does_not_false_positive(
    tmp_path: Path,
):
    """Sanity baseline: two display_texts where "100元" is a substring
    of "1100元". The untampered docx legitimately contains both, and
    the gate must pass."""
    template = _write_template(
        tmp_path / "template.docx",
        ["去年{{ word_0001 }}，今年{{ word_0002 }}。"],
    )
    validation = _write_validation(
        tmp_path / "run_validation.xlsx",
        [
            _row(
                word_id="word_0001",
                raw_token="100元", unit="元",
                generated=100.0, raw_excel=100.0,
                sheet="月度", cell="B2",
            ),
            _row(
                word_id="word_0002",
                raw_token="1100元", unit="元",
                generated=1100.0, raw_excel=1100.0,
                sheet="月度", cell="B3",
            ),
        ],
    )
    out_docx = tmp_path / "out" / "new_report.docx"
    rd = render_docx(template, validation, out_docx)
    assert rd.ok
    log_path = out_docx.parent / "render_log.yml"

    vr = validate_render(out_docx, log_path, validation)
    assert "docx_rendered_text_missing" not in _codes(vr), (
        f"untampered substring overlap falsely tripped the drift gate; "
        f"issues: {[(i.code, i.message) for i in vr.issues]}"
    )


def test_substring_overlap_tampering_is_caught(tmp_path: Path):
    """The smoking-gun case Codex flagged: tampering with the smaller
    display_text leaves it as a substring of the larger, so a naive
    ``count()`` would pass. Longest-first replacement must catch it.
    """
    template = _write_template(
        tmp_path / "template.docx",
        ["去年{{ word_0001 }}，今年{{ word_0002 }}。"],
    )
    validation = _write_validation(
        tmp_path / "run_validation.xlsx",
        [
            _row(
                word_id="word_0001",
                raw_token="100元", unit="元",
                generated=100.0, raw_excel=100.0,
                sheet="月度", cell="B2",
            ),
            _row(
                word_id="word_0002",
                raw_token="1100元", unit="元",
                generated=1100.0, raw_excel=1100.0,
                sheet="月度", cell="B3",
            ),
        ],
    )
    out_docx = tmp_path / "out" / "new_report.docx"
    rd = render_docx(template, validation, out_docx)
    assert rd.ok
    log_path = out_docx.parent / "render_log.yml"

    # Tamper: replace the standalone "100元" (word_0001's rendered
    # value) with "999元". The "1100元" elsewhere still contains "100元"
    # as a substring, so a global count for "100元" would still be 1
    # under naive counting. Longest-first must see "1100元" first,
    # consume its span, and then find 0 occurrences of "100元".
    d = docx.Document(str(out_docx))
    # Surgical edit: rewrite the paragraph so only the standalone
    # "100元" changes, leaving "1100元" intact.
    p = d.paragraphs[0]
    assert "去年100元，今年1100元" in p.text, (
        f"unexpected baseline text: {p.text!r}"
    )
    p.text = p.text.replace("去年100元，", "去年999元，")
    d.save(str(out_docx))

    report = validate_render(out_docx, log_path, validation)
    assert not report.ok, (
        "tampered substring-overlap case must fail; "
        f"issues so far: {[(i.code, i.message) for i in report.issues]}"
    )
    assert "docx_rendered_text_missing" in _codes(report)
    msg = next(
        i.message for i in report.issues
        if i.code == "docx_rendered_text_missing"
    )
    assert "100元" in msg
    assert "word_0001" in msg
    # And the longer text MUST NOT be flagged — it's still in the docx.
    assert "1100元" not in msg or "expected 1x but docx has 1x" not in msg


def test_substring_overlap_tampering_of_longer_text_is_also_caught(
    tmp_path: Path,
):
    """Symmetric case: tamper the LONGER display_text. With longest-first
    counting, the longer one is gone (count=0), and the shorter one's
    count is whatever survives — which still equals its expected count.
    The drift fires only on the longer one. This pins that the
    longest-first sweep doesn't accidentally over-credit the shorter
    when the longer was removed."""
    template = _write_template(
        tmp_path / "template.docx",
        ["去年{{ word_0001 }}，今年{{ word_0002 }}。"],
    )
    validation = _write_validation(
        tmp_path / "run_validation.xlsx",
        [
            _row(
                word_id="word_0001",
                raw_token="100元", unit="元",
                generated=100.0, raw_excel=100.0,
                sheet="月度", cell="B2",
            ),
            _row(
                word_id="word_0002",
                raw_token="1100元", unit="元",
                generated=1100.0, raw_excel=1100.0,
                sheet="月度", cell="B3",
            ),
        ],
    )
    out_docx = tmp_path / "out" / "new_report.docx"
    rd = render_docx(template, validation, out_docx)
    assert rd.ok
    log_path = out_docx.parent / "render_log.yml"

    # Tamper: change the longer "1100元" → "9999元", leave standalone
    # "100元" intact.
    d = docx.Document(str(out_docx))
    p = d.paragraphs[0]
    p.text = p.text.replace("今年1100元", "今年9999元")
    d.save(str(out_docx))

    report = validate_render(out_docx, log_path, validation)
    assert "docx_rendered_text_missing" in _codes(report)
    msg = next(
        i.message for i in report.issues
        if i.code == "docx_rendered_text_missing"
    )
    assert "1100元" in msg
    assert "word_0002" in msg
